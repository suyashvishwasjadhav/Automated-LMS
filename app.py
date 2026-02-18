# =============================================================================
# AI-Powered Learning Management System — Main Application
# =============================================================================
# Copyright (c) 2025 Suyash Vishwas Jadhav. All Rights Reserved.
#
# Author      : Suyash Vishwas Jadhav
# Roles       : Developer, Designer, Architect, Idea Creator,
#               Full Stack Integration, Database Design, Frontend,
#               Backend, Security Implementation, Cloud Deployment (GCP),
#               Testing & QA
#
# Team        : Harsh Gawande (Diagram Evaluation)
#               Nidhi Pawar   (Plagiarism Detection)
#
# Unauthorized copying, modification, distribution, or use of this file,
# via any medium, is strictly prohibited without the express written
# permission of the copyright holder.
# =============================================================================

import os
import random
import string
import smtplib
import traceback
from collections import defaultdict
from datetime import datetime, timedelta, timezone
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from difflib import SequenceMatcher
from functools import wraps

from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_from_directory
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_mail import Mail, Message as MailMessage
from flask_socketio import SocketIO, emit, join_room, leave_room
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from sqlalchemy import create_engine, Column, Integer, String, DateTime, Boolean, Text, ForeignKey, Float, BigInteger, inspect, text, func, desc, UniqueConstraint
from sqlalchemy.orm import sessionmaker, scoped_session, relationship, declarative_base
from authlib.integrations.flask_client import OAuth
from dotenv import load_dotenv



# Add these to your existing imports
import json
import io
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.units import inch
from flask import send_file
import csv

# Import evaluation module
from evaluation import (
    extract_handwritten_text,
    evaluate_answer,
    check_name_mismatch,
    check_plagiarism,
    generate_question_mcq,
    generate_answer_mcq,
    detect_plagiarism_realtime
)

import requests
import numpy as np
import pypdf
from docx import Document as DocxDocument
import google.generativeai as genai
from google.cloud import vision
import re
import base64
from gtts import gTTS
try:
    import whisper
    import torch
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False
    print("Warning: OpenAI Whisper library not installed. Install with: pip install openai-whisper")


# Configure Google Gemini for document Q&A
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    embedding_model = genai.GenerativeModel('models/embedding-001')
    chat_model = genai.GenerativeModel('gemini-2.5-flash')

# Configure GROQ API (Free models: llama-3.3-70b-versatile, mixtral-8x7b-32768)
GROQ_API_KEY = os.environ.get('GROQ_API_KEY')
GROQ_API_URL = 'https://api.groq.com/openai/v1/chat/completions'
GROQ_MODEL = 'llama-3.3-70b-versatile'  # Free model, latest 2025

# Configure DeepSeek API (Free models: deepseek-chat, deepseek-reasoner)
DEEPSEEK_API_KEY = os.environ.get('DEEPSEEK_API_KEY')
DEEPSEEK_API_URL = 'https://api.deepseek.com/v1/chat/completions'
DEEPSEEK_MODEL = 'deepseek-chat'  # Free model, latest 2025

# Initialize OpenAI Whisper for multi-language voice transcription
# NOTE: Whisper model loading is disabled to reduce memory usage in production
# For production, consider using Whisper API or Groq's Whisper endpoint instead
# Supported models: tiny, base, small, medium, large, large-v2, large-v3
whisper_model = None
# Disabled heavy model loading - use API instead
# if WHISPER_AVAILABLE:
#     try:
#         device = "cpu"  # Force CPU to avoid FP16 warning
#         # Using 'base' model - good balance. For better accuracy use 'small' or 'medium'
#         whisper_model = whisper.load_model("base", device=device)
#         print("✅ OpenAI Whisper model loaded successfully (multi-language support enabled)")
#     except Exception as e:
#         print(f"⚠️ Whisper model loading error: {e}")
#         whisper_model = None

# Supported languages for Whisper (ISO 639-1 codes)
WHISPER_SUPPORTED_LANGUAGES = {
    'en': 'English', 'es': 'Spanish', 'fr': 'French', 'de': 'German', 'it': 'Italian',
    'pt': 'Portuguese', 'ru': 'Russian', 'ja': 'Japanese', 'ko': 'Korean', 'zh': 'Chinese',
    'ar': 'Arabic', 'hi': 'Hindi', 'tr': 'Turkish', 'pl': 'Polish', 'nl': 'Dutch',
    'sv': 'Swedish', 'id': 'Indonesian', 'vi': 'Vietnamese', 'he': 'Hebrew', 'th': 'Thai',
    'uk': 'Ukrainian', 'cs': 'Czech', 'ro': 'Romanian', 'hu': 'Hungarian', 'fi': 'Finnish',
    'no': 'Norwegian', 'da': 'Danish', 'el': 'Greek', 'bg': 'Bulgarian', 'hr': 'Croatian',
    'sk': 'Slovak', 'sl': 'Slovenian', 'et': 'Estonian', 'lv': 'Latvian', 'lt': 'Lithuanian',
    'mt': 'Maltese', 'ga': 'Irish', 'cy': 'Welsh', 'is': 'Icelandic', 'mk': 'Macedonian',
    'sq': 'Albanian', 'sr': 'Serbian', 'bs': 'Bosnian', 'ca': 'Catalan', 'eu': 'Basque',
    'gl': 'Galician', 'fa': 'Persian', 'ur': 'Urdu', 'bn': 'Bengali', 'ta': 'Tamil',
    'te': 'Telugu', 'ml': 'Malayalam', 'kn': 'Kannada', 'gu': 'Gujarati', 'pa': 'Punjabi',
    'ne': 'Nepali', 'si': 'Sinhala', 'my': 'Myanmar', 'km': 'Khmer', 'lo': 'Lao',
    'ka': 'Georgian', 'am': 'Amharic', 'sw': 'Swahili', 'af': 'Afrikaans', 'zu': 'Zulu',
    'xh': 'Xhosa', 'yo': 'Yoruba', 'ig': 'Igbo', 'ha': 'Hausa', 'so': 'Somali'
}

from demo import demo_bp

# Import security module
from security import (
    InputValidator, EncryptionService, RateLimiter, SecurityHeaders,
    get_sanitized_form, get_sanitized_json, get_sanitized_args, secure_input
)

# Import plagiarism detection module
from plagiarism_detector import analyze_submission_text

# Import MCQ generation module
from mcq import (
    allowed_file, extract_text_from_file,
    generate_mcqs_from_topic, generate_mcqs_from_document
)

load_dotenv()

def timeago(dt, now=None):
    if now is None:
        now = datetime.now(timezone.utc)
    diff = now - dt
    
    seconds = diff.total_seconds()
    if seconds < 60:
        return "just now"
    elif seconds < 3600:
        minutes = int(seconds / 60)
        return f"{minutes} minute{'s' if minutes != 1 else ''} ago"
    elif seconds < 86400:
        hours = int(seconds / 3600)
        return f"{hours} hour{'s' if hours != 1 else ''} ago"
    elif seconds < 604800:
        days = int(seconds / 86400)
        return f"{days} day{'s' if days != 1 else ''} ago"
    elif seconds < 2592000:
        weeks = int(seconds / 604800)
        return f"{weeks} week{'s' if weeks != 1 else ''} ago"
    elif seconds < 31536000:
        months = int(seconds / 2592000)
        return f"{months} month{'s' if months != 1 else ''} ago"
    else:
        years = int(seconds / 31536000)
        return f"{years} year{'s' if years != 1 else ''} ago"

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SESSION_SECRET', 'your-secret-key-here')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

# Session Configuration for Security
app.config['SESSION_COOKIE_SECURE'] = False  # Set to True in production with HTTPS
app.config['SESSION_COOKIE_HTTPONLY'] = True  # Prevent JavaScript access to session cookie
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # CSRF protection
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=30)  # 30-minute session timeout

# CSRF Protection
def _dummy_generate_csrf():
    return ''

generate_csrf = _dummy_generate_csrf
csrf = None

try:
    from flask_wtf.csrf import CSRFProtect, generate_csrf as _flask_generate_csrf
    generate_csrf = _flask_generate_csrf
    csrf = CSRFProtect(app)
    app.config['WTF_CSRF_ENABLED'] = True
    app.config['WTF_CSRF_TIME_LIMIT'] = 3600
    # Allow CSRF token in JSON requests via headers
    app.config['WTF_CSRF_HEADERS'] = ['X-CSRFToken', 'X-CSRF-Token']
    print("✅ CSRF protection enabled")
except ImportError:
    print("⚠️  Flask-WTF not installed. CSRF protection disabled. Install with: pip install Flask-WTF")
    csrf = None
except Exception as e:
    print(f"⚠️  CSRF setup error: {e}")
    csrf = None

@app.context_processor
def inject_csrf_token():
    def get_csrf_token():
        try:
            return generate_csrf()
        except Exception:
            return ''
    return dict(csrf_token=get_csrf_token)

# Rate Limiting
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    limiter = Limiter(
        app=app,
        key_func=get_remote_address,
        default_limits=["200 per day", "50 per hour"],
        storage_uri="memory://"
    )
    print("✅ Rate limiting enabled")
except ImportError:
    print("⚠️  Flask-Limiter not installed. Rate limiting disabled. Install with: pip install Flask-Limiter")
    limiter = None

# Helper function for conditional rate limiting
def rate_limit_if_available(limit_str):
    """Helper to apply rate limiting only if limiter is available"""
    def decorator(f):
        if limiter:
            return limiter.limit(limit_str)(f)
        return f
    return decorator

# Security Headers Middleware
@app.after_request
def add_security_headers(response):
    """Add security headers to all responses"""
    # Skip security headers for static files to avoid blocking resources
    if request.endpoint and ('static' in str(request.endpoint) or request.path.startswith('/static')):
        return response
    
    headers = SecurityHeaders.get_headers()
    for header, value in headers.items():
        response.headers[header] = value
    return response

# Session Timeout and Validation Middleware
@app.before_request
def check_session_timeout():
    """Check for session timeout and validate session ID"""
    # Skip for static files and public routes
    if request.endpoint and ('static' in str(request.endpoint) or request.path.startswith('/static')):
        return
    
    # Skip for login/register/logout routes
    public_routes = ['login', 'register', 'google_login', 'google_callback', 'logout', 'index', 'demo.demo_page']
    if request.endpoint in public_routes:
        return
    
    if current_user.is_authenticated:
        # Make session permanent to enable timeout
        session.permanent = True
        
        # Check if user's session ID matches the one in database
        if hasattr(current_user, 'session_id') and current_user.session_id:
            current_session_id = session.get('_id')
            if current_session_id != current_user.session_id:
                # Session ID mismatch - user logged in elsewhere
                logout_user()
                session.clear()
                flash('You have been logged out because you logged in from another location.', 'warning')
                return redirect(url_for('login'))
        
        # Check for session timeout (30 minutes of inactivity)
        if hasattr(current_user, 'last_activity') and current_user.last_activity:
            # Ensure both datetimes are timezone-aware for comparison
            last_activity = current_user.last_activity
            if last_activity.tzinfo is None:
                # Make timezone-naive datetime timezone-aware (assume UTC)
                last_activity = last_activity.replace(tzinfo=timezone.utc)
            
            time_since_activity = datetime.now(timezone.utc) - last_activity
            if time_since_activity > timedelta(minutes=30):
                # Session expired due to inactivity
                logout_user()
                session.clear()
                flash('Your session has expired due to inactivity. Please log in again.', 'info')
                return redirect(url_for('login'))
        
        # Update last activity time
        try:
            current_user.last_activity = datetime.now(timezone.utc)
            db_session.commit()
        except Exception as e:
            print(f"Error updating last_activity: {e}")
            db_session.rollback()


# Register custom filters
app.jinja_env.filters['timeago'] = timeago

# Add safe HTML filter for when user content needs to be rendered (already sanitized)
def safe_html_filter(value):
    """Mark HTML as safe for rendering (use only for already-sanitized content)"""
    from markupsafe import Markup
    if value is None:
        return ''
    return Markup(str(value))

app.jinja_env.filters['safe_html'] = safe_html_filter

# Server Configuration
LOCAL_SERVER_NAME = '127.0.0.1:5000'
app.config.update(
    SERVER_NAME=None,  # Don't set SERVER_NAME to avoid routing issues
    APPLICATION_ROOT='/',
    PREFERRED_URL_SCHEME='http'
)

SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
SMTP_EMAIL = os.environ.get("SMTP_EMAIL")
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD")
DATABASE_URL = os.environ.get("DATABASE_URL")
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET")
GOOGLE_DISCOVERY_URL = "https://accounts.google.com/.well-known/openid-configuration"

app.config['MAIL_SERVER'] = SMTP_SERVER
app.config['MAIL_PORT'] = SMTP_PORT
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = SMTP_EMAIL
app.config['MAIL_PASSWORD'] = SMTP_PASSWORD

UPLOAD_FOLDER = 'static/uploads'
LOGO_FOLDER = 'static/logos'
PHOTO_FOLDER = 'static/photos'
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'xls', 'xlsx', 'csv', 'jpg', 'jpeg', 'png', 'mp4', 'webm'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['LOGO_FOLDER'] = LOGO_FOLDER
app.config['PHOTO_FOLDER'] = PHOTO_FOLDER


# Add after existing folder configurations
ASSIGNMENT_UPLOAD_FOLDER = 'static/assignment_uploads'
app.config['ASSIGNMENT_UPLOAD_FOLDER'] = ASSIGNMENT_UPLOAD_FOLDER

# Create folder if it doesn't exist
os.makedirs(ASSIGNMENT_UPLOAD_FOLDER, exist_ok=True)

mail = Mail(app)
# Use threading mode to avoid selector/kevent issues on some macOS/Python setups
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=GOOGLE_CLIENT_ID,
    client_secret=GOOGLE_CLIENT_SECRET,
    server_metadata_url=GOOGLE_DISCOVERY_URL,
    authorize_url='https://accounts.google.com/o/oauth2/v2/auth',
    access_token_url='https://oauth2.googleapis.com/token',
    api_base_url='https://www.googleapis.com/oauth2/v2/',
    userinfo_endpoint='https://www.googleapis.com/oauth2/v2/userinfo',
    client_kwargs={
        'scope': 'https://www.googleapis.com/auth/userinfo.profile https://www.googleapis.com/auth/userinfo.email openid',
        'prompt': 'select_account consent',
        'access_type': 'offline',
        'include_granted_scopes': 'true',
        'redirect_uri': f'http://{LOCAL_SERVER_NAME}/auth/google/callback'
    }
)


# Register demo blueprint
app.register_blueprint(demo_bp)

# Security helper functions
def safe_form_get(key: str, default=None, sanitize: bool = True) -> str:
    """Safely get and sanitize form data"""
    value = request.form.get(key, default)
    if value and sanitize:
        try:
            return InputValidator.sanitize_input(value)
        except ValueError:
            return default or ""
    return value or default or ""

def safe_json_get(key: str = None, default=None):
    """Safely get and sanitize JSON data"""
    if request.is_json:
        data = request.get_json() or {}
        if key:
            value = data.get(key, default)
            if isinstance(value, str):
                try:
                    return InputValidator.sanitize_input(value)
                except ValueError:
                    return default
            return value
        # Sanitize all string values in the dict
        sanitized = {}
        for k, v in data.items():
            if isinstance(v, str):
                try:
                    sanitized[k] = InputValidator.sanitize_input(v)
                except ValueError:
                    sanitized[k] = v
            else:
                sanitized[k] = v
        return sanitized
    return default

def safe_args_get(key: str, default=None) -> str:
    """Safely get and sanitize query arguments"""
    value = request.args.get(key, default)
    if value:
        try:
            return InputValidator.sanitize_input(value)
        except ValueError:
            return default or ""
    return value or default or ""

engine = create_engine(DATABASE_URL)
db_session = scoped_session(sessionmaker(bind=engine))
Base = declarative_base()
Base.query = db_session.query_property()

class User(Base, UserMixin):
    __tablename__ = 'users'
    id = Column(Integer, primary_key=True)
    username = Column(String(100), unique=True, nullable=False)
    email = Column(String(120), unique=True, nullable=False)
    password_hash = Column(String(255))
    full_name = Column(String(200))
    role = Column(String(20), nullable=False)
    organization_name = Column(String(200))
    organization_logo = Column(String(255))
    profile_photo = Column(String(255))
    bio = Column(Text)
    google_id = Column(String(200))
    is_profile_complete = Column(Boolean, default=False)
    is_online = Column(Boolean, default=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    last_seen = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    # Session management fields
    session_id = Column(String(255))  # Current active session ID
    last_activity = Column(DateTime, default=lambda: datetime.now(timezone.utc))  # For session timeout
    
    courses_created = relationship('Course', back_populates='executive', foreign_keys='Course.executive_id')
    courses_as_guide = relationship('Course', back_populates='guide', foreign_keys='Course.guide_id')
    
class Course(Base):
    __tablename__ = 'courses'
    id = Column(Integer, primary_key=True)
    course_name = Column(String(200), nullable=False)
    executive_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    guide_id = Column(Integer, ForeignKey('users.id'))
    organization = Column(String(200))
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    executive = relationship('User', back_populates='courses_created', foreign_keys=[executive_id])
    guide = relationship('User', back_populates='courses_as_guide', foreign_keys=[guide_id])
    learners = relationship('CourseLearner', back_populates='course')
    
class CourseLearner(Base):
    __tablename__ = 'course_learners'
    id = Column(Integer, primary_key=True)
    course_id = Column(Integer, ForeignKey('courses.id'), nullable=False)
    learner_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    
    course = relationship('Course', back_populates='learners')
    learner = relationship('User')
    
class Announcement(Base):
    __tablename__ = 'announcements'
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    title = Column(String(200), nullable=False)
    content = Column(Text)
    file_path = Column(String(255))
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    user = relationship('User')
    
class Upload(Base):
    __tablename__ = 'uploads'
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    course_id = Column(Integer, ForeignKey('courses.id'))
    file_name = Column(String(255), nullable=False)
    file_path = Column(String(255), nullable=False)
    file_type = Column(String(50))
    file_size = Column(Float)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    user = relationship('User')
    course = relationship('Course')
    
class Message(Base):
    __tablename__ = 'messages'
    id = Column(Integer, primary_key=True)
    sender_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    receiver_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    content = Column(Text, nullable=False)
    is_read = Column(Boolean, default=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    sender = relationship('User', foreign_keys=[sender_id])
    receiver = relationship('User', foreign_keys=[receiver_id])
    
class ExecutiveContact(Base):
    __tablename__ = 'executive_contacts'
    id = Column(Integer, primary_key=True)
    executive_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    issue_type = Column(String(50), nullable=False)  # urgent_issue, complaint, suggestions, requirements
    details = Column(Text, nullable=False)
    status = Column(String(20), default='pending')  # pending, in_progress, resolved
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    executive = relationship('User', foreign_keys=[executive_id])

class GuideFeedback(Base):
    __tablename__ = 'guide_feedback'
    id = Column(Integer, primary_key=True)
    guide_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    rating = Column(Integer, nullable=False)  # 1-5 stars
    feedback_text = Column(Text)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    guide = relationship('User', foreign_keys=[guide_id])

class GuideSuggestion(Base):
    __tablename__ = 'guide_suggestions'
    id = Column(Integer, primary_key=True)
    guide_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    suggestion_text = Column(Text, nullable=False)
    status = Column(String(20), default='pending')  # pending, reviewed, implemented
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    guide = relationship('User', foreign_keys=[guide_id])

class LearnerSuggestion(Base):
    __tablename__ = 'learner_suggestions'
    id = Column(Integer, primary_key=True)
    learner_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    suggestion_text = Column(Text, nullable=False)
    status = Column(String(20), default='pending')  # pending, reviewed, implemented
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    learner = relationship('User', foreign_keys=[learner_id])

class OTP(Base):
    __tablename__ = 'otps'
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    otp_code = Column(String(6), nullable=False)
    purpose = Column(String(100))
    expires_at = Column(DateTime, nullable=False)
    is_used = Column(Boolean, default=False)
    
    user = relationship('User')
    
class Issue(Base):
    __tablename__ = 'issues'
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    course_id = Column(Integer, ForeignKey('courses.id'))
    title = Column(String(200), nullable=False)
    description = Column(Text)
    status = Column(String(50), default='Open')
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    user = relationship('User')
    course = relationship('Course')

class DocumentChunk(Base):
    """Store document text chunks with embeddings for RAG"""
    __tablename__ = 'document_chunks'
    id = Column(Integer, primary_key=True)
    upload_id = Column(Integer, ForeignKey('uploads.id'), nullable=False)
    chunk_index = Column(Integer, nullable=False)
    text = Column(Text, nullable=False)
    embedding = Column(Text, nullable=False)  # JSON array of floats
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    upload = relationship('Upload')

class MaterialRead(Base):
    __tablename__ = 'material_reads'
    id = Column(Integer, primary_key=True)
    learner_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    upload_id = Column(Integer, ForeignKey('uploads.id'), nullable=False)
    read_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    learner = relationship('User')
    upload = relationship('Upload')


"""
Add these database models to app.py after the existing models
Place right after the Issue model
"""

class Assignment(Base):
    __tablename__ = 'assignments'
    id = Column(Integer, primary_key=True)
    course_id = Column(Integer, ForeignKey('courses.id'), nullable=False)
    guide_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    name = Column(String(200), nullable=False)
    instructions = Column(Text)
    deadline = Column(DateTime, nullable=False)
    total_marks = Column(Float, nullable=False)
    evaluation_model = Column(String(50), default='gemini')
    scores_visible = Column(Boolean, default=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    course = relationship('Course')
    guide = relationship('User')
    questions = relationship('AssignmentQuestion', back_populates='assignment', cascade='all, delete-orphan')
    submissions = relationship('AssignmentSubmission', back_populates='assignment', cascade='all, delete-orphan')

class AssignmentQuestion(Base):
    __tablename__ = 'assignment_questions'
    id = Column(Integer, primary_key=True)
    assignment_id = Column(Integer, ForeignKey('assignments.id'), nullable=False)
    question_text = Column(Text, nullable=False)
    marks = Column(Float, nullable=False)
    order_num = Column(Integer, nullable=False)
    
    assignment = relationship('Assignment', back_populates='questions')
    answers = relationship('AssignmentAnswer', back_populates='question', cascade='all, delete-orphan')

class AssignmentSubmission(Base):
    __tablename__ = 'assignment_submissions'
    id = Column(Integer, primary_key=True)
    assignment_id = Column(Integer, ForeignKey('assignments.id'), nullable=False)
    learner_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    submitted_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    total_score = Column(Float, default=0)
    is_flagged = Column(Boolean, default=False)
    flag_reason = Column(Text)
    
    assignment = relationship('Assignment', back_populates='submissions')
    learner = relationship('User')
    answers = relationship('AssignmentAnswer', back_populates='submission', cascade='all, delete-orphan')

class AssignmentAnswer(Base):
    __tablename__ = 'assignment_answers'
    id = Column(Integer, primary_key=True)
    submission_id = Column(Integer, ForeignKey('assignment_submissions.id'), nullable=False)
    question_id = Column(Integer, ForeignKey('assignment_questions.id'), nullable=False)
    image_paths = Column(Text)  # JSON array of image paths
    extracted_text = Column(Text)
    relevance_score = Column(Float, default=0)
    grammar_score = Column(Float, default=0)
    size_score = Column(Float, default=0)
    uniqueness_score = Column(Float, default=0)
    total_score = Column(Float, default=0)
    feedback = Column(Text)
    covered_points = Column(Text)  # JSON array
    missing_points = Column(Text)  # JSON array
    word_count = Column(Integer, default=0)
    
    submission = relationship('AssignmentSubmission', back_populates='answers')
    question = relationship('AssignmentQuestion', back_populates='answers')


class CustomScoreboardColumn(Base):
    __tablename__ = 'custom_scoreboard_columns'
    id = Column(Integer, primary_key=True)
    course_id = Column(Integer, ForeignKey('courses.id'), nullable=False)
    name = Column(String(150), nullable=False)
    total_marks = Column(Float, nullable=False, default=0.0)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))

    course = relationship('Course')
    values = relationship('CustomScoreboardValue', back_populates='column', cascade='all, delete-orphan')


class CustomScoreboardValue(Base):
    __tablename__ = 'custom_scoreboard_column_values'
    __table_args__ = (
        UniqueConstraint('column_id', 'learner_id', name='uq_custom_scoreboard_column_learner'),
    )
    id = Column(Integer, primary_key=True)
    column_id = Column(Integer, ForeignKey('custom_scoreboard_columns.id'), nullable=False)
    learner_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    mark = Column(Float, nullable=False, default=0.0)
    updated_at = Column(DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    column = relationship('CustomScoreboardColumn', back_populates='values')
    learner = relationship('User')

"""
Add these new database models to app.py after the Assignment models
"""

class QuestionMCQ(Base):
    __tablename__ = 'question_mcqs'
    id = Column(Integer, primary_key=True)
    question_id = Column(Integer, ForeignKey('assignment_questions.id'), nullable=False)
    mcq_question = Column(Text, nullable=False)
    option_a = Column(Text, nullable=False)
    option_b = Column(Text, nullable=False)
    option_c = Column(Text, nullable=False)
    option_d = Column(Text, nullable=False)
    correct_option = Column(String(1), nullable=False)  # A, B, C, or D
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    question = relationship('AssignmentQuestion')

class Notification(Base):
    __tablename__ = 'notifications'
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    title = Column(String(200), nullable=False)
    message = Column(Text, nullable=False)
    notification_type = Column(String(50), nullable=False)  # assignment_created, submission_received
    related_id = Column(Integer)  # assignment_id or submission_id
    is_read = Column(Boolean, default=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    user = relationship('User')
# Add this after QuestionMCQ model (around line 285)

# Diagram Evaluation Models
class DiagramAssignment(Base):
    __tablename__ = 'diagram_assignments'
    id = Column(Integer, primary_key=True)
    course_id = Column(Integer, ForeignKey('courses.id'), nullable=False)
    topic = Column(String(200), nullable=False)
    description = Column(Text)
    created_by = Column(Integer, ForeignKey('users.id'), nullable=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    course = relationship('Course')
    creator = relationship('User')
    submissions = relationship('DiagramSubmission', back_populates='assignment', cascade='all, delete-orphan')

class DiagramSubmission(Base):
    __tablename__ = 'diagram_submissions'
    id = Column(Integer, primary_key=True)
    assignment_id = Column(Integer, ForeignKey('diagram_assignments.id'), nullable=False)
    learner_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    diagram_path = Column(String(500), nullable=False)
    description = Column(Text)
    evaluation_data = Column(Text)  # JSON string
    total_score = Column(Float, default=0)
    topic_match_score = Column(Float, default=0)
    accuracy_score = Column(Float, default=0)
    structure_score = Column(Float, default=0)
    visual_score = Column(Float, default=0)
    technical_score = Column(Float, default=0)
    is_hand_drawn = Column(Boolean, default=False)
    drawing_type = Column(String(50))
    confidence_hand_drawn = Column(Float, default=0)
    final_marks = Column(Float, default=0)
    marks_status = Column(String(100))
    submitted_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    assignment = relationship('DiagramAssignment', back_populates='submissions')
    learner = relationship('User')

class LearnerMCQResponse(Base):
    __tablename__ = 'learner_mcq_responses'
    id = Column(Integer, primary_key=True)
    answer_id = Column(Integer, ForeignKey('assignment_answers.id'), nullable=False)
    mcq_question = Column(Text, nullable=True)
    option_a = Column(Text, nullable=True)
    option_b = Column(Text, nullable=True)
    option_c = Column(Text, nullable=True)
    option_d = Column(Text, nullable=True)
    correct_option = Column(String(1), nullable=True)
    learner_answer = Column(String(1), nullable=False)  # A, B, C, or D
    is_correct = Column(Boolean, nullable=False, default=False)
    time_taken = Column(Integer, nullable=True)  # seconds
    created_at = Column(DateTime, nullable=True, default=lambda: datetime.now(timezone.utc))
    
    answer = relationship('AssignmentAnswer')


# Update existing PlagiarismMatch model (around line 290)
class PlagiarismMatch(Base):
    __tablename__ = 'plagiarism_matches'
    id = Column(Integer, primary_key=True)
    answer_id = Column(Integer, ForeignKey('assignment_answers.id'), nullable=False)
    matched_answer_id = Column(Integer, ForeignKey('assignment_answers.id'), nullable=False)
    similarity_score = Column(Float, nullable=False)
    status = Column(String(20), default='pending')  # pending, ignored, penalized
    penalty_marks = Column(Float, default=0)
    reviewed_by = Column(Integer, ForeignKey('users.id'))
    reviewed_at = Column(DateTime)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    answer = relationship('AssignmentAnswer', foreign_keys=[answer_id])
    matched_answer = relationship('AssignmentAnswer', foreign_keys=[matched_answer_id])
    reviewer = relationship('User', foreign_keys=[reviewed_by])

class PlagiarismAccuracy(Base):
    """Store plagiarism detection accuracy results for each answer"""
    __tablename__ = 'plagiarism_accuracy'
    id = Column(Integer, primary_key=True)
    answer_id = Column(Integer, ForeignKey('assignment_answers.id'), nullable=False)
    plagiarism_score = Column(Float, nullable=False)  # 0-100
    ai_confidence = Column(Float, nullable=False)  # 0-100
    plagiarism_detected = Column(Boolean, default=False)
    ai_detected = Column(Boolean, default=False)
    details = Column(Text)  # JSON string with matches and indicators
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    answer = relationship('AssignmentAnswer')

class TestMCQ(Base):
    """Store generated test MCQs"""
    __tablename__ = 'test_mcqs'
    id = Column(Integer, primary_key=True)
    guide_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    course_id = Column(Integer, ForeignKey('courses.id'), nullable=True)  # Optional course association
    title = Column(String(200), nullable=False)
    source_type = Column(String(20), nullable=False)  # 'topic' or 'document'
    source_content = Column(Text)  # Topic text or document content preview
    source_filename = Column(String(255))  # Original filename if document
    num_questions = Column(Integer, nullable=False)
    marks_per_question = Column(Float, default=1.0)  # Marks for each question
    total_marks = Column(Float)  # Total marks (num_questions * marks_per_question)
    time_limit_minutes = Column(Integer)  # Time limit in minutes (20, 30, 60, 120, 180, 240, 300)
    difficulty = Column(String(20))  # easy, medium, hard (for topics)
    subcategory = Column(String(200))  # For topics
    option_difficulty = Column(String(20))  # identifiable, hard (for topics)
    mcqs_data = Column(Text)  # JSON string of MCQs
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    is_active = Column(Boolean, default=True)
    
    guide = relationship('User', foreign_keys=[guide_id])
    course = relationship('Course', foreign_keys=[course_id])

class TestMCQAttempt(Base):
    """Store learner attempts on test MCQs"""
    __tablename__ = 'test_mcq_attempts'
    id = Column(Integer, primary_key=True)
    test_mcq_id = Column(Integer, ForeignKey('test_mcqs.id'), nullable=False)
    learner_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    answers = Column(Text)  # JSON string of {question_id: selected_answer}
    score = Column(Float)  # Percentage score
    marks_obtained = Column(Float)  # Marks obtained out of total marks
    total_marks = Column(Float)  # Total marks for the test
    total_questions = Column(Integer)
    correct_answers = Column(Integer)
    started_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    completed_at = Column(DateTime)
    
    test_mcq = relationship('TestMCQ')
    learner = relationship('User', foreign_keys=[learner_id])

# Document Q&A System Models
class DocumentQA(Base):
    """Store documents for Q&A system"""
    __tablename__ = 'document_qa'
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=False)
    course_id = Column(Integer, ForeignKey('courses.id'))
    filename = Column(String(500), nullable=False)
    original_filename = Column(String(500), nullable=False)
    file_type = Column(String(50), nullable=False)
    file_size = Column(BigInteger)
    page_count = Column(Integer)
    extracted_text = Column(Text)
    file_path = Column(String(1000))
    summary = Column(Text)  # Simple summary
    topics = Column(Text)  # JSON array of topics
    upload_time = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    user = relationship('User')
    course = relationship('Course')
    chunks = relationship('DocumentQAChunk', back_populates='document', cascade='all, delete-orphan')
    chat_history = relationship('DocumentQAChat', back_populates='document', cascade='all, delete-orphan')

class DocumentQAChunk(Base):
    """Store document chunks with embeddings for RAG"""
    __tablename__ = 'document_qa_chunks'
    id = Column(Integer, primary_key=True)
    document_id = Column(Integer, ForeignKey('document_qa.id', ondelete='CASCADE'), nullable=False)
    chunk_text = Column(Text, nullable=False)
    chunk_index = Column(Integer, nullable=False)
    embedding = Column(Text)  # JSON array of floats
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    document = relationship('DocumentQA', back_populates='chunks')

class DocumentQAChat(Base):
    """Store chat history for document Q&A"""
    __tablename__ = 'document_qa_chat'
    id = Column(Integer, primary_key=True)
    session_id = Column(String(100), nullable=False)
    document_id = Column(Integer, ForeignKey('document_qa.id', ondelete='CASCADE'), nullable=False)
    user_message = Column(Text, nullable=False)
    bot_response = Column(Text, nullable=False)
    timestamp = Column(DateTime, default=lambda: datetime.now(timezone.utc))
    
    document = relationship('DocumentQA', back_populates='chat_history')

# Update AssignmentAnswer model - add this relationship
# Add to existing AssignmentAnswer class:
"""
    mcq_response = relationship('LearnerMCQResponse', back_populates='answer', uselist=False)
    plagiarism_matches = relationship('PlagiarismMatch', foreign_keys='PlagiarismMatch.answer_id')
"""

# Base.metadata.create_all(engine)

Base.metadata.create_all(engine)

# Ensure plagiarism_matches table has the new columns (for deployments that
# previously created the table without the extended schema). This will run
# at startup and add any missing columns safely using IF NOT EXISTS.
try:
    inspector = inspect(engine)
    if 'plagiarism_matches' in inspector.get_table_names():
        existing_cols = [c['name'] for c in inspector.get_columns('plagiarism_matches')]
        with engine.begin() as conn:
            if 'status' not in existing_cols:
                conn.execute(text("ALTER TABLE plagiarism_matches ADD COLUMN IF NOT EXISTS status VARCHAR(20) DEFAULT 'pending'"))
            if 'penalty_marks' not in existing_cols:
                conn.execute(text("ALTER TABLE plagiarism_matches ADD COLUMN IF NOT EXISTS penalty_marks DOUBLE PRECISION DEFAULT 0"))
            if 'reviewed_by' not in existing_cols:
                conn.execute(text("ALTER TABLE plagiarism_matches ADD COLUMN IF NOT EXISTS reviewed_by INTEGER"))
            if 'reviewed_at' not in existing_cols:
                conn.execute(text("ALTER TABLE plagiarism_matches ADD COLUMN IF NOT EXISTS reviewed_at TIMESTAMP"))
except Exception as e:
    print(f"Schema migration check for plagiarism_matches failed: {e}")

# Register custom filters
app.jinja_env.filters['timeago'] = timeago

@login_manager.user_loader
def load_user(user_id):
    return db_session.get(User, int(user_id))

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_random_password(length=12):
    characters = string.ascii_letters + string.digits + string.punctuation
    return ''.join(random.choice(characters) for i in range(length))

def generate_otp():
    return ''.join(random.choices(string.digits, k=6))

def send_email(to_email, subject, body):
    try:
        msg = MIMEMultipart()
        msg['From'] = SMTP_EMAIL
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'html'))
        
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SMTP_EMAIL, SMTP_PASSWORD)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        print(f"Email error: {e}")
        return False

def check_similar_org_names(org_name):
    existing_orgs = db_session.query(User.organization_name).filter(User.organization_name.isnot(None)).all()
    for existing_org in existing_orgs:
        if existing_org[0]:
            similarity = SequenceMatcher(None, org_name.lower(), existing_org[0].lower()).ratio()
            if similarity > 0.85:
                return True
    return False

def role_required(*roles):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated or current_user.role not in roles:
                flash('Access denied')
                return redirect(url_for('index'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator

@app.route('/')
def index():
    return redirect(url_for('demo.demo_page'))


@app.route('/register', methods=['GET', 'POST'])
@rate_limit_if_available("10 per minute")
def register():
    if current_user.is_authenticated:
        dashboard_url = get_dashboard_url(current_user)
        flash('You are already signed in. Please log out if you need to register a different account.', 'info')
        return redirect(dashboard_url)

    # Pull any Google info from session for pre-filling the form
    google_email = session.get('google_email')
    google_id = session.get('google_id')
    google_name = session.get('google_name')

    if request.method == 'POST':
        # Sanitize and validate inputs
        try:
            full_name = get_sanitized_form('full_name') or request.form.get('full_name', '')
        # If we have a google_email in session, prefer it (user came from Google)
            email_raw = google_email or request.form.get('email', '')
            email = email_raw.strip().lower() if email_raw else ''
            organization_name = get_sanitized_form('organization_name') or request.form.get('organization_name', '')

            # Validate email
            if email and not InputValidator.validate_email(email):
                return jsonify({'success': False, 'message': 'Invalid email format'})

        # If the user connected via Google, we'll auto-generate a password unless one is provided
            provided_password = request.form.get('password', '')
            confirm_password = request.form.get('confirm_password', '')
        except ValueError as e:
            return jsonify({'success': False, 'message': str(e)}), 400

        if not full_name or len(full_name) < 9:
            return jsonify({'success': False, 'message': 'Full name must be at least 9 characters'})

        # If not using Google, password fields are required and must match
        if not google_email:
            if not provided_password:
                return jsonify({'success': False, 'message': 'Password is required'})
            if provided_password != confirm_password:
                return jsonify({'success': False, 'message': 'Passwords do not match'})

        # Prevent duplicate emails
        if db_session.query(User).filter_by(email=email).first():
            return jsonify({'success': False, 'message': 'Email already exists'})

        if check_similar_org_names(organization_name):
            return jsonify({'success': False, 'message': 'Similar organization name already exists'})

        username = email.split('@')[0] + '_' + organization_name.replace(' ', '_')[:10]

        if 'organization_logo' in request.files:
            logo_file = request.files['organization_logo']
            if logo_file and logo_file.filename:
                filename = secure_filename(logo_file.filename)
                logo_path = os.path.join(LOGO_FOLDER, f"{username}_{filename}")
                logo_file.save(logo_path)
            else:
                logo_path = None
        else:
            logo_path = None

        # If registering via Google, generate a password if not provided and mark profile complete
        if google_email:
            password = provided_password or generate_random_password()
            is_profile_complete = True
        else:
            password = provided_password
            is_profile_complete = True

        new_user = User(
            username=username,
            email=email,
            password_hash=generate_password_hash(password, method='pbkdf2:sha256'),
            full_name=full_name,
            role='Executive',
            organization_name=organization_name,
            organization_logo=logo_path,
            google_id=google_id,
            is_profile_complete=is_profile_complete
        )

        db_session.add(new_user)
        db_session.commit()

        # Clean up session Google info so subsequent visits don't reuse it
        session.pop('google_email', None)
        session.pop('google_id', None)
        session.pop('google_name', None)

        return jsonify({'success': True, 'message': 'Registration successful! Please login.'})

    # For GET, render the form with any Google data to pre-fill inputs
    return render_template('register.html', google_email=google_email, google_name=google_name)

@app.route('/login', methods=['GET', 'POST'])
@rate_limit_if_available("5 per minute")
def login():
    if current_user.is_authenticated:
        dashboard_url = get_dashboard_url(current_user)
        already_msg = 'You are already signed in. Please log out before accessing the login form.'
        if request.method == 'POST':
            return jsonify({'success': False, 'message': already_msg}), 400
        flash(already_msg, 'info')
        return redirect(dashboard_url)

    if request.method == 'POST':
        # Get client identifier for rate limiting
        client_id = request.remote_addr or 'unknown'
        
        # Check rate limit
        allowed, remaining = RateLimiter.check_rate_limit(client_id)
        if not allowed:
            return jsonify({
                'success': False, 
                'message': f'Too many login attempts. Please try again in {remaining} seconds.'
            }), 429

        # Sanitize inputs (but don't sanitize password as it may contain special chars)
        try:
            username_or_email_raw = request.form.get('username_or_email', '')
            password_raw = request.form.get('password', '')

            # Sanitize username/email (but allow @ and . for emails)
            username_or_email = InputValidator.sanitize_string(username_or_email_raw).strip()
            password = password_raw.strip()  # Don't sanitize password, just trim

            # Check for injection attempts in username/email
            if InputValidator.check_sql_injection(username_or_email):
                RateLimiter.record_attempt(client_id)
                return jsonify({'success': False, 'message': 'Invalid credentials'}), 400
        except ValueError as e:
            RateLimiter.record_attempt(client_id)
            return jsonify({'success': False, 'message': 'Invalid input'}), 400

        print(f"LOGIN ATTEMPT - received username_or_email={username_or_email!r}, password length={len(password)}")

        user = db_session.query(User).filter(
            (User.username == username_or_email) | (User.email == username_or_email)
        ).first()

        if not user:
            print(f"LOGIN DEBUG - no user found for {username_or_email!r}")
            return jsonify({'success': False, 'message': 'Invalid credentials'})

        print(f"LOGIN DEBUG - found user id={user.id} username={user.username!r} email={user.email!r} stored_hash_len={len(user.password_hash or '')}")

        try:
            valid = check_password_hash(user.password_hash or '', password)
        except Exception as e:
            print(f"LOGIN DEBUG - check_password_hash raised: {e}")
            valid = False

        print(f"LOGIN DEBUG - password check result: {valid}")

        if user and valid:
            # Reset rate limit on successful login
            RateLimiter.reset_attempts(client_id)
            
            # SESSION FIXATION PROTECTION: Regenerate session ID
            # Store old session data
            old_session_data = dict(session)
            # Clear the session
            session.clear()
            # Restore session data (this forces Flask to generate a new session ID)
            session.update(old_session_data)
            # Make session permanent for timeout functionality
            session.permanent = True
            
            # SINGLE ACTIVE SESSION: Store new session ID in database
            # This will invalidate any other active sessions for this user
            new_session_id = session.get('_id')
            if not new_session_id:
                # If _id doesn't exist yet, generate it by modifying session
                session.modified = True
                new_session_id = session.get('_id')
            
            user.session_id = new_session_id
            user.is_online = True
            user.last_seen = datetime.now(timezone.utc)
            user.last_activity = datetime.now(timezone.utc)
            db_session.commit()
            
            # Now login the user
            login_user(user)

            if not user.is_profile_complete:
                return jsonify({'success': True, 'redirect': url_for('complete_profile')})

            if user.role == 'Executive':
                return jsonify({'success': True, 'redirect': url_for('executive_dashboard')})
            elif user.role == 'Guide':
                return jsonify({'success': True, 'redirect': url_for('guide_dashboard')})
            elif user.role == 'Learner':
                return jsonify({'success': True, 'redirect': url_for('learner_dashboard')})
        else:
            # Record failed attempt
            RateLimiter.record_attempt(client_id)

        return jsonify({'success': False, 'message': 'Invalid credentials'})
    
    return render_template('login.html')

@app.route('/auth/google')
def google_login():
    # Use consistent redirect URI
    redirect_uri = f'http://{LOCAL_SERVER_NAME}/auth/google/callback'
    print(f"DEBUG - Google Login Redirect URI: {redirect_uri}")
    print("⚠️ Make sure this URI is exactly configured in Google Cloud Console")
    return google.authorize_redirect(redirect_uri=redirect_uri)

def get_dashboard_url(user):
    """Helper function to get the appropriate dashboard URL based on user role"""
    if user.role == 'Executive':
        return url_for('executive_dashboard')
    elif user.role == 'Guide':
        return url_for('guide_dashboard')
    elif user.role == 'Learner':
        return url_for('learner_dashboard')
    return url_for('index')

@app.route('/auth/google/callback')
def google_callback():
    try:
        token = google.authorize_access_token()
        if not token:
            flash('Failed to get access token. Please try again.')
            return redirect(url_for('login'))

        # Request userinfo explicitly using the token to avoid empty/implicit payloads
        resp = google.get('userinfo', token=token)
        try:
            user_info = resp.json()
        except Exception as e:
            print(f"❌ Failed to decode userinfo JSON in google_callback: {e}")
            print(f"Response status: {getattr(resp, 'status_code', 'N/A')}, text: {getattr(resp, 'text', '')!r}")
            flash('Failed to get user info from Google. Please try again.')
            return redirect(url_for('login'))

        email = user_info.get('email')
        google_id = user_info.get('sub') or user_info.get('id')
        name = user_info.get('name')

        if not email:
            flash('Email not provided by Google. Please try again.')
            return redirect(url_for('login'))
            
        user = db_session.query(User).filter_by(email=email).first()
            
        if user:
            # Existing user - Apply same session security as regular login
            
            # SESSION FIXATION PROTECTION: Regenerate session ID
            old_session_data = dict(session)
            session.clear()
            session.update(old_session_data)
            session.permanent = True
            
            # SINGLE ACTIVE SESSION: Store new session ID
            new_session_id = session.get('_id')
            if not new_session_id:
                session.modified = True
                new_session_id = session.get('_id')
            
            user.google_id = google_id
            user.session_id = new_session_id
            user.is_online = True
            user.last_seen = datetime.now(timezone.utc)
            user.last_activity = datetime.now(timezone.utc)
            db_session.commit()
            login_user(user)
            
            if not user.is_profile_complete:
                # Store Google info in session for profile completion
                session['google_email_profile'] = email
                session['google_name_profile'] = name
                return redirect(url_for('complete_profile'))
            
            return redirect(get_dashboard_url(user))
        else:
            # New user registration
            session['google_email'] = email
            session['google_id'] = google_id
            session['google_name'] = name
            return redirect(url_for('register'))

    except Exception as e:
        print(f"Google auth error: {e}")
        app.logger.error(f"Google auth error: {str(e)}")
        flash('Failed to authenticate with Google. Please try again.')
        return redirect(url_for('login'))

@app.route('/complete_profile', methods=['GET', 'POST'])
@login_required
def complete_profile():
    # Redirect if profile is already complete
    if current_user.is_profile_complete:
        return redirect(get_dashboard_url(current_user))
    
    if request.method == 'POST':
        # Debug: Log form data
        print(f"DEBUG complete_profile - Form data keys: {list(request.form.keys())}")
        print(f"DEBUG complete_profile - Files: {list(request.files.keys())}")
        
        try:
            # Get form data with fallback to direct form access
            try:
                full_name = safe_form_get('full_name')
                if not full_name:
                    full_name = request.form.get('full_name', '').strip()
            except Exception as e:
                print(f"Error getting full_name: {e}")
                full_name = request.form.get('full_name', '').strip()
            
            new_password = request.form.get('new_password', '').strip()  # Don't sanitize password
            
            try:
                bio = safe_form_get('bio', '')
                if not bio:
                    bio = request.form.get('bio', '').strip()
            except Exception as e:
                print(f"Error getting bio: {e}")
                bio = request.form.get('bio', '').strip()
        except ValueError as e:
            return jsonify({'success': False, 'message': f'Invalid input: {str(e)}'}), 400
        
        # Get Google email from session if it exists
        google_email = session.get('google_email_profile')
        google_name = session.get('google_name_profile')
        
        # If this is a Google login and emails don't match
        if google_email and google_email.lower() != current_user.email.lower():
            return jsonify({
                'success': False, 
                'message': 'The Google account email does not match your registered email'
            })
        
        # Update user information
        current_user.full_name = full_name or google_name  # Use Google name as fallback
        if new_password:
            current_user.password_hash = generate_password_hash(new_password, method='pbkdf2:sha256')
        
        if 'profile_photo' in request.files:
            photo_file = request.files['profile_photo']
            if photo_file and photo_file.filename:
                filename = secure_filename(photo_file.filename)
                photo_path = os.path.join(PHOTO_FOLDER, f"{current_user.username}_{filename}")
                photo_file.save(photo_path)
                current_user.profile_photo = photo_path
        
        # Mark profile as complete
        current_user.is_profile_complete = True
        
        # If this was a Google login, store the Google ID
        if google_email:
            current_user.google_id = session.get('google_id')
        
        db_session.commit()
        
        # Clean up session
        session.pop('google_email_profile', None)
        session.pop('google_name_profile', None)
        session.pop('google_id', None)
        
        # Return success response
        return jsonify({
            'success': True,
            'message': 'Profile completed successfully!',
            'redirect': get_dashboard_url(current_user)
        })
    
    # GET request - render profile completion form
    return render_template('complete_profile.html', 
                         google_email=session.get('google_email_profile'),
                         google_name=session.get('google_name_profile'))

@app.route('/update_profile_picture', methods=['POST'])
@login_required
def update_profile_picture():
    """Update profile picture for any role (Learner, Guide, Executive)"""
    try:
        if 'profile_photo' not in request.files:
            return jsonify({'success': False, 'message': 'No file provided'})
        
        photo_file = request.files['profile_photo']
        
        if not photo_file or not photo_file.filename:
            return jsonify({'success': False, 'message': 'No file selected'})
        
        # Validate file type
        if not (photo_file.filename.lower().endswith(('.jpg', '.jpeg', '.png', '.gif'))):
            return jsonify({'success': False, 'message': 'Invalid file type. Please upload an image (JPG, PNG, GIF)'})
        
        # Validate file size (max 5MB)
        photo_file.seek(0, os.SEEK_END)
        file_size = photo_file.tell()
        photo_file.seek(0)
        
        if file_size > 5 * 1024 * 1024:  # 5MB
            return jsonify({'success': False, 'message': 'File size too large. Maximum size is 5MB'})
        
        # Delete old profile photo if exists
        if current_user.profile_photo:
            old_photo_path = os.path.join(app.root_path, current_user.profile_photo)
            if os.path.exists(old_photo_path):
                try:
                    os.remove(old_photo_path)
                except Exception as e:
                    print(f"Error deleting old photo: {e}")
        
        # Save new photo
        filename = secure_filename(photo_file.filename)
        timestamp = datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')
        photo_path = os.path.join(PHOTO_FOLDER, f"{current_user.username}_{timestamp}_{filename}")
        photo_file.save(photo_path)
        
        # Update user record (store relative path)
        current_user.profile_photo = photo_path
        db_session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Profile picture updated successfully',
            'photo_url': url_for('static', filename=photo_path.replace('static/', ''))
        })
        
    except Exception as e:
        db_session.rollback()
        print(f"Error updating profile picture: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})
        
        return jsonify({
            'success': True, 
            'redirect': get_dashboard_url(current_user)
        })
    
    return render_template('complete_profile.html')

@app.route('/auth/google/profile')
@login_required
def google_profile_auth():
    # Use consistent redirect URI for profile completion
    redirect_uri = f'http://{LOCAL_SERVER_NAME}/auth/google/profile/callback'
    print(f"DEBUG - Google Profile Redirect URI: {redirect_uri}")
    print("⚠️ Make sure this URI is exactly configured in Google Cloud Console")
    return google.authorize_redirect(redirect_uri=redirect_uri)

@app.route('/auth/google/profile/callback')
@login_required
def google_profile_callback():
    try:
        print("🔄 Processing Google profile callback...")
        token = google.authorize_access_token()
        if not token:
            print("❌ No token returned from google.authorize_access_token()")
            flash('Failed to get access token from Google', 'error')
            return redirect(url_for('complete_profile'))

        access_token = token.get('access_token')
        if access_token:
            print(f"✓ Got access token: {access_token[:10]}...")
        else:
            print("✓ Token received but no access_token field present")

        # Try to get userinfo from the userinfo endpoint first
        resp = google.get('userinfo', token=token)
        print(f"Userinfo response status: {getattr(resp, 'status_code', 'N/A')}")
        user_info = None
        try:
            user_info = resp.json()
        except Exception as e:
            print(f"❌ Failed decoding userinfo JSON: {e}")
            print(f"Response text: {getattr(resp, 'text', '')!r}")

        # Fallback: try to decode id_token payload if userinfo empty
        if not user_info:
            id_token = token.get('id_token')
            if id_token:
                try:
                    # Decode JWT payload without verification to extract claims
                    parts = id_token.split('.')
                    if len(parts) >= 2:
                        import base64, json
                        payload = parts[1]
                        # Add padding
                        payload += '=' * (-len(payload) % 4)
                        decoded = base64.urlsafe_b64decode(payload.encode())
                        user_info = json.loads(decoded.decode())
                        print('✓ Extracted user_info from id_token payload')
                except Exception as e:
                    print(f"❌ Failed to decode id_token payload: {e}")

        if not user_info:
            flash('Failed to fetch Google user info', 'error')
            return redirect(url_for('complete_profile'))

        print(f"✓ Got user info: {user_info}")

        # Extract main fields
        email = user_info.get('email')
        name = user_info.get('name')
        google_sub = user_info.get('sub') or user_info.get('id')

        # If the current user hasn't completed profile yet, store Google info in session
        # and redirect back to the complete_profile page so the UI updates and user can submit
        if current_user and not current_user.is_profile_complete:
            session['google_email_profile'] = email
            session['google_name_profile'] = name
            session['google_id'] = google_sub or session.get('google_id')
            print(f"➡️ Stored google info in session for user {current_user.email}: {email}")
            return redirect(url_for('complete_profile'))

        # Otherwise (user already has profile), update DB and redirect accordingly
        user = db_session.query(User).filter_by(id=current_user.id).first()
        if not user:
            print("❌ User not found in database")
            flash('User not found', 'error')
            return redirect(url_for('login'))

        user.google_id = google_sub or user.google_id
        user.full_name = name or user.full_name
        user.profile_photo = user_info.get('picture') or user.profile_photo
        user.is_profile_complete = True
        db_session.commit()
        print(f"✓ Updated profile for user {user.email}")

        flash('Profile completed successfully!', 'success')
        if user.role == 'Guide':
            return redirect(url_for('guide_dashboard'))
        else:
            return redirect(url_for('learner_dashboard'))

    except Exception as e:
        print(f"❌ Error in profile callback: {str(e)}")
        traceback.print_exc()
        flash('Error completing profile', 'error')
        return redirect(url_for('complete_profile'))

@app.route('/executive/dashboard')
@login_required
@role_required('Executive')
def executive_dashboard():
    """Merged executive dashboard with all features"""
    # Get all courses created by the executive
    courses = db_session.query(Course).filter_by(executive_id=current_user.id).all()
    
    # Get announcements
    announcements = db_session.query(Announcement).order_by(
        Announcement.created_at.desc()
    ).limit(10).all()
    
    return render_template(
        'executive_dashboard.html', 
        courses=courses,
        announcements=announcements
    )

@app.route('/executive/add_learners', methods=['POST'])
@login_required
@role_required('Executive')
@rate_limit_if_available("20 per minute")
def add_learners_to_course():
    try:
        # Debug: Log form data
        print(f"DEBUG add_learners - Form data keys: {list(request.form.keys())}")
        print(f"DEBUG add_learners - course_id: {request.form.get('course_id', 'NOT FOUND')}")
        print(f"DEBUG add_learners - learner_emails: {request.form.get('learner_emails', 'NOT FOUND')}")
        print(f"DEBUG add_learners - new_learner_emails: {request.form.get('new_learner_emails', 'NOT FOUND')}")
        
        # Get and validate course_id
        course_id_raw = request.form.get('course_id', '').strip()
        if not course_id_raw:
            return jsonify({'success': False, 'message': 'Course ID is required'})
    
        try:
            course_id = int(course_id_raw)
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': 'Invalid course ID format. Please refresh and try again.'})
        
        # Get and validate learner emails (handle both field names)
        learner_emails_raw = request.form.get('learner_emails', '') or request.form.get('new_learner_emails', '')
        if isinstance(learner_emails_raw, str):
            learner_emails_raw = learner_emails_raw.strip()
        else:
            learner_emails_raw = ''
        
        if not learner_emails_raw:
            return jsonify({'success': False, 'message': 'Please provide at least one learner email'}), 400
        
        learner_emails = [email.strip().lower() for email in learner_emails_raw.split(',') if email.strip()]
    
        if not learner_emails:
            return jsonify({'success': False, 'message': 'Please provide at least one valid learner email'})
    
        if len(learner_emails) > 20:
            return jsonify({'success': False, 'message': 'Maximum 20 learners can be added at once'})
        
        # Remove duplicates
        learner_emails = list(dict.fromkeys(learner_emails))  # Preserves order while removing duplicates
    
        # Validate email formats
        import re
        email_pattern = r'^[^\s@]+@[^\s@]+\.[^\s@]+$'
        invalid_emails = [email for email in learner_emails if not re.match(email_pattern, email)]
        if invalid_emails:
            return jsonify({'success': False, 'message': f'Invalid email format: {", ".join(invalid_emails[:5])}'})
    
        # Fetch course with proper error handling
        course = db_session.query(Course).get(course_id)
        if not course:
            return jsonify({'success': False, 'message': 'Course not found. It may have been deleted.'})
        
        # Ensure executive owns this course
        if course.executive_id != current_user.id:
            return jsonify({'success': False, 'message': 'Unauthorized: You do not have access to this course'})
        
        added_learners = []
        already_enrolled = []
        skipped_non_learners = []
        failed_emails = []
        email_errors = []
        
        for learner_email in learner_emails:
            try:
                existing_user = db_session.query(User).filter_by(email=learner_email).first()
            
                if existing_user:
                    # Check if user is already a learner
                    if existing_user.role != 'Learner':
                        skipped_non_learners.append(f"{learner_email} (role: {existing_user.role})")
                        continue
                
                    # Check if already enrolled
                    existing_enrollment = db_session.query(CourseLearner).filter_by(
                        course_id=course.id,
                        learner_id=existing_user.id
                    ).first()
                
                    if existing_enrollment:
                        already_enrolled.append(learner_email)
                        continue
                
                    # Enroll existing learner
                    try:
                        enrollment = CourseLearner(course_id=course.id, learner_id=existing_user.id)
                        db_session.add(enrollment)
                        db_session.flush()  # Flush to catch any errors early
                        added_learners.append(learner_email)
                    except Exception as enroll_error:
                        # Don't rollback here - just track the error and continue
                        failed_emails.append(learner_email)
                        email_errors.append(f"{learner_email}: {str(enroll_error)}")
                        print(f"Error enrolling {learner_email}: {enroll_error}")
                        import traceback
                        traceback.print_exc()
                        continue
            
                else:
                    # Create new learner with unique username
                    learner_password = generate_random_password()
                    timestamp = int(datetime.now(timezone.utc).timestamp())
                    base_username = f"Learner_{learner_email.split('@')[0]}_{course.course_name.replace(' ', '_')[:10]}"
                    learner_username = f"{base_username}_{timestamp}"
                    
                    # Ensure username is unique
                    max_attempts = 10
                    attempt = 0
                    while attempt < max_attempts:
                        existing_username = db_session.query(User).filter_by(username=learner_username).first()
                        if not existing_username:
                            break
                        attempt += 1
                        learner_username = f"{base_username}_{timestamp}_{attempt}"
                    
                    if attempt >= max_attempts:
                        failed_emails.append(learner_email)
                        email_errors.append(f"{learner_email}: Could not generate unique username")
                        continue
                    
                    # Check if email already exists (shouldn't happen, but double-check)
                    if db_session.query(User).filter_by(email=learner_email).first():
                        already_enrolled.append(learner_email)
                        continue
                    
                    try:
                        learner_user = User(
                            username=learner_username,
                            email=learner_email,
                            password_hash=generate_password_hash(learner_password, method='pbkdf2:sha256'),
                            role='Learner',
                            organization_name=course.organization,
                            is_profile_complete=False
                        )
                        db_session.add(learner_user)
                        db_session.flush()  # Get the user ID
                
                        enrollment = CourseLearner(course_id=course.id, learner_id=learner_user.id)
                        db_session.add(enrollment)
                
                        # Send credentials email (don't fail if email fails)
                        email_body = f"""
                <html>
                <body style="font-family: Arial, sans-serif; padding: 20px;">
                    <h2 style="color: #667eea;">Welcome to {course.course_name}!</h2>
                    <p>Your learner account has been created.</p>
                    <div style="background: #f8fafc; padding: 15px; border-radius: 8px; margin: 20px 0;">
                        <p><strong>Username:</strong> {learner_username}</p>
                        <p><strong>Password:</strong> {learner_password}</p>
                    </div>
                    <p>Please login and complete your profile to access course materials.</p>
                    <p style="color: #64748b; font-size: 12px; margin-top: 30px;">
                        If you did not expect this email, please ignore it.
                    </p>
                </body>
                </html>
                """
                        email_sent = send_email(learner_email, f"Your {course.course_name} Credentials", email_body)
                        if not email_sent:
                            print(f"Warning: Failed to send email to {learner_email}, but user was created")
                        
                        added_learners.append(learner_email)
                    except Exception as create_error:
                        # If user creation failed, try to remove the failed user object from session
                        try:
                            if 'learner_user' in locals():
                                db_session.expunge(learner_user)
                        except:
                            pass
                        failed_emails.append(learner_email)
                        email_errors.append(f"{learner_email}: {str(create_error)}")
                        print(f"Error creating learner {learner_email}: {create_error}")
                        import traceback
                        traceback.print_exc()
                        continue
                
            except Exception as email_error:
                # Handle individual email errors without stopping the whole process
                failed_emails.append(learner_email)
                email_errors.append(f"{learner_email}: {str(email_error)}")
                continue
        
        # Commit all successful additions
        try:
            db_session.commit()
        except Exception as commit_error:
            db_session.rollback()
            import traceback
            traceback.print_exc()
            return jsonify({
                'success': False, 
                'message': f'Database error while saving: {str(commit_error)}. Please try again.'
            })
        
        # Build comprehensive success message
        message_parts = []
        if added_learners:
            message_parts.append(f'Successfully added {len(added_learners)} learner(s)')
        if already_enrolled:
            message_parts.append(f'{len(already_enrolled)} already enrolled')
        if skipped_non_learners:
            message_parts.append(f'{len(skipped_non_learners)} skipped (not learners): {", ".join(skipped_non_learners[:3])}')
        if failed_emails:
            message_parts.append(f'{len(failed_emails)} failed: {", ".join(failed_emails[:3])}')
        
        if not added_learners and not already_enrolled:
            return jsonify({
                'success': False,
                'message': 'No learners were added. ' + ('. '.join(message_parts) if message_parts else 'Please check the email addresses and try again.')
            })
        
        return jsonify({
            'success': True,
            'message': '. '.join(message_parts) if message_parts else 'Operation completed',
            'details': {
                'added': len(added_learners),
                'already_enrolled': len(already_enrolled),
                'skipped': len(skipped_non_learners),
                'failed': len(failed_emails)
            }
        })
    
    except Exception as e:
        db_session.rollback()
        import traceback
        error_trace = traceback.format_exc()
        print(f"Error adding learners: {error_trace}")
        return jsonify({
            'success': False, 
            'message': f'An error occurred while adding learners: {str(e)}. Please try again or contact support.'
        })

@app.route('/api/course/<int:course_id>/details')
@login_required
@role_required('Executive')
def get_course_details(course_id):
    """Get detailed information about a course"""
    course = db_session.query(Course).filter_by(
        id=course_id, 
        executive_id=current_user.id
    ).first()
    
    if not course:
        return jsonify({'success': False, 'message': 'Course not found or access denied'})
    
    guide_data = None
    if course.guide:
        guide_data = {
            'id': course.guide.id,
            'username': course.guide.username,
            'full_name': course.guide.full_name,
            'email': course.guide.email,
            'profile_photo': url_for('static', filename=course.guide.profile_photo.replace('static/', '')) if course.guide.profile_photo else None,
            'is_online': course.guide.is_online
        }
    
    learners_data = []
    for course_learner in course.learners:
        learner = course_learner.learner
        learners_data.append({
            'id': learner.id,
            'username': learner.username,
            'full_name': learner.full_name,
            'email': learner.email,
            'profile_photo': url_for('static', filename=learner.profile_photo.replace('static/', '')) if learner.profile_photo else None,
            'is_online': learner.is_online
        })
    
    course_data = {
        'id': course.id,
        'course_name': course.course_name,
        'organization': course.organization,
        'created_at': course.created_at.strftime('%B %d, %Y'),
        'guide': guide_data,
        'learners': learners_data
    }
    
    return jsonify({'success': True, 'course': course_data})

@app.route('/executive/create_course', methods=['POST'])
@login_required
@role_required('Executive')
@rate_limit_if_available("10 per minute")
def create_course():
    """Create a new course with comprehensive validation and error handling"""
    try:
        # Debug: Log form data
        print(f"DEBUG create_course - Form data keys: {list(request.form.keys())}")
        print(f"DEBUG create_course - course_name: {request.form.get('course_name', 'NOT FOUND')}")
        print(f"DEBUG create_course - guide_email: {request.form.get('guide_email', 'NOT FOUND')}")
        
        # Get and sanitize form data (with fallback to direct form access)
        try:
            course_name = safe_form_get('course_name', '').strip()
            if not course_name:
                course_name = request.form.get('course_name', '').strip()
        except Exception as e:
            print(f"Error getting course_name: {e}")
            course_name = request.form.get('course_name', '').strip()
        
        guide_email_raw = request.form.get('guide_email', '').strip().lower()
        
        try:
            learner_emails_input = safe_form_get('learner_emails', '').strip()
            if not learner_emails_input:
                learner_emails_input = request.form.get('learner_emails', '').strip()
        except Exception as e:
            print(f"Error getting learner_emails: {e}")
            learner_emails_input = request.form.get('learner_emails', '').strip()
        
        # Validate course name
        if not course_name:
            return jsonify({'success': False, 'message': 'Course name is required'}), 400
        
        if len(course_name) < 3:
            return jsonify({'success': False, 'message': 'Course name must be at least 3 characters'})
        
        if len(course_name) > 200:
            return jsonify({'success': False, 'message': 'Course name must be less than 200 characters'})
        
        # Validate guide email
        if not guide_email_raw:
            return jsonify({'success': False, 'message': 'Guide email is required'})
        
        if not InputValidator.validate_email(guide_email_raw):
            return jsonify({'success': False, 'message': 'Invalid guide email format'})
        
        guide_email = guide_email_raw
        
        # Parse and validate learner emails
        learner_emails = []
        if learner_emails_input:
            learner_emails = [email.strip().lower() for email in learner_emails_input.split(',') if email.strip()]
            # Remove duplicates
            learner_emails = list(dict.fromkeys(learner_emails))
        
        # Validate email formats
        import re
        email_pattern = r'^[^\s@]+@[^\s@]+\.[^\s@]+$'
        
        if not re.match(email_pattern, guide_email):
            return jsonify({'success': False, 'message': 'Invalid guide email format'})
        
        if learner_emails:
            if len(learner_emails) > 20:
                return jsonify({'success': False, 'message': 'Maximum 20 learners per course'})
            
            invalid_emails = [email for email in learner_emails if not re.match(email_pattern, email)]
            if invalid_emails:
                return jsonify({
                    'success': False, 
                    'message': f'Invalid email format: {", ".join(invalid_emails[:3])}'
                })
            
            # Check if guide email is in learner emails
            if guide_email in learner_emails:
                return jsonify({
                    'success': False, 
                    'message': 'Guide email cannot be the same as a learner email'
                })
        
        # Get organization name (handle None case)
        organization = current_user.organization_name or 'Default Organization'
        
        # Check if guide already exists
        existing_guide = db_session.query(User).filter_by(email=guide_email).first()
        
        if existing_guide:
            if existing_guide.role != 'Guide':
                return jsonify({
                    'success': False, 
                    'message': f'This email is already registered as {existing_guide.role}. Please use a different email for the guide.'
                })
            guide_user = existing_guide
        else:
            # Create new guide account with unique username
            guide_password = generate_random_password()
            timestamp = int(datetime.now(timezone.utc).timestamp())
            base_username = f"Guide_{guide_email.split('@')[0]}_{course_name.replace(' ', '_')[:10]}"
            guide_username = f"{base_username}_{timestamp}"
            
            # Ensure username is unique
            max_attempts = 10
            attempt = 0
            while attempt < max_attempts:
                existing_username = db_session.query(User).filter_by(username=guide_username).first()
                if not existing_username:
                    break
                attempt += 1
                guide_username = f"{base_username}_{timestamp}_{attempt}"
            
            if attempt >= max_attempts:
                return jsonify({
                    'success': False, 
                    'message': 'Could not generate unique username for guide. Please try again.'
                })
            
            try:
                guide_user = User(
                    username=guide_username,
                    email=guide_email,
                    password_hash=generate_password_hash(guide_password, method='pbkdf2:sha256'),
                    role='Guide',
                    organization_name=organization,
                    is_profile_complete=False
                )
                db_session.add(guide_user)
                db_session.flush()  # Get the user ID
            
                # Send guide credentials email (don't fail if email fails)
                guide_email_body = f"""
                <html>
            <body style="font-family: Arial, sans-serif; padding: 20px; background: #f8fafc;">
                <div style="max-width: 600px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1);">
                    <h2 style="color: #667eea; margin-bottom: 20px;">Welcome Guide for {course_name}!</h2>
                    <p style="color: #334155; font-size: 16px; line-height: 1.6;">
                        Your guide account has been created for the course: <strong>{course_name}</strong>
                    </p>
                    <div style="background: #f1f5f9; padding: 20px; border-radius: 8px; margin: 20px 0;">
                        <p style="margin: 10px 0;"><strong>Username:</strong> {guide_username}</p>
                        <p style="margin: 10px 0;"><strong>Password:</strong> {guide_password}</p>
                        <p style="margin: 10px 0;"><strong>Login URL:</strong> <a href="http://127.0.0.1:5000/login" style="color: #667eea;">Login Here</a></p>
                    </div>
                    <p style="color: #64748b; font-size: 14px;">
                        Please login and complete your profile to start managing your course.
                    </p>
                    <hr style="border: none; border-top: 1px solid #e2e8f0; margin: 20px 0;">
                    <p style="color: #94a3b8; font-size: 12px;">
                        If you did not expect this email, please contact the administrator.
                    </p>
                </div>
            </body>
            </html>
            """
                email_sent = send_email(guide_email, f"Guide Credentials for {course_name}", guide_email_body)
                if not email_sent:
                    print(f"Warning: Failed to send email to guide: {guide_email}, but guide was created")
            except Exception as guide_error:
                db_session.rollback()
                import traceback
                traceback.print_exc()
                return jsonify({
                    'success': False, 
                    'message': f'Error creating guide account: {str(guide_error)}. Please check the email and try again.'
                })
        
        # Create course
        try:
            new_course = Course(
                course_name=course_name,
                executive_id=current_user.id,
                guide_id=guide_user.id,
                organization=organization
            )
            db_session.add(new_course)
            db_session.flush()  # Get the course ID
        except Exception as course_error:
            db_session.rollback()
            import traceback
            traceback.print_exc()
            return jsonify({
                'success': False, 
                'message': f'Error creating course: {str(course_error)}. Please try again.'
            })
        
        # Add learners with comprehensive error handling
        added_learners_count = 0
        skipped_learners = []
        failed_learners = []
        
        for learner_email in learner_emails:
            try:
                existing_learner = db_session.query(User).filter_by(email=learner_email).first()
            
                if existing_learner:
                    if existing_learner.role != 'Learner':
                        skipped_learners.append(f"{learner_email} (role: {existing_learner.role})")
                        continue
                    learner_user = existing_learner
                else:
                    # Create new learner account with unique username
                    learner_password = generate_random_password()
                    timestamp = int(datetime.now(timezone.utc).timestamp())
                    base_username = f"Learner_{learner_email.split('@')[0]}_{course_name.replace(' ', '_')[:10]}"
                    learner_username = f"{base_username}_{timestamp}"
                    
                    # Ensure username is unique
                    max_attempts = 10
                    attempt = 0
                    while attempt < max_attempts:
                        existing_username = db_session.query(User).filter_by(username=learner_username).first()
                        if not existing_username:
                            break
                        attempt += 1
                        learner_username = f"{base_username}_{timestamp}_{attempt}"
                    
                    if attempt >= max_attempts:
                        failed_learners.append(f"{learner_email} (username generation failed)")
                        continue
                    
                    try:
                        learner_user = User(
                            username=learner_username,
                            email=learner_email,
                            password_hash=generate_password_hash(learner_password, method='pbkdf2:sha256'),
                            role='Learner',
                            organization_name=organization,
                            is_profile_complete=False
                        )
                        db_session.add(learner_user)
                        db_session.flush()  # Get the user ID
                
                        # Send learner credentials email (don't fail if email fails)
                        learner_email_body = f"""
                        <html>
                        <body style="font-family: Arial, sans-serif; padding: 20px; background: #f8fafc;">
                            <div style="max-width: 600px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1);">
                                <h2 style="color: #667eea; margin-bottom: 20px;">Welcome to {course_name}!</h2>
                                <p style="color: #334155; font-size: 16px; line-height: 1.6;">
                                    Your learner account has been created.
                                </p>
                                <div style="background: #f1f5f9; padding: 20px; border-radius: 8px; margin: 20px 0;">
                                    <p style="margin: 10px 0;"><strong>Username:</strong> {learner_username}</p>
                                    <p style="margin: 10px 0;"><strong>Password:</strong> {learner_password}</p>
                                    <p style="margin: 10px 0;"><strong>Login URL:</strong> <a href="http://127.0.0.1:5000/login" style="color: #667eea;">Login Here</a></p>
                                </div>
                                <p style="color: #64748b; font-size: 14px;">
                                    Please login and complete your profile to access course materials.
                                </p>
                                <hr style="border: none; border-top: 1px solid #e2e8f0; margin: 20px 0;">
                                <p style="color: #94a3b8; font-size: 12px;">
                                    If you did not expect this email, please ignore it.
                                </p>
                            </div>
                        </body>
                        </html>
                        """
                        email_sent = send_email(learner_email, f"Your {course_name} Credentials", learner_email_body)
                        if not email_sent:
                            print(f"Warning: Failed to send email to learner: {learner_email}, but learner was created")
                        
                        added_learners_count += 1
                    except Exception as learner_create_error:
                        failed_learners.append(f"{learner_email} ({str(learner_create_error)})")
                        continue
                
                    # Enroll learner in course (check for duplicate enrollment)
                    existing_enrollment = db_session.query(CourseLearner).filter_by(
                        course_id=new_course.id,
                        learner_id=learner_user.id
                    ).first()
                
                    if existing_enrollment:
                        continue  # Already enrolled, skip
                
                    try:
                        course_learner = CourseLearner(
                            course_id=new_course.id,
                            learner_id=learner_user.id
                        )
                        db_session.add(course_learner)
                        added_learners_count += 1
                    except Exception as enroll_error:
                        failed_learners.append(f"{learner_email} (enrollment failed: {str(enroll_error)})")
                        continue
                    
            except Exception as learner_error:
                failed_learners.append(f"{learner_email} ({str(learner_error)})")
                continue
        
        # Commit all changes
        try:
            db_session.commit()
        except Exception as commit_error:
            db_session.rollback()
            import traceback
            traceback.print_exc()
            return jsonify({
                'success': False, 
                'message': f'Database error while saving course: {str(commit_error)}. Please try again.'
            })
        
        # Build success message
        message_parts = [f'Course "{course_name}" created successfully']
        if added_learners_count > 0:
            message_parts.append(f'with {added_learners_count} learner(s)')
        if skipped_learners:
            message_parts.append(f'{len(skipped_learners)} skipped (not learners)')
        if failed_learners:
            message_parts.append(f'{len(failed_learners)} failed to add')
        message = ' '.join(message_parts) + '!'
        
        return jsonify({
            'success': True, 
            'message': message,
            'course_id': new_course.id
        })
        
    except ValueError as ve:
        db_session.rollback()
        return jsonify({
            'success': False, 
            'message': f'Invalid input: {str(ve)}. Please check your data and try again.'
        })
    except Exception as e:
        db_session.rollback()
        import traceback
        error_trace = traceback.format_exc()
        print(f"Error creating course: {error_trace}")
        return jsonify({
            'success': False, 
            'message': f'An error occurred while creating the course: {str(e)}. Please try again or contact support if the problem persists.'
        })

@app.route('/executive/delete_user/<int:user_id>', methods=['POST'])
@login_required
@role_required('Executive')
def request_delete_user(user_id):
    otp_code = generate_otp()
    expires_at = datetime.now(timezone.utc) + timedelta(minutes=10)
    
    new_otp = OTP(
        user_id=current_user.id,
        otp_code=otp_code,
        purpose=f'delete_user_{user_id}',
        expires_at=expires_at
    )
    db_session.add(new_otp)
    db_session.commit()
    
    email_body = f"""
    <html>
    <body>
        <h2>Delete User Verification</h2>
        <p>Your OTP code is: <strong>{otp_code}</strong></p>
        <p>This code expires in 10 minutes.</p>
    </body>
    </html>
    """
    send_email(current_user.email, "User Deletion OTP", email_body)
    
    return jsonify({'success': True, 'message': 'OTP sent to your email'})

@app.route('/executive/verify_delete_user', methods=['POST'])
@login_required
@role_required('Executive')
def verify_delete_user():
    user_id = request.form.get('user_id')
    otp_code = request.form.get('otp_code')
    
    otp = db_session.query(OTP).filter_by(
        user_id=current_user.id,
        otp_code=otp_code,
        purpose=f'delete_user_{user_id}',
        is_used=False
    ).first()
    
    if not otp:
        return jsonify({'success': False, 'message': 'Invalid OTP'})
    
    # Check OTP expiration - handle timezone-aware and timezone-naive datetimes
    now = datetime.now(timezone.utc)
    expires_at = otp.expires_at
    
    # Make expires_at timezone-aware if it's naive (assume UTC)
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    
    if expires_at < now:
        return jsonify({'success': False, 'message': 'OTP expired'})
    
    user_to_delete = db_session.query(User).get(user_id)
    if user_to_delete:
        try:
            # Delete all related records in the correct order to avoid foreign key violations
            
            # 1. Delete LearnerMCQResponse (references AssignmentAnswer)
            from sqlalchemy import text
            db_session.execute(
                text("DELETE FROM learner_mcq_responses WHERE answer_id IN (SELECT id FROM assignment_answers WHERE submission_id IN (SELECT id FROM assignment_submissions WHERE learner_id = :user_id))"),
                {'user_id': user_id}
            )
            
            # 2. Delete PlagiarismMatch (references AssignmentAnswer)
            db_session.execute(
                text("DELETE FROM plagiarism_matches WHERE answer_id IN (SELECT id FROM assignment_answers WHERE submission_id IN (SELECT id FROM assignment_submissions WHERE learner_id = :user_id)) OR matched_answer_id IN (SELECT id FROM assignment_answers WHERE submission_id IN (SELECT id FROM assignment_submissions WHERE learner_id = :user_id))"),
                {'user_id': user_id}
            )
            
            # 3. Delete AssignmentAnswer (references AssignmentSubmission)
            db_session.execute(
                text("DELETE FROM assignment_answers WHERE submission_id IN (SELECT id FROM assignment_submissions WHERE learner_id = :user_id)"),
                {'user_id': user_id}
            )
            
            # 4. Delete AssignmentSubmission (references User)
            db_session.query(AssignmentSubmission).filter_by(learner_id=user_id).delete()
            
            # 5. Delete DiagramSubmission (references User)
            try:
                db_session.query(DiagramSubmission).filter_by(learner_id=user_id).delete()
            except:
                pass  # Table might not exist
            
            # 6. Delete DiagramAssignment created by user (if user is guide)
            try:
                diagram_assignments = db_session.query(DiagramAssignment).filter_by(created_by=user_id).all()
                for assignment in diagram_assignments:
                    # Delete related submissions first
                    db_session.query(DiagramSubmission).filter_by(assignment_id=assignment.id).delete()
                db_session.query(DiagramAssignment).filter_by(created_by=user_id).delete()
            except:
                pass  # Table might not exist
            
            # 7. Delete DocumentQA and related records (references user_id)
            try:
                documents = db_session.query(DocumentQA).filter_by(user_id=user_id).all()
                for doc in documents:
                    # Delete related DocumentQAChunk and DocumentQAChat (cascade should handle this, but being explicit)
                    db_session.query(DocumentQAChunk).filter_by(document_id=doc.id).delete()
                    db_session.query(DocumentQAChat).filter_by(document_id=doc.id).delete()
                db_session.query(DocumentQA).filter_by(user_id=user_id).delete()
            except:
                pass  # Table might not exist
            
            # 8. Delete CourseLearner records
            db_session.query(CourseLearner).filter_by(learner_id=user_id).delete()
            
            # 9. Delete LearnerSuggestion
            db_session.query(LearnerSuggestion).filter_by(learner_id=user_id).delete()
            
            # 10. Delete MaterialRead
            db_session.query(MaterialRead).filter_by(learner_id=user_id).delete()
            
            # 11. Delete Notification records
            try:
                db_session.query(Notification).filter_by(user_id=user_id).delete()
            except:
                pass  # Table might not exist
            
            # 12. Delete OTP records
            db_session.query(OTP).filter_by(user_id=user_id).delete()
            
            # 13. Delete Messages (as sender or receiver)
            db_session.query(Message).filter(
                (Message.sender_id == user_id) | (Message.receiver_id == user_id)
            ).delete()
            
            # 14. Delete Announcements created by user
            db_session.query(Announcement).filter_by(user_id=user_id).delete()
            
            # 15. Delete Issues created by user
            db_session.query(Issue).filter_by(user_id=user_id).delete()
            
            # 16. Delete Uploads by user (and related DocumentChunks)
            uploads = db_session.query(Upload).filter_by(user_id=user_id).all()
            for upload in uploads:
                # Delete related DocumentChunks first
                db_session.query(DocumentChunk).filter_by(upload_id=upload.id).delete()
            db_session.query(Upload).filter_by(user_id=user_id).delete()
            
            # 17. Delete ExecutiveContact (if user is executive)
            db_session.query(ExecutiveContact).filter_by(executive_id=user_id).delete()
            
            # 18. Delete GuideFeedback (if user is guide)
            db_session.query(GuideFeedback).filter_by(guide_id=user_id).delete()
            
            # 19. Delete GuideSuggestion (if user is guide)
            db_session.query(GuideSuggestion).filter_by(guide_id=user_id).delete()
            
            # 20. Delete CustomScoreboardValue (if user is learner)
            try:
                db_session.query(CustomScoreboardValue).filter_by(learner_id=user_id).delete()
            except:
                pass  # Table might not exist or no records
            
            # 21. Handle PlagiarismMatch where user reviewed (if user is guide)
            # Set reviewed_by to None instead of deleting, as the matches are still valid
            try:
                db_session.query(PlagiarismMatch).filter_by(reviewed_by=user_id).update({'reviewed_by': None}, synchronize_session=False)
            except:
                # If update fails, try to delete (in case reviewed_by is not nullable)
                db_session.query(PlagiarismMatch).filter_by(reviewed_by=user_id).delete()
            
            # 22. Mark OTP as used
            otp.is_used = True
            
            # 23. Finally delete the user
            db_session.delete(user_to_delete)
            
            db_session.commit()
            return jsonify({'success': True, 'message': 'User deleted successfully'})
        except Exception as e:
            db_session.rollback()
            import traceback
            traceback.print_exc()
            print(f"Error deleting user: {str(e)}")
            return jsonify({'success': False, 'message': f'Error deleting user: {str(e)}. Please try again.'})
    
    return jsonify({'success': False, 'message': 'User not found'})

@app.route('/executive/announcements', methods=['POST'])
@login_required
@role_required('Executive')
def create_announcement():
    """Create new announcement (API endpoint)"""
    title = request.form.get('title')
    content = request.form.get('content')
    file_path = None
    
    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename:
            filename = secure_filename(file.filename)
            file_path = os.path.join(
                UPLOAD_FOLDER, 
                f"announcement_{datetime.now(timezone.utc).timestamp()}_{filename}"
            )
            file.save(file_path)
    
    announcement = Announcement(
        user_id=current_user.id,
        title=title,
        content=content,
        file_path=file_path
    )
    db_session.add(announcement)
    db_session.commit()
    
    return jsonify({
        'success': True, 
        'message': 'Announcement posted successfully'
    })

# ==================== DOCUMENT Q&A HELPER FUNCTIONS ====================

DOCUMENT_QA_UPLOAD_FOLDER = 'static/document_qa_uploads'
os.makedirs(DOCUMENT_QA_UPLOAD_FOLDER, exist_ok=True)

def extract_text_from_pdf_qa(file_path):
    """Extract text from PDF"""
    text = ""
    page_count = 0
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            page_count = len(pdf_reader.pages)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"PDF extraction error: {e}")
    return text, page_count

def extract_text_from_docx_qa(file_path):
    """Extract text from DOCX"""
    text = ""
    page_count = 0
    try:
        doc = DocxDocument(file_path)
        page_count = len(doc.paragraphs) // 25
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"DOCX extraction error: {e}")
    return text, page_count

def extract_text_from_txt_qa(file_path):
    """Extract text from TXT"""
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except:
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
        except Exception as e:
            print(f"TXT extraction error: {e}")
    return text, 1

def extract_text_from_csv_qa(file_path):
    """Extract text from CSV"""
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except Exception as e:
        print(f"CSV extraction error: {e}")
    return text, 1

def extract_text_from_image_qa(file_path):
    """Extract text from image using Google Vision OCR"""
    try:
        client = vision.ImageAnnotatorClient()
        with open(file_path, 'rb') as image_file:
            content = image_file.read()
        image = vision.Image(content=content)
        response = client.text_detection(image=image)
        texts = response.text_annotations
        if texts:
            return texts[0].description, 1
        return "", 1
    except Exception as e:
        print(f"Image OCR error: {e}")
        return "", 1

def chunk_text_qa(text, chunk_size=1000, overlap=200):
    """Split text into overlapping chunks"""
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        if chunk:
            chunks.append(chunk)
    return chunks

def get_embedding_qa(text):
    """Get embedding from Gemini"""
    try:
        if not text or len(text.strip()) == 0 or not GEMINI_API_KEY:
            return None
        
        result = genai.embed_content(
            model="models/text-embedding-004",
            content=text,
            task_type="retrieval_document"
        )
        
        embedding = np.array(result['embedding'], dtype=np.float32)
        
        if embedding is None or len(embedding) == 0:
            return None
        if np.isnan(embedding).any() or np.isinf(embedding).any():
            print(f"Warning: Embedding contains NaN or Inf values")
            return None
        
        return embedding.tolist()  # Convert to list for JSON storage
    except Exception as e:
        print(f"Embedding error: {e}")
        return None

def store_embeddings_qa(document_id, text):
    """Create and store embeddings for document chunks"""
    chunks = chunk_text_qa(text)
    
    stored_count = 0
    for idx, chunk in enumerate(chunks):
        if not chunk or len(chunk.strip()) == 0:
            continue
            
        embedding = get_embedding_qa(chunk)
        if embedding is not None:
            chunk_obj = DocumentQAChunk(
                document_id=document_id,
                chunk_text=chunk,
                chunk_index=idx,
                embedding=json.dumps(embedding)
            )
            db_session.add(chunk_obj)
            stored_count += 1
    
    db_session.commit()
    print(f"Stored {stored_count} embeddings for document {document_id}")

def generate_document_summary(text):
    """Generate a simple summary and extract topics using Gemini"""
    if not GEMINI_API_KEY or not text:
        return "Summary not available", []
    
    try:
        prompt = f"""Analyze this document and provide:
1. A simple 2-3 sentence summary
2. A list of main topics (as a JSON array)

Document text (first 5000 characters):
{text[:5000]}

Respond in this format:
SUMMARY: [your summary here]
TOPICS: [JSON array of topics]

Topics should be 1-3 words each, like: ["Cloud Storage", "Security", "Networking"]"""
        
        response = chat_model.generate_content(prompt)
        response_text = response.text
        
        # Parse response
        summary = ""
        topics = []
        
        if "SUMMARY:" in response_text:
            summary = response_text.split("SUMMARY:")[1].split("TOPICS:")[0].strip()
        if "TOPICS:" in response_text:
            topics_str = response_text.split("TOPICS:")[1].strip()
            try:
                topics = json.loads(topics_str)
            except:
                # Try to extract topics manually
                topics = re.findall(r'"([^"]+)"', topics_str)
        
        return summary, topics
    except Exception as e:
        print(f"Summary generation error: {e}")
        return "Summary generation failed", []

def markdown_to_html_qa(text):
    """Convert markdown to HTML for proper display"""
    if not text:
        return ""
    
    text = text.replace('```', '')
    lines = text.split('\n')
    html_lines = []
    in_list = False
    
    for line in lines:
        stripped = line.strip()
        is_bullet = (stripped.startswith('* ') and not stripped.startswith('**')) or (stripped.startswith('- ') and len(stripped) > 2)
        
        if is_bullet:
            if not in_list:
                html_lines.append('<ul style="margin: 10px 0; padding-left: 20px;">')
                in_list = True
            content = stripped[2:].strip()
            content = re.sub(r'\*\*([^*]+?)\*\*', r'<strong>\1</strong>', content)
            content = re.sub(r'(?<!\*)\*([^*\n]+?)\*(?!\*)', r'<em>\1</em>', content)
            html_lines.append(f'<li style="margin: 5px 0; line-height: 1.5;">{content}</li>')
        else:
            if in_list:
                html_lines.append('</ul>')
                in_list = False
            if stripped:
                para_text = stripped
                para_text = re.sub(r'\*\*([^*]+?)\*\*', r'<strong>\1</strong>', para_text)
                para_text = re.sub(r'(?<!\*)\*([^*\n]+?)\*(?!\*)', r'<em>\1</em>', para_text)
                html_lines.append(f'<p style="margin: 10px 0; line-height: 1.6;">{para_text}</p>')
            else:
                if html_lines and not html_lines[-1].endswith('</ul>'):
                    html_lines.append('<br>')
    
    if in_list:
        html_lines.append('</ul>')
    
    result = '\n'.join(html_lines)
    result = result.replace('<strong><strong>', '<strong>').replace('</strong></strong>', '</strong>')
    result = result.replace('<p style="margin: 10px 0; line-height: 1.6;"></p>', '')
    return result

# ==================== END DOCUMENT Q&A HELPER FUNCTIONS ====================

@app.route('/guide/dashboard')
@login_required
@role_required('Guide')
def guide_dashboard():
    if not current_user.is_profile_complete:
        return redirect(url_for('complete_profile'))
    
    # Fetch all data needed for all tabs
    announcements = db_session.query(Announcement).order_by(Announcement.created_at.desc()).limit(10).all()
    courses = db_session.query(Course).filter_by(guide_id=current_user.id).all()  # Get all courses, not just 5
    uploads = db_session.query(Upload).filter_by(user_id=current_user.id).order_by(Upload.created_at.desc()).all()
    
    # Fetch diagram evaluation data
    diagram_assignments = db_session.query(DiagramAssignment).join(Course).filter(
        Course.guide_id == current_user.id
    ).order_by(DiagramAssignment.created_at.desc()).all()
    
    diagram_submissions = db_session.query(DiagramSubmission).join(DiagramAssignment).join(Course).filter(
        Course.guide_id == current_user.id
    ).order_by(DiagramSubmission.submitted_at.desc()).all()
    
    return render_template('guide_dashboard.html', 
                         announcements=announcements, 
                         courses=courses,
                         uploads=uploads,
                         diagram_assignments=diagram_assignments,
                         diagram_submissions=diagram_submissions)

@app.route('/guide/upload', methods=['POST'])
@login_required
@role_required('Guide')
def guide_upload():
    course_id = request.form.get('course_id')
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file uploaded'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No file selected'})
    
    if file and allowed_file(file.filename):
        file_size = len(file.read())
        file.seek(0)
        
        if file_size > 50 * 1024 * 1024:
            return jsonify({'success': False, 'message': 'File size exceeds 50MB'})
        
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, f"{current_user.username}_{datetime.now(timezone.utc).timestamp()}_{filename}")
        file.save(file_path)
        
        upload = Upload(
            user_id=current_user.id,
            course_id=course_id,
            file_name=filename,
            file_path=file_path,
            file_type=filename.rsplit('.', 1)[1].lower(),
            file_size=file_size / (1024 * 1024)
        )
        db_session.add(upload)
        db_session.commit()
        
        return jsonify({'success': True, 'message': 'File uploaded successfully'})
    
    return jsonify({'success': False, 'message': 'Invalid file type'})

@app.route('/learner/dashboard')
@login_required
@role_required('Learner')
def learner_dashboard():
    if not current_user.is_profile_complete:
        return redirect(url_for('complete_profile'))
    
    # Fetch all data for the merged dashboard
    announcements = db_session.query(Announcement).order_by(Announcement.created_at.desc()).limit(10).all()
    
    # Get courses for the learner
    courses = db_session.query(Course).join(CourseLearner).filter(
        CourseLearner.learner_id == current_user.id
    ).all()
    
    # Fetch diagram evaluation data
    course_ids = [c.id for c in courses]
    diagram_assignments = db_session.query(DiagramAssignment).filter(
        DiagramAssignment.course_id.in_(course_ids)
    ).order_by(DiagramAssignment.created_at.desc()).all()
    
    diagram_submissions = db_session.query(DiagramSubmission).filter_by(
        learner_id=current_user.id
    ).order_by(DiagramSubmission.submitted_at.desc()).all()
    
    return render_template('learner_dashboard.html', 
                         announcements=announcements, 
                         courses=courses,
                         diagram_assignments=diagram_assignments,
                         diagram_submissions=diagram_submissions)

@app.route('/api/course/<int:course_id>/learners')
@login_required
@role_required('Guide')
def get_course_learners(course_id):
    course = db_session.query(Course).filter_by(id=course_id, guide_id=current_user.id).first()
    if not course:
        return jsonify({'success': False, 'message': 'Course not found'})
    
    learners = []
    for course_learner in course.learners:
        learner = course_learner.learner
        learners.append({
            'id': learner.id,
            'username': learner.username,
            'full_name': learner.full_name,
            'email': learner.email,
            'profile_photo': learner.profile_photo,
            'is_online': learner.is_online
        })
    
    return jsonify({'success': True, 'learners': learners})

@app.route('/api/course/<int:course_id>/materials')
@login_required
@role_required('Learner')
def get_course_materials(course_id):
    """API endpoint to get materials for a specific course"""
    # Verify learner is enrolled in the course
    enrollment = db_session.query(CourseLearner).filter_by(
        course_id=course_id,
        learner_id=current_user.id
    ).first()
    
    if not enrollment:
        return jsonify({'success': False, 'message': 'Not enrolled in this course'}), 403
    
    # Get course materials
    uploads = db_session.query(Upload).filter_by(
        course_id=course_id
    ).order_by(Upload.created_at.desc()).all()
    
    materials = []
    for upload in uploads:
        materials.append({
            'id': upload.id,
            'file_name': upload.file_name,
            'file_path': upload.file_path,
            'file_type': upload.file_type,
            'file_size': upload.file_size,
            'created_at': upload.created_at.strftime('%Y-%m-%d')
        })
    
    return jsonify({'success': True, 'materials': materials})

@app.route('/api/material/<int:material_id>/mark-read', methods=['POST'])
@login_required
@role_required('Learner')
def mark_material_read(material_id):
    """Mark a material as read by the learner"""
    upload = db_session.query(Upload).get(material_id)
    
    if not upload:
        return jsonify({'success': False, 'message': 'Material not found'}), 404
    
    # Verify learner has access to this material
    enrollment = db_session.query(CourseLearner).filter_by(
        course_id=upload.course_id,
        learner_id=current_user.id
    ).first()
    
    if not enrollment:
        return jsonify({'success': False, 'message': 'Access denied'}), 403
    
    # Here you can add logic to track read status
    # For now, just return success
    return jsonify({'success': True, 'message': 'Material marked as read'})




    

@app.route('/messages')
@login_required
def messages():
    if current_user.role == 'Guide':
        # Get learners for courses where current user is the guide.
        # Specify explicit ON clauses to avoid ambiguous foreign key joins.
        learners = db_session.query(User).join(
            CourseLearner, User.id == CourseLearner.learner_id
        ).join(
            Course, Course.id == CourseLearner.course_id
        ).filter(
            Course.guide_id == current_user.id
        ).all()
        return render_template('messages.html', contacts=learners)
    elif current_user.role == 'Learner':
        # Get guides for courses where the current user is enrolled as a learner.
        guides = db_session.query(User).join(
            Course, User.id == Course.guide_id
        ).join(
            CourseLearner, CourseLearner.course_id == Course.id
        ).filter(
            CourseLearner.learner_id == current_user.id
        ).all()
        return render_template('messages.html', contacts=guides)
    elif current_user.role == 'Executive':
        # Get guides for courses created by the executive.
        # Filter out courses with NULL guide_id and ensure we only get Guide role users
        guides = db_session.query(User).join(
            Course, User.id == Course.guide_id
        ).filter(
            Course.executive_id == current_user.id,
            Course.guide_id.isnot(None),
            User.role == 'Guide'
        ).distinct().all()
        return render_template('messages.html', contacts=guides)
    
    return redirect(url_for('index'))

@app.route('/api/executive/learners-by-date')
@login_required
@role_required('Executive')
def get_learners_by_date():
    """Get count of learners added on a specific date for executive's courses"""
    try:
        date_str = request.args.get('date')
        if not date_str:
            return jsonify({'success': False, 'message': 'Date parameter required'})
        
        from datetime import datetime, timedelta
        from sqlalchemy import func, cast, Date
        
        # Parse date
        try:
            if date_str == 'today':
                target_date = datetime.now(timezone.utc).date()
            elif date_str == 'yesterday':
                target_date = (datetime.now(timezone.utc) - timedelta(days=1)).date()
            elif 'week' in date_str.lower():
                # Start of current week
                today = datetime.now(timezone.utc).date()
                start_of_week = today - timedelta(days=today.weekday())
                target_date = start_of_week
            elif 'month' in date_str.lower():
                # Start of current month
                today = datetime.now(timezone.utc).date()
                target_date = today.replace(day=1)
            else:
                # Try to parse as date string
                target_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        except:
            return jsonify({'success': False, 'message': 'Invalid date format. Use YYYY-MM-DD'})
        
        # Get all courses created by executive
        executive_courses = db_session.query(Course.id).filter(
            Course.executive_id == current_user.id
        ).subquery()
        
        # Count learners in executive's courses who were created on target date
        # Using User.created_at as proxy for when learner was added
        count = db_session.query(func.count(User.id)).join(
            CourseLearner, User.id == CourseLearner.learner_id
        ).filter(
            CourseLearner.course_id.in_(db_session.query(executive_courses.c.id)),
            cast(User.created_at, Date) == target_date,
            User.role == 'Learner'
        ).scalar() or 0
        
        # Get unique courses
        courses = db_session.query(Course).join(
            CourseLearner, Course.id == CourseLearner.course_id
        ).join(
            User, CourseLearner.learner_id == User.id
        ).filter(
            Course.executive_id == current_user.id,
            cast(User.created_at, Date) == target_date,
            User.role == 'Learner'
        ).distinct().all()
        
        return jsonify({
            'success': True,
            'count': count,
            'date': str(target_date),
            'courses': [{'id': c.id, 'name': c.course_name} for c in courses]
        })
    except Exception as e:
        print(f"Error fetching learners by date: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/executive/contact', methods=['POST'])
@login_required
@role_required('Executive')
def submit_executive_contact():
    """Submit contact form from executive"""
    try:
        data = request.get_json()
        issue_type = data.get('issue_type')
        details = data.get('details')
        
        if not issue_type:
            return jsonify({'success': False, 'message': 'Issue type is required'})
        
        if not details or not details.strip():
            return jsonify({'success': False, 'message': 'Details are required'})
        
        # Validate issue type
        valid_types = ['urgent_issue', 'complaint', 'suggestions', 'requirements']
        if issue_type not in valid_types:
            return jsonify({'success': False, 'message': 'Invalid issue type'})
        
        # Create contact record
        contact = ExecutiveContact(
            executive_id=current_user.id,
            issue_type=issue_type,
            details=details.strip(),
            status='pending'
        )
        db_session.add(contact)
        db_session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Your message has been submitted successfully. We will get back to you soon!'
        })
    except Exception as e:
        db_session.rollback()
        print(f"Error submitting contact: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/guide/feedback', methods=['POST'])
@login_required
@role_required('Guide')
def submit_guide_feedback():
    """Submit feedback and rating from guide"""
    try:
        data = request.get_json()
        rating = data.get('rating')
        feedback_text = data.get('feedback_text', '').strip()
        
        if not rating:
            return jsonify({'success': False, 'message': 'Rating is required'})
        
        if rating < 1 or rating > 5:
            return jsonify({'success': False, 'message': 'Rating must be between 1 and 5'})
        
        # Create feedback record
        feedback = GuideFeedback(
            guide_id=current_user.id,
            rating=rating,
            feedback_text=feedback_text if feedback_text else None
        )
        db_session.add(feedback)
        db_session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Thank you for your feedback!'
        })
    except Exception as e:
        db_session.rollback()
        print(f"Error submitting feedback: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/guide/suggestion', methods=['POST'])
@login_required
@role_required('Guide')
def submit_guide_suggestion():
    """Submit suggestion from guide"""
    try:
        data = request.get_json()
        suggestion_text = data.get('suggestion_text', '').strip()
        
        if not suggestion_text:
            return jsonify({'success': False, 'message': 'Suggestion text is required'})
        
        # Create suggestion record
        suggestion = GuideSuggestion(
            guide_id=current_user.id,
            suggestion_text=suggestion_text,
            status='pending'
        )
        db_session.add(suggestion)
        db_session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Thank you for your suggestion! We will review it.'
        })
    except Exception as e:
        db_session.rollback()
        print(f"Error submitting suggestion: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/learner/suggestion', methods=['POST'])
@login_required
@role_required('Learner')
def submit_learner_suggestion():
    """Submit suggestion from learner"""
    try:
        data = request.get_json()
        suggestion_text = data.get('suggestion_text', '').strip()
        
        if not suggestion_text:
            return jsonify({'success': False, 'message': 'Suggestion text is required'})
        
        # Create suggestion record
        suggestion = LearnerSuggestion(
            learner_id=current_user.id,
            suggestion_text=suggestion_text,
            status='pending'
        )
        db_session.add(suggestion)
        db_session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Thank you for your suggestion! We will review it.'
        })
    except Exception as e:
        db_session.rollback()
        print(f"Error submitting suggestion: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/messages/<int:user_id>')
@login_required
def get_messages(user_id):
    messages = db_session.query(Message).filter(
        ((Message.sender_id == current_user.id) & (Message.receiver_id == user_id)) |
        ((Message.sender_id == user_id) & (Message.receiver_id == current_user.id))
    ).order_by(Message.created_at).all()
    
    return jsonify([{
        'id': msg.id,
        'sender_id': msg.sender_id,
        'receiver_id': msg.receiver_id,
        'content': msg.content,
        'created_at': msg.created_at.isoformat(),
        'is_read': msg.is_read
    } for msg in messages])

@app.route('/download/<path:filename>')
@login_required
def download_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)

@app.route('/forgot-password', methods=['GET'])
def forgot_password():
    """Render forgot password page"""
    return render_template('forgot_password.html')

@app.route('/forgot-password/send-otp', methods=['POST'])
def send_forgot_password_otp():
    """Send OTP to user's email for password reset"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip()
        
        if not email:
            return jsonify({'success': False, 'message': 'Email is required'})
        
        # Check if user exists
        user = db_session.query(User).filter_by(email=email).first()
        
        if not user:
            return jsonify({'success': False, 'message': 'No account found with this email'})
        
        # Delete any existing unused OTPs for this user
        db_session.query(OTP).filter_by(
            user_id=user.id,
            purpose='password_reset',
            is_used=False
        ).delete()
        
        # Generate new OTP
        otp_code = generate_otp()
        expires_at = datetime.now(timezone.utc) + timedelta(minutes=10)
        
        new_otp = OTP(
            user_id=user.id,
            otp_code=otp_code,
            purpose='password_reset',
            expires_at=expires_at
        )
        db_session.add(new_otp)
        db_session.commit()
        
        # Send email with OTP
        email_body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; padding: 20px; background: #f8fafc;">
            <div style="max-width: 600px; margin: 0 auto; background: white; padding: 40px; border-radius: 15px; box-shadow: 0 4px 20px rgba(0,0,0,0.1);">
                <h2 style="color: #667eea; text-align: center; margin-bottom: 20px;">🔑 Password Reset Request</h2>
                
                <p style="color: #334155; font-size: 16px; line-height: 1.6;">
                    Hello {user.full_name or 'there'},
                </p>
                
                <p style="color: #334155; font-size: 16px; line-height: 1.6;">
                    We received a request to reset your password. Use the verification code below:
                </p>
                
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                            padding: 30px; 
                            border-radius: 12px; 
                            text-align: center; 
                            margin: 30px 0;">
                    <p style="color: rgba(255,255,255,0.9); 
                               font-size: 14px; 
                               margin: 0 0 10px 0; 
                               text-transform: uppercase; 
                               letter-spacing: 1px;">
                        Your Verification Code
                    </p>
                    <p style="color: #fff; 
                               font-size: 42px; 
                               font-weight: 700; 
                               letter-spacing: 8px; 
                               margin: 0;
                               text-shadow: 0 2px 10px rgba(0,0,0,0.3);">
                        {otp_code}
                    </p>
                </div>
                
                <div style="background: #fff3cd; 
                            border-left: 4px solid #ffc107; 
                            padding: 15px; 
                            border-radius: 8px; 
                            margin: 20px 0;">
                    <p style="color: #856404; margin: 0; font-size: 14px;">
                        ⚠️ <strong>Important:</strong> This code expires in 10 minutes.
                    </p>
                </div>
                
                <p style="color: #334155; font-size: 16px; line-height: 1.6;">
                    If you didn't request a password reset, please ignore this email or contact support if you have concerns.
                </p>
                
                <hr style="border: none; border-top: 1px solid #e2e8f0; margin: 30px 0;">
                
                <p style="color: #94a3b8; font-size: 12px; text-align: center;">
                    This is an automated message from Learning Platform. Please do not reply to this email.
                </p>
            </div>
        </body>
        </html>
        """
        
        if send_email(user.email, "Password Reset - Verification Code", email_body):
            return jsonify({
                'success': True, 
                'message': 'Verification code sent to your email'
            })
        else:
            return jsonify({
                'success': False, 
                'message': 'Failed to send email. Please try again.'
            })
            
    except Exception as e:
        db_session.rollback()
        print(f"Error sending OTP: {str(e)}")
        return jsonify({
            'success': False, 
            'message': 'An error occurred. Please try again.'
        })

@app.route('/forgot-password/verify-otp', methods=['POST'])
def verify_forgot_password_otp():
    """Verify OTP for password reset"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip()
        otp_code = data.get('otp', '').strip()
        
        if not email or not otp_code:
            return jsonify({'success': False, 'message': 'Email and OTP are required'})
        
        if len(otp_code) != 6:
            return jsonify({'success': False, 'message': 'Invalid OTP format'})
        
        # Find user
        user = db_session.query(User).filter_by(email=email).first()
        
        if not user:
            return jsonify({'success': False, 'message': 'User not found'})
        
        # Find valid OTP
        otp = db_session.query(OTP).filter_by(
            user_id=user.id,
            otp_code=otp_code,
            purpose='password_reset',
            is_used=False
        ).first()
        
        if not otp:
            return jsonify({'success': False, 'message': 'Invalid or expired verification code'})
        
        # Check if OTP is expired - handle timezone-aware and timezone-naive datetimes
        now = datetime.now(timezone.utc)
        expires_at = otp.expires_at
        
        # Make expires_at timezone-aware if it's naive (assume UTC)
        if expires_at.tzinfo is None:
            expires_at = expires_at.replace(tzinfo=timezone.utc)
        
        if expires_at < now:
            return jsonify({'success': False, 'message': 'Verification code has expired'})
        
        # Mark OTP as used (but don't delete yet - we'll verify it again during reset)
        # This prevents reuse but allows password reset
        return jsonify({
            'success': True, 
            'message': 'Verification code verified successfully'
        })
        
    except Exception as e:
        print(f"Error verifying OTP: {str(e)}")
        return jsonify({
            'success': False, 
            'message': 'An error occurred. Please try again.'
        })

@app.route('/forgot-password/reset', methods=['POST'])
def reset_password():
    """Reset user password after OTP verification"""
    try:
        data = request.get_json()
        email = data.get('email', '').strip()
        new_password = data.get('new_password', '').strip()
        
        if not email or not new_password:
            return jsonify({'success': False, 'message': 'Email and new password are required'})
        
        if len(new_password) < 8:
            return jsonify({'success': False, 'message': 'Password must be at least 8 characters'})
        
        # Find user
        user = db_session.query(User).filter_by(email=email).first()
        
        if not user:
            return jsonify({'success': False, 'message': 'User not found'})
        
        # Verify that there's a valid unused OTP (extra security check)
        # Fetch OTP and check expiration in Python to handle timezone issues
        otp = db_session.query(OTP).filter_by(
            user_id=user.id,
            purpose='password_reset',
            is_used=False
        ).first()
        
        if not otp:
            return jsonify({
                'success': False, 
                'message': 'Invalid session. Please request a new verification code.'
            })
        
        # Check expiration - handle timezone-aware and timezone-naive datetimes
        now = datetime.now(timezone.utc)
        expires_at = otp.expires_at
        
        # Make expires_at timezone-aware if it's naive (assume UTC)
        if expires_at.tzinfo is None:
            expires_at = expires_at.replace(tzinfo=timezone.utc)
        
        if expires_at < now:
            return jsonify({
                'success': False, 
                'message': 'Verification code has expired. Please request a new one.'
            })
        
        # Update password
        user.password_hash = generate_password_hash(new_password, method='pbkdf2:sha256')
        
        # Mark OTP as used
        otp.is_used = True
        
        db_session.commit()
        
        # Send confirmation email
        confirmation_email_body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; padding: 20px; background: #f8fafc;">
            <div style="max-width: 600px; margin: 0 auto; background: white; padding: 40px; border-radius: 15px; box-shadow: 0 4px 20px rgba(0,0,0,0.1);">
                <h2 style="color: #56ab2f; text-align: center; margin-bottom: 20px;">✅ Password Reset Successful</h2>
                
                <p style="color: #334155; font-size: 16px; line-height: 1.6;">
                    Hello {user.full_name or 'there'},
                </p>
                
                <p style="color: #334155; font-size: 16px; line-height: 1.6;">
                    Your password has been successfully reset. You can now login with your new password.
                </p>
                
                <div style="text-align: center; margin: 30px 0;">
                    <a href="http://127.0.0.1:5000/login" 
                       style="display: inline-block;
                              background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                              color: white;
                              padding: 15px 40px;
                              text-decoration: none;
                              border-radius: 12px;
                              font-weight: 600;
                              font-size: 16px;
                              box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);">
                        Login Now
                    </a>
                </div>
                
                <div style="background: #fff3cd; 
                            border-left: 4px solid #ffc107; 
                            padding: 15px; 
                            border-radius: 8px; 
                            margin: 20px 0;">
                    <p style="color: #856404; margin: 0; font-size: 14px;">
                        ⚠️ If you didn't make this change, please contact support immediately.
                    </p>
                </div>
                
                <hr style="border: none; border-top: 1px solid #e2e8f0; margin: 30px 0;">
                
                <p style="color: #94a3b8; font-size: 12px; text-align: center;">
                    This is an automated message from Learning Platform. Please do not reply to this email.
                </p>
            </div>
        </body>
        </html>
        """
        
        send_email(user.email, "Password Reset Successful", confirmation_email_body)
        
        return jsonify({
            'success': True, 
            'message': 'Password reset successfully! You can now login.'
        })
        
    except Exception as e:
        db_session.rollback()
        print(f"Error resetting password: {str(e)}")
        return jsonify({
            'success': False, 
            'message': 'An error occurred. Please try again.'
        })

@app.route('/logout')
@login_required
def logout():
    # Clear session data from database
    if current_user.is_authenticated:
        current_user.is_online = False
        current_user.session_id = None  # Clear session ID to invalidate session
        current_user.last_activity = None
        db_session.commit()
    
    # Logout user and clear session
    logout_user()
    session.clear()  # Clear all session data
    
    flash('You have been logged out successfully.', 'success')
    return redirect(url_for('index'))

@socketio.on('connect')
def handle_connect():
    if current_user.is_authenticated:
        current_user.is_online = True
        db_session.commit()
        join_room(f"user_{current_user.id}")

@socketio.on('disconnect')
def handle_disconnect():
    if current_user.is_authenticated:
        current_user.is_online = False
        current_user.last_seen = datetime.now(timezone.utc)
        db_session.commit()

@socketio.on('send_message')
def handle_send_message(data):
    receiver_id = data.get('receiver_id')
    content = data.get('content')
    
    if not receiver_id or not content:
        return
    
    message = Message(
        sender_id=current_user.id,
        receiver_id=receiver_id,
        content=content
    )
    db_session.add(message)
    db_session.commit()
    
    emit('new_message', {
        'id': message.id,
        'sender_id': message.sender_id,
        'receiver_id': message.receiver_id,
        'content': message.content,
        'created_at': message.created_at.isoformat()
    }, room=f"user_{receiver_id}")
    
    emit('new_message', {
        'id': message.id,
        'sender_id': message.sender_id,
        'receiver_id': message.receiver_id,
        'content': message.content,
        'created_at': message.created_at.isoformat()
    }, room=f"user_{current_user.id}")

@socketio.on('mark_as_read')
def handle_mark_as_read(data):
    message_id = data.get('message_id')
    message = db_session.query(Message).get(message_id)
    if message and message.receiver_id == current_user.id:
        message.is_read = True
        db_session.commit()

@app.teardown_appcontext
def shutdown_session(exception=None):
    db_session.remove()




########################################################
###### Assignment Routes #################################
########################################################
"""
Add these routes to app.py
Make sure to:
1. Import evaluation module: from evaluation import *
2. Add database models from evaluation.py
3. Import required libraries for PDF/CSV generation
"""

import json
import io
from datetime import datetime
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
import csv

# Configure upload folder for assignment images
ASSIGNMENT_UPLOAD_FOLDER = 'static/assignment_uploads'
os.makedirs(ASSIGNMENT_UPLOAD_FOLDER, exist_ok=True)

# ========================================
# GUIDE ASSIGNMENT ROUTES
# ========================================

@app.route('/guide/assignments/create', methods=['POST'])
@login_required
@role_required('Guide')
def create_assignment():
    """Create new assignment with questions and MCQs - comprehensive error handling"""
    try:
        # Get and validate form data
        course_id = safe_form_get('course_id')
        name = safe_form_get('name')
        instructions = safe_form_get('instructions', '').strip()
        deadline_str = safe_form_get('deadline')
        evaluation_model = safe_form_get('evaluation_model', 'gemini')
        
        question_texts = request.form.getlist('question_text[]')
        question_marks = request.form.getlist('question_marks[]')
        
        # Validation checks
        if not course_id:
            return jsonify({
                'success': False, 
                'message': 'Please select a course',
                'field': 'course_id'
            }), 400
        
        if not name:
            return jsonify({
                'success': False, 
                'message': 'Assignment name is required',
                'field': 'name'
            }), 400
        
        # Validate name length
        if len(name.strip()) < 3:
            return jsonify({
                'success': False, 
                'message': 'Assignment name must be at least 3 characters long',
                'field': 'name'
            }), 400
        
        if len(name.strip()) > 200:
            return jsonify({
                'success': False, 
                'message': 'Assignment name must be less than 200 characters',
                'field': 'name'
            }), 400
        
        if not deadline_str:
            return jsonify({
                'success': False, 
                'message': 'Deadline is required',
                'field': 'deadline'
            }), 400
        
        # Validate questions
        if not question_texts or len(question_texts) == 0:
            return jsonify({
                'success': False, 
                'message': 'At least one question is required'
            }), 400
        
        # Validate question texts are not empty
        empty_questions = [i+1 for i, q in enumerate(question_texts) if not q.strip()]
        if empty_questions:
            return jsonify({
                'success': False, 
                'message': f'Question {empty_questions[0]} text cannot be empty'
            }), 400
        
        # Validate marks
        if len(question_marks) != len(question_texts):
            return jsonify({
                'success': False, 
                'message': 'Marks must be provided for all questions'
            }), 400
        
        # Validate marks are positive numbers
        try:
            marks_list = [float(m) for m in question_marks]
            if any(m <= 0 for m in marks_list):
                return jsonify({
                    'success': False, 
                    'message': 'All question marks must be greater than 0'
                }), 400
        except (ValueError, TypeError):
            return jsonify({
                'success': False, 
                'message': 'Invalid marks format. Please enter valid numbers'
            }), 400
        
        # Validate course_id is integer
        try:
            course_id_int = int(course_id)
        except (ValueError, TypeError):
            return jsonify({
                'success': False, 
                'message': 'Invalid course selected',
                'field': 'course_id'
            }), 400
        
        # Verify course belongs to guide
        course = db_session.query(Course).filter_by(
            id=course_id_int,
            guide_id=current_user.id
        ).first()
        
        if not course:
            return jsonify({
                'success': False, 
                'message': 'Course not found or you do not have access to this course',
                'field': 'course_id'
            }), 403
        
        # Validate deadline format and check it's in the future
        try:
            deadline = datetime.strptime(deadline_str, '%Y-%m-%dT%H:%M')
            if deadline <= datetime.now():
                return jsonify({
                    'success': False, 
                    'message': 'Deadline must be in the future',
                    'field': 'deadline'
                }), 400
        except ValueError:
            return jsonify({
                'success': False, 
                'message': 'Invalid deadline format. Please use the date picker',
                'field': 'deadline'
            }), 400
        
        # Validate evaluation model
        valid_models = ['gemini', 'deepseek', 'groq']
        if evaluation_model not in valid_models:
            evaluation_model = 'gemini'  # Default fallback
        
        # Calculate total marks
        total_marks = sum(marks_list)
        
        # Check for duplicate assignment name in same course
        existing = db_session.query(Assignment).filter_by(
            course_id=course_id_int,
            name=name.strip()
        ).first()
        
        if existing:
            return jsonify({
                'success': False, 
                'message': f'An assignment with name "{name}" already exists in this course',
                'field': 'name'
            }), 409
        
        # Create assignment
        assignment = Assignment(
            course_id=course.id,
            guide_id=current_user.id,
            name=name.strip(),
            instructions=instructions if instructions else None,
            deadline=deadline,
            total_marks=total_marks,
            evaluation_model=evaluation_model
        )
        db_session.add(assignment)
        db_session.flush()
        
        # Add questions with MCQs
        for idx, (q_text, q_marks) in enumerate(zip(question_texts, marks_list)):
            question = AssignmentQuestion(
                assignment_id=assignment.id,
                question_text=q_text.strip(),
                marks=q_marks,
                order_num=idx + 1
            )
            db_session.add(question)
            db_session.flush()
            
            # Generate MCQ about question topic (with error handling)
            try:
                mcq_data = generate_question_mcq(q_text.strip(), evaluation_model)
                if mcq_data:
                    question_mcq = QuestionMCQ(
                        question_id=question.id,
                        mcq_question=mcq_data['question'],
                        option_a=mcq_data['option_a'],
                        option_b=mcq_data['option_b'],
                        option_c=mcq_data['option_c'],
                        option_d=mcq_data['option_d'],
                        correct_option=mcq_data['correct']
                    )
                    db_session.add(question_mcq)
            except Exception as mcq_error:
                print(f"Warning: Could not generate MCQ for question {idx+1}: {mcq_error}")
                # Continue without MCQ - assignment creation should not fail
        
        db_session.commit()
        
        return jsonify({
            'success': True, 
            'message': f'Assignment "{name}" created successfully with {len(question_texts)} question(s)!',
            'assignment_id': assignment.id
        }), 201
        
    except ValueError as e:
        db_session.rollback()
        return jsonify({
            'success': False, 
            'message': f'Invalid input: {str(e)}',
            'field': 'general'
        }), 400
    except Exception as e:
        db_session.rollback()
        import traceback
        error_trace = traceback.format_exc()
        print(f"Error creating assignment: {error_trace}")
        return jsonify({
            'success': False, 
            'message': 'An error occurred while creating the assignment. Please try again.',
            'field': 'general',
            'error': str(e) if app.debug else None
        }), 500

@app.route('/guide/assignments/list')
@login_required
@role_required('Guide')
def list_assignments():
    """Get all assignments for guide's courses"""
    try:
        assignments = db_session.query(Assignment).filter_by(
            guide_id=current_user.id
        ).order_by(Assignment.created_at.desc()).all()
        
        result = []
        for assignment in assignments:
            submissions_count = db_session.query(AssignmentSubmission).filter_by(
                assignment_id=assignment.id
            ).count()
            
            result.append({
                'id': assignment.id,
                'name': assignment.name,
                'course_name': assignment.course.course_name,
                'deadline': assignment.deadline.strftime('%B %d, %Y at %I:%M %p'),
                'total_marks': assignment.total_marks,
                'evaluation_model': assignment.evaluation_model.upper(),
                'submissions_count': submissions_count,
                'scores_visible': assignment.scores_visible
            })
        
        return jsonify({'success': True, 'assignments': result})
        
    except Exception as e:
        print(f"Error listing assignments: {str(e)}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/guide/assignments/<int:assignment_id>/submissions')
@login_required
@role_required('Guide')
def get_assignment_submissions(assignment_id):
    """Get all submissions for an assignment with detailed evaluation"""
    try:
        # Rollback any failed transaction first
        try:
            db_session.rollback()
        except:
            pass
        
        assignment = db_session.query(Assignment).filter_by(
            id=assignment_id,
            guide_id=current_user.id
        ).first()
        
        if not assignment:
            return jsonify({'success': False, 'message': 'Assignment not found'})
        
        submissions = db_session.query(AssignmentSubmission).filter_by(
            assignment_id=assignment_id
        ).all()
        
        result = []
        for submission in submissions:
            answers_data = []
            for answer in submission.answers:
                # Pre-fetch question to avoid lazy loading issues
                try:
                    question = db_session.query(AssignmentQuestion).filter_by(id=answer.question_id).first()
                    if not question:
                        print(f"Warning: Question {answer.question_id} not found for answer {answer.id}")
                        continue
                    question_text = question.question_text
                    question_marks = question.marks
                except Exception as q_error:
                    db_session.rollback()
                    print(f"Error fetching question for answer {answer.id}: {q_error}")
                    # Try again after rollback
                    try:
                        question = db_session.query(AssignmentQuestion).filter_by(id=answer.question_id).first()
                        if question:
                            question_text = question.question_text
                            question_marks = question.marks
                        else:
                            continue
                    except:
                        print(f"Failed to fetch question after rollback")
                        continue
                images = json.loads(answer.image_paths) if answer.image_paths else []
                covered = json.loads(answer.covered_points) if answer.covered_points else []
                missing = json.loads(answer.missing_points) if answer.missing_points else []
                
                # Add MCQ result - handle missing mcq_type column gracefully
                mcq_result = None
                try:
                    mcq_response = db_session.query(LearnerMCQResponse).filter_by(
                        answer_id=answer.id
                    ).first()
                
                    if mcq_response:
                        # Handle missing mcq_type column gracefully
                        mcq_type = getattr(mcq_response, 'mcq_type', 'answer')
                        mcq_result = {
                            'type': mcq_type,
                            'is_correct': mcq_response.is_correct,
                            'time_taken': mcq_response.time_taken,
                            'learner_answer': mcq_response.learner_answer,
                            'correct_answer': mcq_response.correct_option
                        }
                except Exception as mcq_error:
                    # Rollback on error
                    try:
                        db_session.rollback()
                    except:
                        pass
                    
                    # If query fails due to missing column or transaction error, skip MCQ data
                    if 'mcq_type' in str(mcq_error).lower() or 'undefinedcolumn' in str(mcq_error).lower() or 'infailed' in str(mcq_error).lower():
                        # Skip MCQ data if table structure is different
                        mcq_result = None
                        try:
                            db_session.rollback()
                        except:
                            pass
                    else:
                        print(f"Error fetching MCQ response: {mcq_error}")
                        mcq_result = None
                
                answers_data.append({
                    'question_text': question_text,
                    'max_marks': question_marks,
                    'images': images,
                    'extracted_text': answer.extracted_text,
                    'relevance_score': answer.relevance_score,
                    'grammar_score': answer.grammar_score,
                    'size_score': answer.size_score,
                    'uniqueness_score': answer.uniqueness_score,
                    'total_score': answer.total_score,
                    'answer_id': answer.id,
                    'feedback': answer.feedback,
                    'covered_points': covered,
                    'missing_points': missing,
                    'word_count': answer.word_count,
                    'mcq_result': mcq_result
                })
            
            # Count MCQ responses - handle missing mcq_type column
            mcq_correct = 0
            mcq_total = 0
            try:
                answer_ids = [a.id for a in submission.answers]
                if answer_ids:
                    mcq_correct = sum(1 for answer in submission.answers 
                         if db_session.query(LearnerMCQResponse).filter_by(
                             answer_id=answer.id, is_correct=True
                         ).first())
                    mcq_total = db_session.query(LearnerMCQResponse).filter(
                        LearnerMCQResponse.answer_id.in_(answer_ids)
                    ).count()
            except Exception as mcq_count_error:
                # Rollback on error
                try:
                    db_session.rollback()
                except:
                    pass
                
                # If query fails due to missing column or transaction error, use raw SQL
                if 'mcq_type' in str(mcq_count_error).lower() or 'undefinedcolumn' in str(mcq_count_error).lower() or 'infailed' in str(mcq_count_error).lower():
                    try:
                        from sqlalchemy import text
                        answer_ids = [a.id for a in submission.answers]
                        if answer_ids:
                            count_result = db_session.execute(
                                text("SELECT COUNT(*) FROM learner_mcq_responses WHERE answer_id = ANY(:answer_ids) AND is_correct = true"),
                                {'answer_ids': answer_ids}
                            ).scalar()
                            mcq_correct = count_result or 0
                            count_result = db_session.execute(
                                text("SELECT COUNT(*) FROM learner_mcq_responses WHERE answer_id = ANY(:answer_ids)"),
                                {'answer_ids': answer_ids}
                            ).scalar()
                            mcq_total = count_result or 0
                    except Exception as raw_error:
                        try:
                            db_session.rollback()
                        except:
                            pass
                        print(f"Error counting MCQs with raw SQL: {raw_error}")
                        mcq_correct = 0
                        mcq_total = 0
                else:
                    print(f"Error counting MCQ responses: {mcq_count_error}")
                    mcq_correct = 0
                    mcq_total = 0

            result.append({
                'submission_id': submission.id,
                'learner_name': submission.learner.full_name or submission.learner.username,
                'learner_id': submission.learner.id,
                'submitted_at': submission.submitted_at.strftime('%B %d, %Y at %I:%M %p'),
                'total_score': submission.total_score,
                'is_flagged': submission.is_flagged,
                'flag_reason': submission.flag_reason,
                'mcq_summary': {
                    'correct': mcq_correct,
                    'total': mcq_total,
                    'all_correct': mcq_correct == mcq_total if mcq_total > 0 else False
                } if mcq_total > 0 else None,
                'answers': answers_data
            })
        
        return jsonify({
            'success': True,
            'submissions': result,
            'assignment': {
                'name': assignment.name,
                'total_marks': assignment.total_marks
            }
        })
        
    except Exception as e:
        # Always rollback on exception
        try:
            db_session.rollback()
        except:
            pass
        
        import traceback
        error_trace = traceback.format_exc()
        print(f"Error getting submissions: {error_trace}")
        return jsonify({'success': False, 'message': f'Error loading submissions: {str(e)}'})


def build_scoreboard_for_course(course):
    """Return scoreboard dataset for a given guide-owned course."""
    if not course:
        return None

    course_id = course.id
    learners = db_session.query(User).join(
        CourseLearner, CourseLearner.learner_id == User.id
    ).filter(
        CourseLearner.course_id == course_id
    ).order_by(User.full_name).all()

    assignment_total = db_session.query(
        func.coalesce(func.sum(Assignment.total_marks), 0)
    ).filter(
        Assignment.course_id == course_id
    ).scalar() or 0.0

    assignment_scores = dict(
        db_session.query(
            AssignmentSubmission.learner_id,
            func.coalesce(func.sum(AssignmentSubmission.total_score), 0)
        ).join(Assignment).filter(
            Assignment.course_id == course_id
        ).group_by(
            AssignmentSubmission.learner_id
        ).all()
    )

    mcq_total = db_session.query(
        func.coalesce(func.sum(func.coalesce(TestMCQ.total_marks, TestMCQ.num_questions)), 0)
    ).filter(
        TestMCQ.course_id == course_id
    ).scalar() or 0.0

    mcq_scores = dict(
        db_session.query(
            TestMCQAttempt.learner_id,
            func.coalesce(func.sum(TestMCQAttempt.marks_obtained), 0)
        ).join(TestMCQ).filter(
            TestMCQ.course_id == course_id,
            TestMCQAttempt.completed_at.isnot(None)
        ).group_by(
            TestMCQAttempt.learner_id
        ).all()
    )

    columns = db_session.query(CustomScoreboardColumn).filter_by(
        course_id=course_id
    ).order_by(
        CustomScoreboardColumn.created_at.asc()
    ).all()

    column_values = []
    column_ids = [col.id for col in columns]
    if column_ids:
        column_values = db_session.query(CustomScoreboardValue).filter(
            CustomScoreboardValue.column_id.in_(column_ids)
        ).all()

    value_map = defaultdict(dict)
    for value in column_values:
        value_map[value.column_id][value.learner_id] = float(value.mark or 0.0)

    students = []
    for learner in learners:
        custom_scores = []
        for column in columns:
            custom_scores.append({
                'column_id': column.id,
                'mark': value_map[column.id].get(learner.id, 0.0)
            })

        students.append({
            'learner_id': learner.id,
            'learner_name': learner.full_name or learner.username,
            'assignment_marks': float(assignment_scores.get(learner.id, 0.0)),
            'mcq_marks': float(mcq_scores.get(learner.id, 0.0)),
            'custom_scores': custom_scores
        })

    return {
        'course': {
            'id': course.id,
            'name': course.course_name
        },
        'assignment_total': float(assignment_total),
        'mcq_total': float(mcq_total),
        'columns': [{
            'id': column.id,
            'name': column.name,
            'total_marks': float(column.total_marks)
        } for column in columns],
        'students': students
    }


@app.route('/guide/scoreboard/course/<int:course_id>/data')
@login_required
@role_required('Guide')
def get_universal_scoreboard_data(course_id):
    """Return universal scoreboard dataset for a course."""
    course = db_session.query(Course).filter_by(
        id=course_id,
        guide_id=current_user.id
    ).first()

    if not course:
        return jsonify({'success': False, 'error': 'Course not found'}), 404

    dataset = build_scoreboard_for_course(course)
    if dataset is None:
        return jsonify({'success': False, 'error': 'Could not build dataset'}), 500

    return jsonify({'success': True, **dataset})


@app.route('/guide/scoreboard/columns', methods=['POST'])
@login_required
@role_required('Guide')
def create_custom_scoreboard_column():
    """Create a new custom column for the universal scoreboard."""
    payload = request.get_json(silent=True) or {}
    course_id = payload.get('course_id')
    name = (payload.get('name') or '').strip()
    total_marks = payload.get('total_marks')

    try:
        course_id = int(course_id)
    except (TypeError, ValueError):
        return jsonify({'success': False, 'error': 'Invalid course selected'}), 400

    if not name:
        return jsonify({'success': False, 'error': 'Course and column name are required'}), 400

    try:
        total_marks = float(total_marks)
    except (TypeError, ValueError):
        return jsonify({'success': False, 'error': 'Total marks must be a valid number'}), 400

    if total_marks < 0:
        return jsonify({'success': False, 'error': 'Total marks cannot be negative'}), 400

    course = db_session.query(Course).filter_by(
        id=course_id,
        guide_id=current_user.id
    ).first()

    if not course:
        return jsonify({'success': False, 'error': 'Course not found'}), 404

    column = CustomScoreboardColumn(
        course_id=course_id,
        name=name,
        total_marks=total_marks
    )
    db_session.add(column)
    db_session.commit()

    return jsonify({
        'success': True,
        'column': {
            'id': column.id,
            'name': column.name,
            'total_marks': float(column.total_marks)
        }
    })


@app.route('/guide/scoreboard/columns/<int:column_id>/value', methods=['POST'])
@login_required
@role_required('Guide')
def update_custom_scoreboard_value(column_id):
    """Update or insert a mark for a learner in a custom column."""
    payload = request.get_json(silent=True) or {}
    learner_id = payload.get('learner_id')
    mark = payload.get('mark')

    try:
        learner_id = int(learner_id)
    except (TypeError, ValueError):
        return jsonify({'success': False, 'error': 'Invalid learner selected'}), 400

    if mark is None:
        return jsonify({'success': False, 'error': 'Learner and mark are required'}), 400

    try:
        mark = float(mark)
    except (TypeError, ValueError):
        return jsonify({'success': False, 'error': 'Mark must be a number'}), 400

    column = db_session.query(CustomScoreboardColumn).filter_by(id=column_id).first()
    if not column or column.course.guide_id != current_user.id:
        return jsonify({'success': False, 'error': 'Column not found'}), 404

    if mark < 0 or mark > float(column.total_marks):
        return jsonify({'success': False, 'error': 'Mark must be between 0 and column total'}), 400

    enrollment = db_session.query(CourseLearner).filter_by(
        course_id=column.course_id,
        learner_id=learner_id
    ).first()
    if not enrollment:
        return jsonify({'success': False, 'error': 'Learner is not enrolled in this course'}), 400

    value = db_session.query(CustomScoreboardValue).filter_by(
        column_id=column_id,
        learner_id=learner_id
    ).first()

    if not value:
        value = CustomScoreboardValue(
            column_id=column_id,
            learner_id=learner_id,
            mark=mark
        )
        db_session.add(value)
    else:
        value.mark = mark

    db_session.commit()

    return jsonify({'success': True, 'mark': float(mark)})


@app.route('/guide/scoreboard/course/<int:course_id>/pdf')
@login_required
@role_required('Guide')
def download_scoreboard_pdf(course_id):
    """Return a PDF report for the universal scoreboard."""
    course = db_session.query(Course).filter_by(
        id=course_id,
        guide_id=current_user.id
    ).first()

    if not course:
        return "Access denied", 403

    dataset = build_scoreboard_for_course(course)
    if not dataset:
        return "No data available", 404

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=30, rightMargin=30, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    elements = []

    title = Paragraph(f"<b>Universal Scoreboard</b><br/>{course.course_name}", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 12))

    summary_text = f"""
    <b>Assignments Total:</b> {dataset['assignment_total']:.1f}<br/>
    <b>MCQ Tests Total:</b> {dataset['mcq_total']:.1f}<br/>
    <b>Learners:</b> {len(dataset['students'])}
    """
    elements.append(Paragraph(summary_text, styles['Normal']))
    elements.append(Spacer(1, 16))

    table_data = []
    header = [
        'Learner',
        f"Assignments ({dataset['assignment_total']:.1f})",
        f"MCQ Tests ({dataset['mcq_total']:.1f})"
    ]
    for column in dataset['columns']:
        header.append(f"{column['name']} ({column['total_marks']:.1f})")
    header.append('Total')
    table_data.append(header)

    for student in dataset['students']:
        row = [
            student['learner_name'],
            f"{student['assignment_marks']:.1f}",
            f"{student['mcq_marks']:.1f}"
        ]
        custom_total = 0.0
        for custom in student['custom_scores']:
            row.append(f"{float(custom['mark']):.1f}")
            custom_total += float(custom['mark'])
        total = student['assignment_marks'] + student['mcq_marks'] + custom_total
        row.append(f"{total:.1f}")
        table_data.append(row)

    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f2937')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
        ('BOX', (0, 0), (-1, -1), 1, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey])
    ]))
    elements.append(table)

    doc.build(elements)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{course.course_name}_Universal_Scoreboard.pdf",
        mimetype='application/pdf'
    )


@app.route('/guide/assignments/answers/<int:answer_id>/score', methods=['POST'])
@login_required
@role_required('Guide')
def update_assignment_answer_score(answer_id):
    """Allow guides to override the final mark for a single answer."""
    try:
        payload = request.get_json(silent=True) or {}
        total_score = payload.get('total_score')
        if total_score is None:
            return jsonify({'success': False, 'error': 'Score is required'}), 400

        try:
            total_score = float(total_score)
        except (TypeError, ValueError):
            return jsonify({'success': False, 'error': 'Invalid score value'}), 400

        if total_score < 0:
            return jsonify({'success': False, 'error': 'Score cannot be negative'}), 400

        answer = db_session.query(AssignmentAnswer).filter_by(id=answer_id).first()
        if not answer:
            return jsonify({'success': False, 'error': 'Answer not found'}), 404

        assignment = answer.submission.assignment
        if assignment.guide_id != current_user.id:
            return jsonify({'success': False, 'error': 'Access denied'}), 403

        max_marks = answer.question.marks if answer.question and answer.question.marks is not None else None
        if max_marks is not None and total_score > float(max_marks):
            return jsonify({
                'success': False,
                'error': 'Score cannot exceed question max marks'
            }), 400

        answer.total_score = total_score
        db_session.commit()

        return jsonify({'success': True, 'total_score': total_score})
    except Exception as e:
        db_session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/learner/assignments/mcq/<int:answer_id>', methods=['GET'])
@login_required
@role_required('Learner')
def get_mcq_for_answer(answer_id):
    """Get random MCQ (question or answer type) after submission"""
    try:
        # Rollback any failed transaction first
        try:
            db_session.rollback()
        except:
            pass
        
        print(f"DEBUG get_mcq_for_answer - answer_id: {answer_id}, user_id: {current_user.id}")
        
        # Use fresh query with explicit session handling
        try:
            answer = db_session.query(AssignmentAnswer).filter_by(id=answer_id).first()
        except Exception as query_error:
            db_session.rollback()
            print(f"DEBUG: Error querying answer: {query_error}")
            answer = db_session.query(AssignmentAnswer).filter_by(id=answer_id).first()
        
        if not answer:
            print(f"DEBUG: Answer {answer_id} not found")
            return jsonify({'success': False, 'message': 'Answer not found'}), 404
        
        # Refresh to ensure we have latest data
        try:
            db_session.refresh(answer)
        except:
            pass
        
        if answer.submission.learner_id != current_user.id:
            print(f"DEBUG: Access denied - answer belongs to learner {answer.submission.learner_id}, current user is {current_user.id}")
            return jsonify({'success': False, 'message': 'Access denied'}), 403
        
        # Check if already answered - handle transaction errors
        existing = None
        try:
            existing = db_session.query(LearnerMCQResponse).filter_by(
                answer_id=answer_id
            ).first()
        except Exception as mcq_check_error:
            db_session.rollback()  # Rollback on error
            # If query fails due to missing column, use raw SQL
            if 'mcq_type' in str(mcq_check_error).lower() or 'undefinedcolumn' in str(mcq_check_error).lower() or 'infailed' in str(mcq_check_error).lower():
                try:
                    from sqlalchemy import text
                    result = db_session.execute(
                        text("SELECT id FROM learner_mcq_responses WHERE answer_id = :answer_id LIMIT 1"),
                        {'answer_id': answer_id}
                    ).fetchone()
                    if result:
                        existing = True  # Just mark as existing
                except Exception as raw_error:
                    db_session.rollback()
                    print(f"Error with raw SQL check: {raw_error}")
                    existing = None
            else:
                print(f"Error checking existing MCQ: {mcq_check_error}")
                # Try again after rollback
                try:
                    existing = db_session.query(LearnerMCQResponse).filter_by(
                        answer_id=answer_id
                    ).first()
                except:
                    existing = None
        
        if existing:
            return jsonify({'success': False, 'message': 'MCQ already answered'})
        
        import random
        mcq_type = random.choice(['question', 'answer'])
        mcq_data = None
        
        if mcq_type == 'question':
            # Get stored question MCQ - handle transaction errors
            try:
                question_mcq = db_session.query(QuestionMCQ).filter_by(
                    question_id=answer.question_id
                ).first()
            except Exception as q_error:
                db_session.rollback()
                print(f"Error querying QuestionMCQ: {q_error}")
                # Try again after rollback
                try:
                    question_mcq = db_session.query(QuestionMCQ).filter_by(
                        question_id=answer.question_id
                    ).first()
                except:
                    question_mcq = None
            
            if question_mcq:
                mcq_data = {
                    'type': 'question',
                    'question': question_mcq.mcq_question,
                    'options': {
                        'A': question_mcq.option_a,
                        'B': question_mcq.option_b,
                        'C': question_mcq.option_c,
                        'D': question_mcq.option_d
                    },
                    'correct': question_mcq.correct_option
                }
                print(f"DEBUG: Using question MCQ from database")
            else:
                mcq_type = 'answer'  # Fallback
                print(f"DEBUG: No question MCQ found, falling back to answer MCQ")
        
        if mcq_type == 'answer' or not mcq_data:
            # Generate MCQ about their answer
            try:
                evaluation_model = answer.submission.assignment.evaluation_model
            except:
                # If we can't access assignment, rollback and try again
                db_session.rollback()
                db_session.refresh(answer)
                evaluation_model = answer.submission.assignment.evaluation_model or 'gemini'
            
            print(f"DEBUG: Generating answer MCQ with model: {evaluation_model}")
            mcq_result = generate_answer_mcq(
                answer.question.question_text,
                answer.extracted_text,
                evaluation_model
            )
            
            if mcq_result:
                mcq_data = {
                    'type': 'answer',
                    'question': mcq_result['question'],
                    'options': {
                        'A': mcq_result['option_a'],
                        'B': mcq_result['option_b'],
                        'C': mcq_result['option_c'],
                        'D': mcq_result['option_d']
                    },
                    'correct': mcq_result['correct']
                }
                print(f"DEBUG: MCQ generated successfully")
            else:
                # Provide specific error message based on model
                model_name = evaluation_model or 'gemini'
                api_key_name = {
                    'gemini': 'GEMINI_API_KEY',
                    'deepseek': 'DEEPSEEK_API_KEY',
                    'groq': 'GROQ_API_KEY'
                }.get(model_name.lower(), 'GEMINI_API_KEY')
                
                error_msg = f'MCQ generation failed: {api_key_name} not configured. Please configure the API key for {model_name.upper()} model in your .env file.'
                print(f"DEBUG: {error_msg}")
                return jsonify({
                    'success': False, 
                    'message': error_msg
                }), 400
        
        if not mcq_data:
            print("DEBUG: mcq_data not set, falling back to error")
            return jsonify({
                'success': False,
                'message': 'Failed to generate MCQ. Please try again.'
            }), 500
        
        print(f"DEBUG: Returning MCQ data successfully")
        return jsonify({
            'success': True,
            'answer_id': answer_id,
            'mcq': mcq_data
        })
        
    except Exception as e:
        # Always rollback on exception
        try:
            db_session.rollback()
        except:
            pass
        
        import traceback
        error_trace = traceback.format_exc()
        print(f"MCQ fetch error: {error_trace}")
        return jsonify({'success': False, 'message': f'Error fetching MCQ: {str(e)}'}), 500


@app.route('/learner/assignments/mcq/submit', methods=['POST'])
@login_required
@role_required('Learner')
def submit_mcq_answer():
    """Submit MCQ answer"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Invalid request data'}), 400
            
        answer_id = data.get('answer_id')
        mcq_type = data.get('mcq_type')
        selected_option = data.get('selected_option')
        correct_option = data.get('correct_option')
        time_taken = data.get('time_taken', 10)
        mcq_question = data.get('mcq_question', '')
        options = data.get('options', {})
        
        if not answer_id:
            return jsonify({'success': False, 'message': 'Answer ID is required'}), 400
        
        answer = db_session.query(AssignmentAnswer).get(answer_id)
        
        if not answer:
            return jsonify({'success': False, 'message': 'Answer not found'}), 404
            
        if answer.submission.learner_id != current_user.id:
            return jsonify({'success': False, 'message': 'Access denied'}), 403
        
        # Calculate if correct
        is_correct = (selected_option == correct_option)
        
        # Check if MCQ already exists - use minimal query
        from sqlalchemy import text
        existing = None
        try:
            db_session.rollback()  # Clear any failed transaction
            result = db_session.execute(
                text("SELECT id FROM learner_mcq_responses WHERE answer_id = :answer_id LIMIT 1"),
                {'answer_id': answer_id}
            ).fetchone()
            existing = result is not None
        except Exception as check_error:
            db_session.rollback()
            print(f"Error checking existing MCQ: {check_error}")
            existing = False
        
        if existing:
            # Update existing record - only update columns we know exist
            try:
                db_session.rollback()
                db_session.execute(
                    text("""
                        UPDATE learner_mcq_responses 
                        SET learner_answer = :learner_answer,
                            is_correct = :is_correct,
                            time_taken = :time_taken
                        WHERE answer_id = :answer_id
                    """),
                    {
                        'answer_id': answer_id,
                        'learner_answer': selected_option,
                        'is_correct': is_correct,
                        'time_taken': time_taken
                    }
                )
                db_session.commit()
                print(f"DEBUG: MCQ response updated for answer_id {answer_id}")
            except Exception as update_error:
                db_session.rollback()
                print(f"Error updating MCQ: {update_error}")
                return jsonify({
                    'success': False,
                    'message': 'Failed to update MCQ response. Please try again.'
                }), 500
        else:
            # Create new record - try full insert first, then fallback to minimal
            try:
                db_session.rollback()
                # Try with all columns
                db_session.execute(
                    text("""
                        INSERT INTO learner_mcq_responses 
                        (answer_id, mcq_question, option_a, option_b, option_c, option_d, 
                         correct_option, learner_answer, is_correct, time_taken, created_at)
                        VALUES 
                        (:answer_id, :mcq_question, :option_a, :option_b, :option_c, :option_d,
                         :correct_option, :learner_answer, :is_correct, :time_taken, :created_at)
                    """),
                    {
                        'answer_id': answer_id,
                        'mcq_question': mcq_question,
                        'option_a': options.get('A', ''),
                        'option_b': options.get('B', ''),
                        'option_c': options.get('C', ''),
                        'option_d': options.get('D', ''),
                        'correct_option': correct_option,
                        'learner_answer': selected_option,
                        'is_correct': is_correct,
                        'time_taken': time_taken,
                        'created_at': datetime.now(timezone.utc)
                    }
                )
                db_session.commit()
                print(f"DEBUG: MCQ response created (full) for answer_id {answer_id}")
            except Exception as full_error:
                db_session.rollback()
                print(f"Full insert failed, trying minimal: {full_error}")
                
                # Fallback to minimal columns
                try:
                    db_session.execute(
                        text("""
                            INSERT INTO learner_mcq_responses 
                            (answer_id, learner_answer, is_correct, time_taken)
                            VALUES 
                            (:answer_id, :learner_answer, :is_correct, :time_taken)
                        """),
                        {
                            'answer_id': answer_id,
                            'learner_answer': selected_option,
                            'is_correct': is_correct,
                            'time_taken': time_taken
                        }
                    )
                    db_session.commit()
                    print(f"DEBUG: MCQ response created (minimal) for answer_id {answer_id}")
                except Exception as minimal_error:
                    db_session.rollback()
                    print(f"Error creating MCQ (minimal): {minimal_error}")
                    return jsonify({
                        'success': False,
                        'message': 'Database error. Please contact administrator to run database migrations.'
                    }), 500
        
        # Flag submission if wrong
        if not is_correct:
            submission = answer.submission
            if not submission.is_flagged:
                submission.is_flagged = True
                submission.flag_reason = "Failed MCQ verification"
            else:
                if "MCQ" not in (submission.flag_reason or ''):
                    submission.flag_reason += " | Failed MCQ verification"
            db_session.commit()
        
        return jsonify({
            'success': True,
            'is_correct': is_correct,
            'message': 'Correct!' if is_correct else 'Incorrect answer'
        })
    
    except Exception as e:
        db_session.rollback()
        import traceback
        error_trace = traceback.format_exc()
        print(f"ERROR in submit_mcq_answer: {error_trace}")
        return jsonify({
            'success': False,
            'message': f'Error submitting MCQ: {str(e)}'
        }), 500

@app.route('/guide/assignments/<int:assignment_id>/scores')
@login_required
@role_required('Guide')
def get_assignment_scores(assignment_id):
    """Get scores summary for an assignment"""
    try:
        assignment = db_session.query(Assignment).filter_by(
            id=assignment_id,
            guide_id=current_user.id
        ).first()
        
        if not assignment:
            return jsonify({'success': False, 'message': 'Assignment not found'})
        
        submissions = db_session.query(AssignmentSubmission).filter_by(
            assignment_id=assignment_id
        ).all()
        
        scores = []
        for submission in submissions:
            percentage = (submission.total_score / assignment.total_marks * 100) if assignment.total_marks > 0 else 0
            
            scores.append({
                'submission_id': submission.id,
                'learner_name': submission.learner.full_name or submission.learner.username,
                'learner_id': submission.learner.id,
                'submitted_at': submission.submitted_at.strftime('%b %d, %Y'),
                'total_score': round(submission.total_score, 2),
                'percentage': round(percentage, 1),
                'is_flagged': submission.is_flagged
            })
        
        return jsonify({
            'success': True,
            'scores': scores,
            'assignment': {
                'id': assignment.id,
                'name': assignment.name,
                'total_marks': assignment.total_marks,
                'scores_visible': assignment.scores_visible
            }
        })
        
    except Exception as e:
        print(f"Error getting scores: {str(e)}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/guide/assignments/<int:assignment_id>/visibility', methods=['POST'])
@login_required
@role_required('Guide')
def toggle_scores_visibility(assignment_id):
    """Toggle score visibility for learners"""
    try:
        assignment = db_session.query(Assignment).filter_by(
            id=assignment_id,
            guide_id=current_user.id
        ).first()
        
        if not assignment:
            return jsonify({'success': False, 'message': 'Assignment not found'})
        
        data = request.get_json()
        assignment.scores_visible = data.get('visible', False)
        db_session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Scores are now {"visible" if assignment.scores_visible else "hidden"} to learners'
        })
        
    except Exception as e:
        db_session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/guide/assignments/submission/<int:submission_id>/report')
@login_required
@role_required('Guide')
def download_individual_report(submission_id):
    """Download individual learner report as PDF"""
    try:
        submission = db_session.query(AssignmentSubmission).get(submission_id)
        
        if not submission or submission.assignment.guide_id != current_user.id:
            return "Access denied", 403
        
        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title = Paragraph(f"<b>{submission.assignment.name} - Report</b>", styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 12))
        
        # Learner info
        info = f"<b>Learner:</b> {submission.learner.full_name or submission.learner.username}<br/>"
        info += f"<b>Submitted:</b> {submission.submitted_at.strftime('%B %d, %Y at %I:%M %p')}<br/>"
        info += f"<b>Total Score:</b> {submission.total_score} / {submission.assignment.total_marks}<br/>"
        
        if submission.is_flagged:
            info += f"<b><font color='red'>Flagged:</font></b> {submission.flag_reason}<br/>"
        
        elements.append(Paragraph(info, styles['Normal']))
        elements.append(Spacer(1, 20))
        
        # Answers
        for idx, answer in enumerate(submission.answers):
            elements.append(Paragraph(f"<b>Question {idx + 1}:</b> {answer.question.question_text}", styles['Heading3']))
            elements.append(Spacer(1, 6))
            
            # Score breakdown table
            score_data = [
                ['Relevance', 'Grammar', 'Size', 'Uniqueness', 'Total'],
                [
                    f"{answer.relevance_score}/50",
                    f"{answer.grammar_score}/10",
                    f"{answer.size_score}/30",
                    f"{answer.uniqueness_score}/10",
                    f"{answer.total_score}/{answer.question.marks}"
                ]
            ]
            
            score_table = Table(score_data)
            score_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(score_table)
            elements.append(Spacer(1, 12))
            
            # Feedback
            elements.append(Paragraph(f"<b>Feedback:</b> {answer.feedback}", styles['Normal']))
            elements.append(Spacer(1, 6))
            
            # Word count
            elements.append(Paragraph(f"<i>Word Count: {answer.word_count}</i>", styles['Normal']))
            elements.append(Spacer(1, 20))
        
        doc.build(elements)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{submission.learner.username}_{submission.assignment.name}_report.pdf",
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"Error generating report: {str(e)}")
        import traceback
        traceback.print_exc()
        return "Error generating report", 500

@app.route('/guide/assignments/<int:assignment_id>/reports/pdf')
@login_required
@role_required('Guide')
def download_all_reports_pdf(assignment_id):
    """Download all reports combined in one PDF"""
    try:
        assignment = db_session.query(Assignment).filter_by(
            id=assignment_id,
            guide_id=current_user.id
        ).first()
        
        if not assignment:
            return "Access denied", 403
        
        submissions = db_session.query(AssignmentSubmission).filter_by(
            assignment_id=assignment_id
        ).all()
        
        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        
        # Main title
        title = Paragraph(f"<b>{assignment.name} - All Reports</b>", styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 20))
        
        for submission in submissions:
            # Learner section
            learner_title = Paragraph(
                f"<b>{submission.learner.full_name or submission.learner.username}</b>",
                styles['Heading2']
            )
            elements.append(learner_title)
            
            info = f"Score: {submission.total_score}/{assignment.total_marks} | "
            info += f"Submitted: {submission.submitted_at.strftime('%b %d, %Y')}"
            
            if submission.is_flagged:
                info += f" | <font color='red'>FLAGGED</font>"
            
            elements.append(Paragraph(info, styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Summary table
            summary_data = [['Question', 'Score', 'Relevance', 'Grammar', 'Size', 'Uniqueness']]
            
            for idx, answer in enumerate(submission.answers):
                summary_data.append([
                    f"Q{idx + 1}",
                    f"{answer.total_score}/{answer.question.marks}",
                    str(answer.relevance_score),
                    str(answer.grammar_score),
                    str(answer.size_score),
                    str(answer.uniqueness_score)
                ])
            
            summary_table = Table(summary_data)
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(summary_table)
            elements.append(Spacer(1, 30))
        
        doc.build(elements)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{assignment.name}_all_reports.pdf",
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"Error generating all reports: {str(e)}")
        return "Error generating reports", 500

@app.route('/guide/assignments/<int:assignment_id>/scores/csv')
@login_required
@role_required('Guide')
def download_scores_csv(assignment_id):
    """Download scores as CSV"""
    try:
        assignment = db_session.query(Assignment).filter_by(
            id=assignment_id,
            guide_id=current_user.id
        ).first()
        
        if not assignment:
            return "Access denied", 403
        
        submissions = db_session.query(AssignmentSubmission).filter_by(
            assignment_id=assignment_id
        ).all()
        
        # Create CSV
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Header
        headers = ['Learner Name', 'Email', 'Submitted At', 'Total Score', 'Percentage', 'Flagged']
        
        # Add question headers
        if submissions and len(submissions) > 0:
            for idx, answer in enumerate(submissions[0].answers):
                headers.append(f'Q{idx + 1} Score')
        
        writer.writerow(headers)
        
        # Data rows
        for submission in submissions:
            percentage = (submission.total_score / assignment.total_marks * 100) if assignment.total_marks > 0 else 0
            
            row = [
                submission.learner.full_name or submission.learner.username,
                submission.learner.email,
                submission.submitted_at.strftime('%Y-%m-%d %H:%M'),
                round(submission.total_score, 2),
                round(percentage, 1),
                'Yes' if submission.is_flagged else 'No'
            ]
            
            # Add question scores
            for answer in submission.answers:
                row.append(round(answer.total_score, 2))
            
            writer.writerow(row)
        
        output.seek(0)
        
        return send_file(
            io.BytesIO(output.getvalue().encode()),
            as_attachment=True,
            download_name=f"{assignment.name}_scores.csv",
            mimetype='text/csv'
        )
        
    except Exception as e:
        print(f"Error generating CSV: {str(e)}")
        return "Error generating CSV", 500

# ========================================
# LEARNER ASSIGNMENT ROUTES
# ========================================

@app.route('/learner/assignments')
@login_required
@role_required('Learner')
def get_learner_assignments():
    """Get all assignments for learner's enrolled courses"""
    try:
        # Get courses learner is enrolled in
        course_ids = db_session.query(CourseLearner.course_id).filter_by(
            learner_id=current_user.id
        ).all()
        course_ids = [c[0] for c in course_ids]
        
        # Get assignments for these courses
        assignments = db_session.query(Assignment).filter(
            Assignment.course_id.in_(course_ids)
        ).order_by(Assignment.deadline).all()
        
        result = []
        for assignment in assignments:
            # Check if already submitted
            submission = db_session.query(AssignmentSubmission).filter_by(
                assignment_id=assignment.id,
                learner_id=current_user.id
            ).first()
            
            question_count = db_session.query(AssignmentQuestion).filter_by(
                assignment_id=assignment.id
            ).count()
            
            result.append({
                'id': assignment.id,
                'name': assignment.name,
                'course_name': assignment.course.course_name,
                'instructions': assignment.instructions,
                'deadline': assignment.deadline.isoformat(),
                'deadline_formatted': assignment.deadline.strftime('%B %d, %Y at %I:%M %p'),
                'total_marks': assignment.total_marks,
                'question_count': question_count,
                'is_submitted': submission is not None,
                'submission_id': submission.id if submission else None
            })
        
        return jsonify({'success': True, 'assignments': result})
        
    except Exception as e:
        print(f"Error getting learner assignments: {str(e)}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/learner/progress')
@login_required
@role_required('Learner')
def get_learner_progress():
    """Get learner's learning progress statistics"""
    try:
        # Get courses learner is enrolled in
        course_ids = db_session.query(CourseLearner.course_id).filter_by(
            learner_id=current_user.id
        ).all()
        course_ids = [c[0] for c in course_ids]
        
        if not course_ids:
            return jsonify({
                'success': True,
                'total_assignments': 0,
                'submitted_assignments': 0,
                'progress_percentage': 0,
                'message': 'No courses enrolled yet'
            })
        
        # Get total assignments for enrolled courses
        total_assignments = db_session.query(Assignment).filter(
            Assignment.course_id.in_(course_ids)
        ).count()
        
        # Get submitted assignments count
        submitted_count = db_session.query(AssignmentSubmission).join(Assignment).filter(
            AssignmentSubmission.learner_id == current_user.id,
            Assignment.course_id.in_(course_ids)
        ).count()
        
        # Calculate progress percentage
        progress_percentage = round((submitted_count / total_assignments * 100) if total_assignments > 0 else 0, 1)
        
        return jsonify({
            'success': True,
            'total_assignments': total_assignments,
            'submitted_assignments': submitted_count,
            'pending_assignments': total_assignments - submitted_count,
            'progress_percentage': progress_percentage,
            'message': f'{submitted_count} of {total_assignments} assignments submitted'
        })
        
    except Exception as e:
        print(f"Error getting learner progress: {str(e)}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/learner/assignments/<int:assignment_id>')
@login_required
@role_required('Learner')
def get_assignment_details(assignment_id):
    """Get assignment details with questions"""
    try:
        # Verify learner has access
        assignment = db_session.query(Assignment).get(assignment_id)
        
        if not assignment:
            return jsonify({'success': False, 'message': 'Assignment not found'})
        
        enrollment = db_session.query(CourseLearner).filter_by(
            course_id=assignment.course_id,
            learner_id=current_user.id
        ).first()
        
        if not enrollment:
            return jsonify({'success': False, 'message': 'Access denied'})
        
        # Get questions
        questions = db_session.query(AssignmentQuestion).filter_by(
            assignment_id=assignment_id
        ).order_by(AssignmentQuestion.order_num).all()
        
        questions_data = [{
            'id': q.id,
            'question_text': q.question_text,
            'marks': q.marks,
            'order': q.order_num
        } for q in questions]
        
        return jsonify({
            'success': True,
            'assignment': {
                'id': assignment.id,
                'name': assignment.name,
                'instructions': assignment.instructions,
                'deadline': assignment.deadline.strftime('%B %d, %Y at %I:%M %p'),
                'total_marks': assignment.total_marks,
                'questions': questions_data
            }
        })
        
    except Exception as e:
        print(f"Error getting assignment details: {str(e)}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/learner/assignments/submit', methods=['POST'])
@login_required
@role_required('Learner')
def submit_assignment():
    """Submit assignment with image uploads and AI evaluation"""
    try:
        # Get assignment_id from form data
        assignment_id = request.form.get('assignment_id')
        
        # Debug logging (can be removed in production)
        if app.debug:
            print(f"DEBUG: Received assignment_id: {assignment_id}")
            print(f"DEBUG: Form keys: {list(request.form.keys())}")
            print(f"DEBUG: Files received: {list(request.files.keys())}")
        
        if not assignment_id:
            return jsonify({
                'success': False, 
                'message': 'Assignment ID is required. Please refresh the page and try again.'
            }), 400
        
        # Validate assignment_id is a valid integer
        try:
            assignment_id_int = int(assignment_id)
        except (ValueError, TypeError):
            return jsonify({
                'success': False, 
                'message': 'Invalid assignment ID. Please refresh the page and try again.'
            }), 400
        
        assignment = db_session.query(Assignment).get(assignment_id_int)
        
        if not assignment:
            return jsonify({
                'success': False, 
                'message': 'Assignment not found. It may have been deleted or you may not have access.'
            }), 404
        
        # Verify access
        enrollment = db_session.query(CourseLearner).filter_by(
            course_id=assignment.course_id,
            learner_id=current_user.id
        ).first()
        
        if not enrollment:
            return jsonify({
                'success': False, 
                'message': 'Access denied. You are not enrolled in this course.'
            }), 403
        
        # Check if already submitted
        existing = db_session.query(AssignmentSubmission).filter_by(
            assignment_id=assignment.id,
            learner_id=current_user.id
        ).first()
        
        if existing:
            return jsonify({
                'success': False, 
                'message': 'This assignment has already been submitted. You cannot submit it again.'
            }), 409
        
        # Check deadline - properly handle timezone-aware and timezone-naive datetimes
        now = datetime.now(timezone.utc)
        deadline = assignment.deadline
        
        # Make deadline timezone-aware if it's naive (assume UTC)
        if deadline.tzinfo is None:
            deadline = deadline.replace(tzinfo=timezone.utc)
        
        # Check if deadline has passed - BLOCK submission if deadline passed
        if deadline < now:
            return jsonify({
                'success': False, 
                'message': f'Assignment deadline has passed. Deadline was: {deadline.strftime("%B %d, %Y at %I:%M %p")} UTC. You cannot submit this assignment.'
            }), 400
        
        # Create submission
        submission = AssignmentSubmission(
            assignment_id=assignment.id,
            learner_id=current_user.id
        )
        db_session.add(submission)
        db_session.flush()
        
        # Get questions
        questions = db_session.query(AssignmentQuestion).filter_by(
            assignment_id=assignment.id
        ).order_by(AssignmentQuestion.order_num).all()
        
        total_score = 0
        all_extracted_texts = []
        results_data = []
        first_answer_text = None
        
        # Process each question
        for question in questions:
            # Get uploaded files for this question
            files = request.files.getlist(f'question_{question.id}')
            
            # Validate files exist and have content
            valid_files = [f for f in files if f and f.filename and f.filename.strip()]
            
            if not valid_files or len(valid_files) == 0:
                db_session.rollback()
                return jsonify({
                    'success': False, 
                    'message': f'Please upload at least one image for question {question.order_num}: "{question.question_text[:50]}..."'
                }), 400
            
            # Save images and extract text
            image_paths = []
            extracted_texts = []
            
            for file in valid_files:
                try:
                    if not file.filename or not file.filename.strip():
                        continue
                    
                    # Validate file type
                    if not file.filename.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                        db_session.rollback()
                        return jsonify({
                            'success': False, 
                            'message': f'Invalid file type for question {question.order_num}. Please upload image files (JPG, PNG, GIF, etc.)'
                        }), 400
                    
                    filename = secure_filename(file.filename)
                    if not filename:
                        # Generate filename if secure_filename returns empty
                        filename = f"image_{int(datetime.now(timezone.utc).timestamp())}.jpg"
                    
                    timestamp = int(datetime.now(timezone.utc).timestamp())
                    filepath = os.path.join(
                        ASSIGNMENT_UPLOAD_FOLDER,
                        f"{current_user.id}_{assignment.id}_{question.id}_{timestamp}_{filename}"
                    )
                    
                    # Ensure directory exists
                    os.makedirs(os.path.dirname(filepath), exist_ok=True)
                    
                    file.save(filepath)
                    
                    # Verify file was saved
                    if not os.path.exists(filepath):
                        raise Exception(f"Failed to save file: {filename}")
                    
                    image_paths.append(filepath)
                    
                    # Extract text from image
                    extracted_text = extract_handwritten_text(filepath)
                    extracted_texts.append(extracted_text if extracted_text else "")
                    
                except Exception as file_error:
                    print(f"Error processing file for question {question.id}: {file_error}")
                    # Continue with other files, but log the error
                    continue
            
            # Validate that at least one file was successfully saved
            if not image_paths:
                db_session.rollback()
                return jsonify({
                    'success': False, 
                    'message': f'Failed to save images for question {question.order_num}. Please try uploading again.'
                }), 400
            
            # Combine all extracted texts for this question
            combined_text = ' '.join(extracted_texts).strip()
            
            if first_answer_text is None:
                first_answer_text = combined_text
            
            # Store for plagiarism check
            all_extracted_texts.append({
                'learner_id': current_user.id,
                'learner_name': current_user.full_name or current_user.username,
                'text': combined_text
            })
            
            # Evaluate answer using AI
            evaluation = evaluate_answer(
                question.question_text,
                combined_text,
                question.marks,
                assignment.evaluation_model
            )
            
            # Create answer record
            answer = AssignmentAnswer(
                submission_id=submission.id,
                question_id=question.id,
                image_paths=json.dumps(image_paths),
                extracted_text=combined_text,
                relevance_score=evaluation['relevance_score'],
                grammar_score=evaluation['grammar_score'],
                size_score=evaluation['size_score'],
                uniqueness_score=evaluation['uniqueness_score'],
                total_score=evaluation['total_score'],
                feedback=evaluation['feedback'],
                covered_points=json.dumps(evaluation['covered_points']),
                missing_points=json.dumps(evaluation['missing_points']),
                word_count=evaluation['word_count']
            )
            db_session.add(answer)
            db_session.flush()  # ensure answer.id is available
            
            # Store answer ID for MCQ
            results_data.append({
                'question_text': question.question_text,
                'max_marks': question.marks,
                'extracted_text': combined_text,
                'total_score': evaluation['total_score'],
                'feedback': evaluation['feedback'],
                'covered_points': evaluation['covered_points'],
                'missing_points': evaluation['missing_points'],
                'answer_id': answer.id  # Store answer_id for MCQ
            })

            # Real-time plagiarism detection (existing)
            try:
                plagiarism_results = detect_plagiarism_realtime(
                    answer.id,
                    question.id,
                    combined_text,
                    db_session
                )
            except Exception as e:
                print(f"Plagiarism detection error: {e}")
                plagiarism_results = []

            # Store plagiarism matches
            for match in plagiarism_results:
                plagiarism_match = PlagiarismMatch(
                    answer_id=answer.id,
                    matched_answer_id=match['matched_answer_id'],
                    similarity_score=match['similarity_score']
                )
                db_session.add(plagiarism_match)

                # Flag submission if high similarity
                if match['similarity_score'] >= 85:
                    if not submission.is_flagged:
                        submission.is_flagged = True
                        submission.flag_reason = f"High plagiarism detected ({match['similarity_score']}% match)"
                    else:
                        if "plagiarism" not in (submission.flag_reason or '').lower():
                            submission.flag_reason += f" | Plagiarism: {match['similarity_score']}%"
            
            # New: Plagiarism Accuracy Detection
            try:
                accuracy_result = analyze_submission_text(combined_text)
                if accuracy_result.get('success'):
                    plagiarism_data = accuracy_result.get('plagiarism', {})
                    ai_data = accuracy_result.get('ai_detection', {})
                    
                    accuracy_record = PlagiarismAccuracy(
                        answer_id=answer.id,
                        plagiarism_score=plagiarism_data.get('score', 0.0),
                        ai_confidence=ai_data.get('confidence', 0.0),
                        plagiarism_detected=plagiarism_data.get('plagiarism_detected', False),
                        ai_detected=ai_data.get('ai_detected', False),
                        details=json.dumps({
                            'plagiarism_matches': plagiarism_data.get('matches', []),
                            'ai_indicators': ai_data.get('indicators', {}),
                            'details': plagiarism_data.get('details', '')
                        })
                    )
                    db_session.add(accuracy_record)
                    
                    # Flag if high plagiarism or AI detected
                    if plagiarism_data.get('score', 0) >= 50 or ai_data.get('confidence', 0) >= 50:
                        if not submission.is_flagged:
                            submission.is_flagged = True
                            flag_msg = []
                            if plagiarism_data.get('score', 0) >= 50:
                                flag_msg.append(f"Plagiarism: {plagiarism_data.get('score', 0):.1f}%")
                            if ai_data.get('confidence', 0) >= 50:
                                flag_msg.append(f"AI Content: {ai_data.get('confidence', 0):.1f}%")
                            submission.flag_reason = " | ".join(flag_msg)
                        else:
                            flag_msg = []
                            if plagiarism_data.get('score', 0) >= 50:
                                flag_msg.append(f"Plagiarism: {plagiarism_data.get('score', 0):.1f}%")
                            if ai_data.get('confidence', 0) >= 50:
                                flag_msg.append(f"AI: {ai_data.get('confidence', 0):.1f}%")
                            if flag_msg:
                                existing = submission.flag_reason or ""
                                if "Plagiarism Accuracy" not in existing:
                                    submission.flag_reason = f"{existing} | {' | '.join(flag_msg)}"
            except Exception as e:
                print(f"Plagiarism accuracy detection error: {e}")
                import traceback
                traceback.print_exc()

            total_score += evaluation['total_score']
        
        # Update submission total score
        submission.total_score = total_score
        
        # Check for name mismatch flagging
        learner_name = current_user.full_name or current_user.username
        name_mismatch = check_name_mismatch(first_answer_text, learner_name)
        
        if name_mismatch:
            submission.is_flagged = True
            submission.flag_reason = "First three words do not match learner's name"
        
        # Check for plagiarism
        suspicious = check_plagiarism(
            all_extracted_texts[0]['text'],
            []  # Would need to load other submissions here
        )
        
        if suspicious and len(suspicious) > 0:
            submission.is_flagged = True
            if submission.flag_reason:
                submission.flag_reason += " | Suspicious activity: exact match with other submission(s)"
            else:
                submission.flag_reason = "Suspicious activity: exact match with other submission(s)"
        
        db_session.commit()
        
        # Collect answer IDs from results_data (we stored answer_id there)
        answer_ids = [r.get('answer_id') for r in results_data if r.get('answer_id')]
        
        # Fallback: Try to get from submission.answers if available
        if not answer_ids:
            try:
                db_session.refresh(submission)
                if submission.answers:
                    answer_ids = [answer.id for answer in submission.answers]
            except Exception as e:
                print(f"Error refreshing submission: {e}")
                # Last resort: Query directly
                try:
                    from sqlalchemy import text
                    result = db_session.execute(
                        text("SELECT id FROM assignment_answers WHERE submission_id = :submission_id ORDER BY id"),
                        {'submission_id': submission.id}
                    ).fetchall()
                    answer_ids = [row[0] for row in result]
                except Exception as query_error:
                    print(f"Error querying answer_ids: {query_error}")
        
        print(f"DEBUG submit_assignment - Submission ID: {submission.id}, Answer IDs: {answer_ids}")
        
        return jsonify({
            'success': True,
            'message': 'Assignment submitted successfully!',
            'answer_ids': answer_ids,
            'results': {
                'total_score': round(total_score, 2),
                'total_marks': assignment.total_marks,
                'scores_visible': assignment.scores_visible,
                'is_flagged': submission.is_flagged,
                'flag_reason': submission.flag_reason,
                'answers': results_data
            }
        })
        
    except ValueError as e:
        db_session.rollback()
        return jsonify({
            'success': False, 
            'message': f'Invalid input: {str(e)}. Please check your submission and try again.'
        }), 400
    except Exception as e:
        db_session.rollback()
        import traceback
        error_trace = traceback.format_exc()
        print(f"Error submitting assignment: {error_trace}")
        return jsonify({
            'success': False, 
            'message': f'An error occurred while submitting your assignment: {str(e)}. Please try again or contact support if the problem persists.'
        }), 500

@app.route('/learner/assignments/<int:assignment_id>/submission')
@login_required
@role_required('Learner')
def view_learner_submission(assignment_id):
    """View learner's own submission results"""
    try:
        submission = db_session.query(AssignmentSubmission).filter_by(
            assignment_id=assignment_id,
            learner_id=current_user.id
        ).first()
        
        if not submission:
            return jsonify({'success': False, 'message': 'Submission not found'})
        
        assignment = submission.assignment
        results_data = []
        
        for answer in submission.answers:
            covered = json.loads(answer.covered_points) if answer.covered_points else []
            missing = json.loads(answer.missing_points) if answer.missing_points else []
            
            results_data.append({
                'question_text': answer.question.question_text,
                'max_marks': answer.question.marks,
                'extracted_text': answer.extracted_text,
                'total_score': answer.total_score,
                'feedback': answer.feedback,
                'covered_points': covered,
                'missing_points': missing
            })
        
        return jsonify({
            'success': True,
            'results': {
                'total_score': round(submission.total_score, 2),
                'total_marks': assignment.total_marks,
                'scores_visible': assignment.scores_visible,
                'is_flagged': submission.is_flagged,
                'flag_reason': submission.flag_reason,
                'answers': results_data
            }
        })
        
    except Exception as e:
        print(f"Error viewing submission: {str(e)}")
        return jsonify({'success': False, 'message': str(e)})



@app.route('/guide/assignments/<int:assignment_id>/mcq-results')
@login_required
@role_required('Guide')
def get_mcq_results(assignment_id):
    """Get MCQ results for all learners"""
    db_session.rollback() # Clear any failed transactions
    try:
        
        assignment = db_session.query(Assignment).filter_by(
            id=assignment_id,
            guide_id=current_user.id
        ).first()
        
        if not assignment:
            return jsonify({'success': False, 'message': 'Assignment not found'})
        
        submissions = db_session.query(AssignmentSubmission).filter_by(
            assignment_id=assignment_id
        ).all()
        
        results = []
        for submission in submissions:
            learner_data = {
                'learner_id': submission.learner.id,
                'learner_name': submission.learner.full_name or submission.learner.username,
                'answers': []
            }
            
            print(f"DEBUG: Processing submission {submission.id} for learner {learner_data['learner_name']}")
            print(f"DEBUG: Submission has {len(submission.answers)} answers")
            
            for idx, answer in enumerate(submission.answers):
                # Pre-fetch question to avoid lazy loading
                try:
                    question = db_session.query(AssignmentQuestion).filter_by(id=answer.question_id).first()
                    if not question:
                        continue
                    question_text = question.question_text
                except Exception as q_error:
                    try:
                        db_session.rollback()
                        question = db_session.query(AssignmentQuestion).filter_by(id=answer.question_id).first()
                        if question:
                            question_text = question.question_text
                        else:
                            continue
                    except:
                        continue
                
                # Fetch MCQ response - use raw SQL to avoid mcq_type column issue
                mcq_response = None
                try:
                    from sqlalchemy import text
                    result = db_session.execute(
                        text("""
                            SELECT id, answer_id, mcq_question, option_a, option_b, option_c, option_d,
                                   correct_option, learner_answer, is_correct, time_taken, created_at
                            FROM learner_mcq_responses 
                            WHERE answer_id = :answer_id 
                            LIMIT 1
                        """),
                        {'answer_id': answer.id}
                    ).fetchone()
                    
                    if result:
                        # Create a simple object to hold the data
                        class MCQResponse:
                            def __init__(self, row):
                                self.id = row[0]
                                self.answer_id = row[1]
                                self.mcq_question = row[2] if row[2] else ''
                                self.option_a = row[3] if row[3] else ''
                                self.option_b = row[4] if row[4] else ''
                                self.option_c = row[5] if row[5] else ''
                                self.option_d = row[6] if row[6] else ''
                                self.correct_option = row[7] if row[7] else ''
                                self.learner_answer = row[8] if row[8] else ''
                                self.is_correct = row[9] if row[9] is not None else False
                                self.time_taken = row[10] if row[10] else 0
                                self.created_at = row[11] if row[11] else None
                        
                        mcq_response = MCQResponse(result)
                        print(f"DEBUG: Found MCQ response for answer_id {answer.id}: is_correct={mcq_response.is_correct}, learner_answer={mcq_response.learner_answer}")
                    else:
                        print(f"DEBUG: No MCQ response found for answer_id {answer.id}")
                except Exception as mcq_fetch_error:
                    db_session.rollback()
                    print(f"Error fetching MCQ: {mcq_fetch_error}")
                    mcq_response = None
                
                if mcq_response:
                    try:
                        mcq_question = getattr(mcq_response, 'mcq_question', '') or 'MCQ Question'
                        
                        learner_data['answers'].append({
                            'question_number': idx + 1,
                            'question_text': question_text,
                            'mcq_result': {
                                'type': 'answer',
                                'is_correct': mcq_response.is_correct,
                                'time_taken': mcq_response.time_taken or 0,
                                'learner_answer': mcq_response.learner_answer or '',
                                'correct_answer': mcq_response.correct_option or '',
                                'mcq_question': mcq_question,
                                'option_a': getattr(mcq_response, 'option_a', ''),
                                'option_b': getattr(mcq_response, 'option_b', ''),
                                'option_c': getattr(mcq_response, 'option_c', ''),
                                'option_d': getattr(mcq_response, 'option_d', '')
                            }
                        })
                        print(f"DEBUG: Added MCQ result for question {idx + 1}")
                    except Exception as attr_error:
                        print(f"Error accessing MCQ attributes: {attr_error}")
                        continue
            
            if learner_data['answers']:  # Only include if they have MCQ responses
                print(f"DEBUG: Adding learner {learner_data['learner_name']} with {len(learner_data['answers'])} MCQ responses")
                results.append(learner_data)
            else:
                print(f"DEBUG: Learner {learner_data['learner_name']} has no MCQ responses")
        
        print(f"DEBUG get_mcq_results: Returning {len(results)} learners with MCQ responses")
        
        return jsonify({
            'success': True,
            'results': results,
            'assignment': {
                'id': assignment.id,
                'name': assignment.name
            }
        })
        
    except Exception as e:
        # Always rollback on exception
        try:
            db_session.rollback()
        except:
            pass
        
        import traceback
        error_trace = traceback.format_exc()
        print(f"Error getting MCQ results: {error_trace}")
        return jsonify({'success': False, 'message': f'Error loading MCQ results: {str(e)}'})

@app.route('/guide/assignments/<int:assignment_id>/plagiarism-accuracy')
@login_required
@role_required('Guide')
def get_plagiarism_accuracy(assignment_id):
    """Get plagiarism accuracy results for all learners"""
    db_session.rollback()
    try:
        assignment = db_session.query(Assignment).filter_by(
            id=assignment_id,
            guide_id=current_user.id
        ).first()
        
        if not assignment:
            return jsonify({'success': False, 'message': 'Assignment not found'}), 404
        
        submissions = db_session.query(AssignmentSubmission).filter_by(
            assignment_id=assignment_id
        ).all()
        
        results = []
        for submission in submissions:
            learner_data = {
                'learner_id': submission.learner.id,
                'learner_name': submission.learner.full_name or submission.learner.username,
                'answers': []
            }
            
            for idx, answer in enumerate(submission.answers):
                try:
                    question = db_session.query(AssignmentQuestion).filter_by(id=answer.question_id).first()
                    if not question:
                        continue
                    question_text = question.question_text
                except:
                    continue
                
                # Fetch plagiarism accuracy results
                try:
                    accuracy = db_session.query(PlagiarismAccuracy).filter_by(
                        answer_id=answer.id
                    ).first()
                    
                    if accuracy:
                        try:
                            details = json.loads(accuracy.details) if accuracy.details else {}
                        except:
                            details = {}
                        
                        learner_data['answers'].append({
                            'question_number': idx + 1,
                            'question_text': question_text,
                            'plagiarism_score': accuracy.plagiarism_score,
                            'ai_confidence': accuracy.ai_confidence,
                            'plagiarism_detected': accuracy.plagiarism_detected,
                            'ai_detected': accuracy.ai_detected,
                            'matches': details.get('plagiarism_matches', []),
                            'ai_indicators': details.get('ai_indicators', {}),
                            'details': details.get('details', ''),
                            'extracted_text': answer.extracted_text[:200] + "..." if answer.extracted_text and len(answer.extracted_text) > 200 else (answer.extracted_text or "")
                        })
                except Exception as e:
                    print(f"Error processing answer {answer.id}: {e}")
                    continue
            
            if learner_data['answers']:
                results.append(learner_data)
        
        return jsonify({
            'success': True,
            'results': results,
            'assignment': {
                'id': assignment.id,
                'name': assignment.name
            }
        })
        
    except Exception as e:
        db_session.rollback()
        import traceback
        error_trace = traceback.format_exc()
        print(f"ERROR in get_plagiarism_accuracy: {error_trace}")
        return jsonify({
            'success': False,
            'message': f'Error fetching plagiarism accuracy: {str(e)}'
        }), 500

# ============= TEST MCQ GENERATION ROUTES =============

@app.route('/guide/test-mcq/upload', methods=['POST'])
@login_required
@role_required('Guide')
def upload_mcq_content():
    """Upload content for MCQ generation"""
    try:
        input_type = request.form.get('inputType', 'topic')
        
        if input_type == 'topic':
            # Check if topicText is in form or JSON
            topic_text = request.form.get('topicText', '').strip()
            if not topic_text:
                # Try to get from JSON if form data is empty
                try:
                    json_data = request.get_json(silent=True)
                    if json_data:
                        topic_text = json_data.get('topicText', '').strip()
                except:
                    pass
            
            if not topic_text:
                return jsonify({'success': False, 'error': 'Please enter a topic'}), 400
            
            return jsonify({
                'success': True,
                'type': 'topic',
                'content': topic_text,
                'requiresOptions': True
            })
        
        else:  # document
            if 'file' not in request.files:
                return jsonify({'success': False, 'error': 'No file uploaded'}), 400
            
            file = request.files['file']
            if file.filename == '':
                return jsonify({'success': False, 'error': 'No file selected'}), 400
            
            if not allowed_file(file.filename):
                return jsonify({'success': False, 'error': 'Unsupported file type. Allowed: PDF, CSV, TXT, DOCX'}), 400
            
            try:
                content = extract_text_from_file(file, file.filename)
                if not content or len(content.strip()) < 50:
                    return jsonify({'success': False, 'error': 'File content is too short or could not be extracted'}), 400
                
                return jsonify({
                    'success': True,
                    'type': 'document',
                    'content': content,
                    'filename': secure_filename(file.filename)
                })
            except Exception as e:
                return jsonify({'success': False, 'error': f'Error reading file: {str(e)}'}), 400
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/guide/test-mcq/generate', methods=['POST'])
@login_required
@role_required('Guide')
def generate_test_mcqs():
    """Generate MCQs using Gemini"""
    try:
        data = request.get_json()
        content = data.get('content')
        content_type = data.get('type')
        title = data.get('title', 'Untitled MCQ Test')
        
        if not content:
            return jsonify({'success': False, 'error': 'Content is required'}), 400
        
        # Get provider selection
        provider = data.get('provider', 'gemini')
        
        if provider not in ['gemini', 'groq']:
            provider = 'gemini'
        
        if content_type == 'topic':
            num_questions = int(data.get('numQuestions', 5))
            subcategory = data.get('subcategory', '')
            difficulty = data.get('difficulty', 'medium')
            option_difficulty = data.get('optionDifficulty', 'identifiable')
            
            if num_questions < 1 or num_questions > 50:
                return jsonify({'success': False, 'error': 'Number of questions must be between 1 and 50'}), 400
            
            try:
                mcqs = generate_mcqs_from_topic(content, num_questions, subcategory, difficulty, option_difficulty, provider)
            except Exception as e:
                return jsonify({'success': False, 'error': f'Error generating MCQs: {str(e)}'}), 500
        
        else:  # document
            max_questions = data.get('maxQuestions', False)
            num_questions = int(data.get('numQuestions', 10)) if not max_questions else None
            
            if not max_questions and (num_questions < 1 or num_questions > 100):
                return jsonify({'success': False, 'error': 'Number of questions must be between 1 and 100'}), 400
            
            try:
                mcqs = generate_mcqs_from_document(content, num_questions, max_questions, provider)
            except Exception as e:
                return jsonify({'success': False, 'error': f'Error generating MCQs: {str(e)}'}), 500
        
        if not mcqs or len(mcqs) == 0:
            return jsonify({'success': False, 'error': 'No MCQs were generated. Please try again.'}), 500
        
        # Get course ID (required)
        course_id = data.get('courseId')
        if not course_id:
            return jsonify({'success': False, 'error': 'Please select a course'}), 400
        
        # Verify course belongs to guide
        course = db_session.query(Course).filter_by(id=course_id, guide_id=current_user.id).first()
        if not course:
            return jsonify({'success': False, 'error': 'Invalid course selected'}), 400
        
        # Get marks per question (default 1.0)
        marks_per_question = float(data.get('marksPerQuestion', 1.0))
        total_marks = len(mcqs) * marks_per_question
        time_limit_minutes = int(data.get('timeLimitMinutes', 60))  # Default 60 minutes
        
        # Save to database
        test_mcq = TestMCQ(
            guide_id=current_user.id,
            course_id=course_id,
            title=title,
            source_type=content_type,
            source_content=content[:1000],  # Store preview
            source_filename=data.get('filename'),
            num_questions=len(mcqs),
            marks_per_question=marks_per_question,
            total_marks=total_marks,
            time_limit_minutes=time_limit_minutes,
            difficulty=data.get('difficulty'),
            subcategory=data.get('subcategory', ''),
            option_difficulty=data.get('optionDifficulty', ''),
            mcqs_data=json.dumps(mcqs),
            is_active=True
        )
        db_session.add(test_mcq)
        db_session.commit()
        
        return jsonify({
            'success': True,
            'mcqs': mcqs,
            'test_id': test_mcq.id,
            'title': title
        })
        
    except Exception as e:
        db_session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/guide/courses')
@login_required
@role_required('Guide')
def get_guide_courses():
    """Get all courses for the current guide"""
    try:
        courses = db_session.query(Course).filter_by(guide_id=current_user.id).all()
        return jsonify({
            'success': True,
            'courses': [{
                'id': c.id,
                'name': c.course_name,
                'course_name': c.course_name
            } for c in courses]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/guide/courses')
@login_required
@role_required('Guide')
def guide_course_list():
    """Get all courses managed by the guide"""
    try:
        courses = db_session.query(Course).filter_by(guide_id=current_user.id).all()
        return jsonify({
            'success': True,
            'courses': [{
                'id': c.id,
                'name': c.course_name
            } for c in courses]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/guide/test-mcq/list')
@login_required
@role_required('Guide')
def list_test_mcqs():
    """List all test MCQs created by guide"""
    try:
        mcqs = db_session.query(TestMCQ).filter_by(
            guide_id=current_user.id,
            is_active=True
        ).order_by(TestMCQ.created_at.desc()).all()
        
        return jsonify({
            'success': True,
            'mcqs': [{
                'id': m.id,
                'title': m.title,
                'source_type': m.source_type,
                'num_questions': m.num_questions,
                'total_marks': float(m.total_marks) if m.total_marks else m.num_questions,
                'marks_per_question': float(m.marks_per_question) if m.marks_per_question else 1.0,
                'created_at': m.created_at.isoformat() if m.created_at else None
            } for m in mcqs]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/guide/test-mcq/<int:test_mcq_id>/details')
@login_required
@role_required('Guide')
def get_test_mcq_details(test_mcq_id):
    """Return generated MCQ payload for preview"""
    try:
        test_mcq = db_session.query(TestMCQ).filter_by(
            id=test_mcq_id,
            guide_id=current_user.id,
            is_active=True
        ).first()
        if not test_mcq:
            return jsonify({'success': False, 'error': 'Test not found'}), 404

        mcqs = json.loads(test_mcq.mcqs_data) if test_mcq.mcqs_data else []
        return jsonify({
            'success': True,
            'test': {'id': test_mcq.id, 'title': test_mcq.title},
            'mcqs': mcqs
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/guide/test-mcq/<int:test_mcq_id>/update', methods=['POST'])
@login_required
@role_required('Guide')
def update_test_mcqs(test_mcq_id):
    """Allow the guide to edit/add generated MCQs before learners attempt the test."""
    try:
        payload = request.get_json(silent=True) or {}
        mcqs_input = payload.get('mcqs')
        if not isinstance(mcqs_input, list):
            return jsonify({'success': False, 'error': 'Invalid MCQ payload'}), 400

        test_mcq = db_session.query(TestMCQ).filter_by(
            id=test_mcq_id,
            guide_id=current_user.id,
            is_active=True
        ).first()

        if not test_mcq:
            return jsonify({'success': False, 'error': 'MCQ test not found'}), 404

        sanitized_mcqs = []
        for idx, mcq in enumerate(mcqs_input, start=1):
            question = (mcq.get('question', '') or '').strip()
            if not question:
                return jsonify({'success': False, 'error': f'Question {idx} is missing text'}), 400

            options = mcq.get('options') or []
            normalized_options = []
            for option in options:
                option_text = (option or '').strip()
                if option_text:
                    normalized_options.append(option_text)
            if len(normalized_options) < 2:
                return jsonify({'success': False, 'error': f'Question {idx} needs at least two options'}), 400

            answer = (mcq.get('answer', '') or '').strip().upper()
            valid_letters = [chr(65 + i) for i in range(len(normalized_options))]
            if answer not in valid_letters:
                answer = valid_letters[0]

            sanitized_mcqs.append({
                'question': question,
                'options': normalized_options,
                'answer': answer,
                'explanation': mcq.get('explanation', '') or ''
            })

        test_mcq.mcqs_data = json.dumps(sanitized_mcqs)
        db_session.commit()

        return jsonify({'success': True})
    except Exception as e:
        db_session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/guide/test-mcq/<int:test_mcq_id>/scores')
@login_required
@role_required('Guide')
def get_test_mcq_scores(test_mcq_id):
    """Return scores and summaries for a generated MCQ test"""
    try:
        test_mcq = db_session.query(TestMCQ).filter_by(
            id=test_mcq_id,
            guide_id=current_user.id,
            is_active=True
        ).first()

        if not test_mcq:
            return jsonify({'success': False, 'error': 'Test not found'}), 404

        attempts = db_session.query(TestMCQAttempt).filter_by(
            test_mcq_id=test_mcq_id
        ).order_by(TestMCQAttempt.completed_at.desc()).all()

        attempts_data = []
        total_score = 0
        total_marks_obtained = 0
        counted = 0

        for attempt in attempts:
            if not attempt.completed_at:
                continue

            counted += 1
            score_value = float(attempt.score or 0)
            marks_value = float(attempt.marks_obtained or 0)
            total_score += score_value
            total_marks_obtained += marks_value

            total_marks = float(attempt.total_marks) if attempt.total_marks is not None else (
                float(test_mcq.total_marks) if test_mcq.total_marks else float(test_mcq.num_questions or 0)
            )

            attempts_data.append({
                'id': attempt.id,
                'learner_name': attempt.learner.full_name or attempt.learner.username,
                'score': score_value,
                'marks_obtained': marks_value,
                'total_marks': total_marks,
                'correct_answers': attempt.correct_answers or 0,
                'total_questions': attempt.total_questions or 0,
                'completed_at': attempt.completed_at.isoformat() if attempt.completed_at else None
            })

        average_score = round(total_score / counted, 1) if counted else 0
        average_marks = round(total_marks_obtained / counted, 1) if counted else 0
        test_total_marks = float(test_mcq.total_marks) if test_mcq.total_marks else float(test_mcq.num_questions or 0)

        return jsonify({
            'success': True,
            'test': {
                'id': test_mcq.id,
                'title': test_mcq.title,
                'num_questions': test_mcq.num_questions,
                'total_marks': test_total_marks
            },
            'attempts': attempts_data,
            'average_score': average_score,
            'average_marks': average_marks
        })
    except Exception as e:
        db_session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/guide/test-mcq/<int:test_mcq_id>/scores/pdf')
@login_required
@role_required('Guide')
def download_test_mcq_scores_pdf(test_mcq_id):
    """Download MCQ test scores as PDF"""
    try:
        test_mcq = db_session.query(TestMCQ).filter_by(
            id=test_mcq_id,
            guide_id=current_user.id,
            is_active=True
        ).first()
        if not test_mcq:
            return "Access denied", 403

        attempts = db_session.query(TestMCQAttempt).filter(
            TestMCQAttempt.test_mcq_id == test_mcq_id,
            TestMCQAttempt.completed_at.isnot(None)
        ).order_by(TestMCQAttempt.completed_at.desc()).all()

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()

        title = Paragraph(f"<b>MCQ Scores Report</b><br/>{test_mcq.title}", styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 16))

        total_attempts = len(attempts)
        average_score = round(
            sum(float(a.score or 0) for a in attempts) / total_attempts, 1
        ) if total_attempts else 0
        stats_text = f"""
        <b>Summary</b><br/>
        Total Attempts: {total_attempts}<br/>
        Average Score: {average_score:.1f}%<br/>
        Total Questions: {test_mcq.num_questions or 0}<br/>
        Total Marks: {float(test_mcq.total_marks or 0):.1f}<br/>
        """
        elements.append(Paragraph(stats_text, styles['Normal']))
        elements.append(Spacer(1, 16))

        table_data = [['Learner', 'Score %', 'Marks', 'Correct', 'Total Q', 'Completed At']]
        for attempt in attempts:
            score = round(float(attempt.score or 0), 1)
            marks = f"{float(attempt.marks_obtained or 0):.1f}/{float(attempt.total_marks or 0):.1f}"
            completed_at = attempt.completed_at.strftime('%Y-%m-%d %H:%M:%S') if attempt.completed_at else 'N/A'
            table_data.append([
                attempt.learner.full_name or attempt.learner.username,
                f"{score:.1f}%",
                marks,
                attempt.correct_answers or 0,
                attempt.total_questions or 0,
                completed_at
            ])

        if len(table_data) > 1:
            table = Table(table_data, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4f46e5')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey])
            ]))
            elements.append(table)
        else:
            elements.append(Paragraph("No completed attempts yet.", styles['Normal']))

        doc.build(elements)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{test_mcq.title}_scores.pdf",
            mimetype='application/pdf'
        )
    except Exception as e:
        print(f"Error generating MCQ PDF: {e}")
        traceback.print_exc()
        return "Error generating PDF", 500

@app.route('/learner/test-mcq/list')
@login_required
@role_required('Learner')
def list_available_test_mcqs():
    """List all available test MCQs for learners"""
    try:
        # Get all active MCQs from courses the learner is enrolled in
        enrollments = db_session.query(CourseLearner).filter_by(learner_id=current_user.id).all()
        course_ids = [e.course_id for e in enrollments]
        
        if not course_ids:
            return jsonify({'success': True, 'mcqs': []})
        
        # Get MCQs from courses the learner is enrolled in
        mcqs = db_session.query(TestMCQ).filter(
            TestMCQ.course_id.in_(course_ids),
            TestMCQ.is_active == True
        ).order_by(TestMCQ.created_at.desc()).all()
        
        # Check which ones learner has attempted
        attempted_ids = set()
        attempts = db_session.query(TestMCQAttempt.test_mcq_id).filter_by(learner_id=current_user.id).all()
        attempted_ids = {a[0] for a in attempts}
        
        return jsonify({
            'success': True,
            'mcqs': [{
                'id': m.id,
                'title': m.title,
                'num_questions': m.num_questions,
                'total_marks': float(m.total_marks) if m.total_marks else m.num_questions,
                'time_limit_minutes': m.time_limit_minutes or 60,
                'created_at': m.created_at.isoformat() if m.created_at else None,
                'attempted': m.id in attempted_ids,
                'guide_name': m.guide.full_name or m.guide.username
            } for m in mcqs]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/learner/test-mcq/<int:test_mcq_id>')
@login_required
@role_required('Learner')
def get_test_mcq(test_mcq_id):
    """Get test MCQ for learner to attempt"""
    try:
        test_mcq = db_session.query(TestMCQ).filter_by(id=test_mcq_id, is_active=True).first()
        
        if not test_mcq:
            return jsonify({'success': False, 'error': 'Test MCQ not found'}), 404
        
        # Check if learner has access (guide teaches in learner's course)
        enrollments = db_session.query(CourseLearner).filter_by(learner_id=current_user.id).all()
        course_ids = [e.course_id for e in enrollments]
        
        if course_ids:
            guide_courses = db_session.query(Course).filter(
                Course.id.in_(course_ids),
                Course.guide_id == test_mcq.guide_id
            ).first()
            
            if not guide_courses:
                return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        mcqs = json.loads(test_mcq.mcqs_data) if test_mcq.mcqs_data else []
        
        # Remove correct answers and explanations for the attempt
        mcqs_for_attempt = []
        for idx, mcq in enumerate(mcqs):
            mcqs_for_attempt.append({
                'id': idx + 1,
                'question': mcq.get('question', ''),
                'options': mcq.get('options', [])
            })
        
        test_total_marks = float(test_mcq.total_marks) if test_mcq.total_marks else float(test_mcq.num_questions or len(mcqs_for_attempt))
        test_payload = {
            'id': test_mcq.id,
            'title': test_mcq.title,
            'num_questions': len(mcqs_for_attempt),
            'total_marks': round(test_total_marks, 1),
            'time_limit_minutes': test_mcq.time_limit_minutes or 60,
            'marks_per_question': float(test_mcq.marks_per_question or 1.0)
        }

        completed_attempt = db_session.query(TestMCQAttempt).filter(
            TestMCQAttempt.test_mcq_id == test_mcq.id,
            TestMCQAttempt.learner_id == current_user.id,
            TestMCQAttempt.completed_at.isnot(None)
        ).order_by(TestMCQAttempt.completed_at.desc()).first()

        already_attempted = bool(completed_attempt)
        previous_score = round(float(completed_attempt.score or 0), 1) if completed_attempt and completed_attempt.score is not None else None

        return jsonify({
            'success': True,
            'test_mcq': {
                'id': test_mcq.id,
                'title': test_mcq.title,
                'num_questions': len(mcqs_for_attempt)
            },
            'mcqs': mcqs_for_attempt,
            'test': test_payload,
            'already_attempted': already_attempted,
            'previous_score': previous_score
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/learner/test-mcq/<int:test_mcq_id>/start')
@login_required
@role_required('Learner')
def start_test_mcq(test_mcq_id):
    """Create or resume an attempt for a learner and return sanitized test data"""
    try:
        test_mcq = db_session.query(TestMCQ).filter_by(id=test_mcq_id, is_active=True).first()
        if not test_mcq:
            return jsonify({'success': False, 'error': 'Test MCQ not found'}), 404

        enrollments = db_session.query(CourseLearner).filter_by(learner_id=current_user.id).all()
        course_ids = [e.course_id for e in enrollments]
        if course_ids:
            guide_courses = db_session.query(Course).filter(
                Course.id.in_(course_ids),
                Course.guide_id == test_mcq.guide_id
            ).first()

            if not guide_courses:
                return jsonify({'success': False, 'error': 'Access denied'}), 403

        mcqs = json.loads(test_mcq.mcqs_data) if test_mcq.mcqs_data else []
        mcqs_payload = [{
            'question': mcq.get('question', ''),
            'options': mcq.get('options', [])
        } for mcq in mcqs]

        test_total_marks = float(test_mcq.total_marks) if test_mcq.total_marks else float(test_mcq.num_questions or len(mcqs_payload))
        test_payload = {
            'id': test_mcq.id,
            'title': test_mcq.title,
            'num_questions': len(mcqs_payload),
            'total_marks': round(test_total_marks, 1),
            'time_limit_minutes': test_mcq.time_limit_minutes or 60,
            'marks_per_question': float(test_mcq.marks_per_question or 1.0)
        }

        completed_attempt = db_session.query(TestMCQAttempt).filter(
            TestMCQAttempt.test_mcq_id == test_mcq_id,
            TestMCQAttempt.learner_id == current_user.id,
            TestMCQAttempt.completed_at.isnot(None)
        ).order_by(TestMCQAttempt.completed_at.desc()).first()

        if completed_attempt:
            return jsonify({
                'success': False,
                'error': 'You have already attempted this test',
                'previous_score': round(float(completed_attempt.score or 0), 1)
            }), 400

        attempt = db_session.query(TestMCQAttempt).filter(
            TestMCQAttempt.test_mcq_id == test_mcq_id,
            TestMCQAttempt.learner_id == current_user.id,
            TestMCQAttempt.completed_at.is_(None)
        ).order_by(TestMCQAttempt.started_at.desc()).first()

        if not attempt:
            attempt = TestMCQAttempt(
                test_mcq_id=test_mcq_id,
                learner_id=current_user.id,
                answers=json.dumps({}),
                score=0.0,
                marks_obtained=0.0,
                total_marks=test_total_marks,
                total_questions=len(mcqs_payload),
                correct_answers=0,
                started_at=datetime.now(timezone.utc)
            )
            db_session.add(attempt)
            db_session.commit()

        return jsonify({
            'success': True,
            'test': test_payload,
            'mcqs': mcqs_payload,
            'attempt_id': attempt.id
        })

    except Exception as e:
        db_session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/learner/test-mcq/<int:test_mcq_id>/submit', methods=['POST'])
@login_required
@role_required('Learner')
def submit_test_mcq_attempt(test_mcq_id):
    """Submit test MCQ attempt"""
    try:
        test_mcq = db_session.query(TestMCQ).filter_by(id=test_mcq_id, is_active=True).first()
        
        if not test_mcq:
            return jsonify({'success': False, 'error': 'Test MCQ not found'}), 404
        
        data = request.get_json()
        answers = data.get('answers', {})  # {question_id: selected_answer}
        attempt_id = data.get('attempt_id')  # Get attempt ID
        
        # Find the attempt record
        attempt = None
        if attempt_id:
            attempt = db_session.query(TestMCQAttempt).filter_by(
                id=attempt_id,
                test_mcq_id=test_mcq_id,
                learner_id=current_user.id
            ).first()
        
        if not attempt:
            return jsonify({'success': False, 'error': 'Attempt not found. Please start the test again.'}), 404
        
        mcqs = json.loads(test_mcq.mcqs_data) if test_mcq.mcqs_data else []
        
        if not mcqs:
            return jsonify({'success': False, 'error': 'MCQ data not found'}), 404
        
        # Calculate score
        total_questions = len(mcqs)
        correct_count = 0
        
        for idx, mcq in enumerate(mcqs):
            question_id = str(idx + 1)
            selected = answers.get(question_id, '').upper()
            correct = mcq.get('answer', '').upper()
            
            if selected == correct:
                correct_count += 1
        
        # Calculate marks and percentage
        marks_per_question = test_mcq.marks_per_question or 1.0
        total_marks = test_mcq.total_marks or (total_questions * marks_per_question)
        marks_obtained = correct_count * marks_per_question
        score = (marks_obtained / total_marks * 100) if total_marks > 0 else 0
        
        # Update existing attempt record
        attempt.answers = json.dumps(answers)
        attempt.score = score
        attempt.marks_obtained = marks_obtained
        attempt.total_marks = total_marks
        attempt.total_questions = total_questions
        attempt.correct_answers = correct_count
        attempt.completed_at = datetime.now(timezone.utc)
        
        db_session.commit()
        
        # Return results with correct answers
        results = []
        for idx, mcq in enumerate(mcqs):
            question_id = str(idx + 1)
            selected = answers.get(question_id, '').upper()
            correct = mcq.get('answer', '').upper()
            
            results.append({
                'question_id': question_id,
                'question': mcq.get('question', ''),
                'selected': selected,
                'correct': correct,
                'is_correct': selected == correct,
                'explanation': mcq.get('explanation', '')
            })
        
        return jsonify({
            'success': True,
            'score': round(score, 2),
            'marks_obtained': round(marks_obtained, 2),
            'total_marks': round(total_marks, 2),
            'correct_answers': correct_count,
            'total_questions': total_questions,
            'results': results
        })
        
    except Exception as e:
        db_session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/learner/test-mcq/attempts')
@login_required
@role_required('Learner')
def get_my_test_mcq_attempts():
    """Get learner's test MCQ attempts"""
    try:
        attempts = db_session.query(TestMCQAttempt).filter_by(
            learner_id=current_user.id
        ).order_by(TestMCQAttempt.completed_at.desc()).all()
        
        return jsonify({
            'success': True,
            'attempts': [{
                'id': a.id,
                'test_mcq_id': a.test_mcq_id,
                'test_title': a.test_mcq.title,
                'score': a.score,
                'marks_obtained': float(a.marks_obtained) if a.marks_obtained else 0,
                'total_marks': float(a.total_marks) if a.total_marks else a.total_questions,
                'correct_answers': a.correct_answers,
                'total_questions': a.total_questions,
                'completed_at': a.completed_at.isoformat() if a.completed_at else None
            } for a in attempts]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    

@app.route('/guide/assignments/<int:assignment_id>/mcq-report/pdf')
@login_required
@role_required('Guide')
def download_mcq_report_pdf(assignment_id):
    """Download MCQ results as PDF"""
    try:
        assignment = db_session.query(Assignment).filter_by(
            id=assignment_id,
            guide_id=current_user.id
        ).first()
        
        if not assignment:
            return "Access denied", 403
        
        submissions = db_session.query(AssignmentSubmission).filter_by(
            assignment_id=assignment_id
        ).all()
        
        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title = Paragraph(f"<b>MCQ Verification Report - {assignment.name}</b>", styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 20))
        
        # Statistics
        total_mcqs = 0
        correct_mcqs = 0
        
        for submission in submissions:
            for answer in submission.answers:
                # Fetch MCQ - handle missing mcq_type column
                mcq = None
                try:
                    mcq = db_session.query(LearnerMCQResponse).filter_by(answer_id=answer.id).first()
                except Exception as mcq_stat_error:
                    # If query fails due to missing column, use raw SQL
                    if 'mcq_type' in str(mcq_stat_error).lower() or 'undefinedcolumn' in str(mcq_stat_error).lower():
                        try:
                            from sqlalchemy import text
                            result = db_session.execute(
                                text("SELECT is_correct FROM learner_mcq_responses WHERE answer_id = :answer_id LIMIT 1"),
                                {'answer_id': answer.id}
                            ).fetchone()
                            if result:
                                class SimpleMCQ:
                                    def __init__(self, is_correct):
                                        self.is_correct = is_correct
                                mcq = SimpleMCQ(result[0])
                        except Exception:
                            mcq = None
                    else:
                        print(f"Error fetching MCQ for stats: {mcq_stat_error}")
                
                if mcq:
                    total_mcqs += 1
                    if mcq.is_correct:
                        correct_mcqs += 1
        
        accuracy = (correct_mcqs / total_mcqs * 100) if total_mcqs > 0 else 0
        
        stats_text = f"""
        <b>Overall Statistics</b><br/>
        Total Learners: {len(submissions)}<br/>
        Total MCQs: {total_mcqs}<br/>
        Correct Answers: {correct_mcqs}<br/>
        Accuracy: {accuracy:.1f}%<br/>
        """
        elements.append(Paragraph(stats_text, styles['Normal']))
        elements.append(Spacer(1, 30))
        
        # Detailed results
        for submission in submissions:
            elements.append(Paragraph(f"<b>{submission.learner.full_name or submission.learner.username}</b>", styles['Heading3']))
            
            mcq_data = [['Question', 'Type', 'Result', 'Time', 'Answer', 'Correct']]
            
            for idx, answer in enumerate(submission.answers):
                # Fetch MCQ - handle missing mcq_type column
                mcq = None
                try:
                    mcq = db_session.query(LearnerMCQResponse).filter_by(answer_id=answer.id).first()
                except Exception as mcq_pdf_error:
                    # If query fails due to missing column, use raw SQL
                    if 'mcq_type' in str(mcq_pdf_error).lower() or 'undefinedcolumn' in str(mcq_pdf_error).lower():
                        try:
                            from sqlalchemy import text
                            result = db_session.execute(
                                text("SELECT id, answer_id, mcq_question, option_a, option_b, option_c, option_d, correct_option, learner_answer, is_correct, time_taken, created_at FROM learner_mcq_responses WHERE answer_id = :answer_id LIMIT 1"),
                                {'answer_id': answer.id}
                            ).fetchone()
                            if result:
                                class SimpleMCQ:
                                    def __init__(self, row):
                                        self.id = row[0]
                                        self.answer_id = row[1]
                                        self.mcq_question = row[2]
                                        self.option_a = row[3]
                                        self.option_b = row[4]
                                        self.option_c = row[5]
                                        self.option_d = row[6]
                                        self.correct_option = row[7]
                                        self.learner_answer = row[8]
                                        self.is_correct = row[9]
                                        self.time_taken = row[10]
                                        self.created_at = row[11]
                                mcq = SimpleMCQ(result)
                        except Exception:
                            mcq = None
                    else:
                        print(f"Error fetching MCQ for PDF: {mcq_pdf_error}")
                
                if mcq:
                    # Handle missing mcq_type column gracefully
                    mcq_type = getattr(mcq, 'mcq_type', 'answer')
                    mcq_data.append([
                        f"Q{idx + 1}",
                        'Topic' if mcq_type == 'question' else 'Answer',
                        '✓' if mcq.is_correct else '✗',
                        f"{mcq.time_taken}s",
                        mcq.learner_answer or '-',
                        mcq.correct_option
                    ])
            
            if len(mcq_data) > 1:
                table = Table(mcq_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                elements.append(table)
                elements.append(Spacer(1, 20))
        
        doc.build(elements)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{assignment.name}_MCQ_Report.pdf",
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"Error generating MCQ PDF: {e}")
        return "Error generating report", 500


@app.route('/guide/assignments/<int:assignment_id>/plagiarism')
@login_required
@role_required('Guide')
def get_plagiarism_results(assignment_id):
    """Get plagiarism detection results"""
    try:
        assignment = db_session.query(Assignment).filter_by(
            id=assignment_id,
            guide_id=current_user.id
        ).first()
        
        if not assignment:
            return jsonify({'success': False, 'message': 'Assignment not found'})
        
        submissions = db_session.query(AssignmentSubmission).filter_by(
            assignment_id=assignment_id
        ).all()
        
        results = []
        
        for submission in submissions:
            learner_matches = []
            
            for idx, answer in enumerate(submission.answers):
                matches = db_session.query(PlagiarismMatch).filter_by(
                    answer_id=answer.id
                ).all()
                
                for match in matches:
                    matched_answer = db_session.query(AssignmentAnswer).get(match.matched_answer_id)
                    
                    learner_matches.append({
                        'match_id': match.id,
                        'answer_id': answer.id,
                        'question_number': idx + 1,
                        'matched_learner_name': matched_answer.submission.learner.full_name or matched_answer.submission.learner.username,
                        'similarity_score': match.similarity_score,
                        'status': match.status,
                        'penalty_marks': match.penalty_marks,
                        'max_marks': answer.question.marks,
                        'answer_text': answer.extracted_text[:500] + '...' if len(answer.extracted_text) > 500 else answer.extracted_text,
                        'matched_text': matched_answer.extracted_text[:500] + '...' if len(matched_answer.extracted_text) > 500 else matched_answer.extracted_text
                    })
            
            if learner_matches:
                results.append({
                    'learner_id': submission.learner.id,
                    'learner_name': submission.learner.full_name or submission.learner.username,
                    'plagiarism_matches': learner_matches
                })
        
        return jsonify({
            'success': True,
            'results': results,
            'assignment': {
                'id': assignment.id,
                'name': assignment.name
            }
        })
        
    except Exception as e:
        print(f"Error getting plagiarism results: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})



@app.route('/guide/plagiarism/apply-penalty', methods=['POST'])
@login_required
@role_required('Guide')
def apply_plagiarism_penalty():
    """Apply penalty for plagiarism"""
    try:
        data = request.get_json()
        match_id = data.get('match_id')
        answer_id = data.get('answer_id')
        penalty_marks = float(data.get('penalty_marks', 0))
        
        # Update plagiarism match
        match = db_session.query(PlagiarismMatch).get(match_id)
        if not match:
            return jsonify({'success': False, 'message': 'Match not found'})
        
        match.status = 'penalized'
        match.penalty_marks = penalty_marks
        match.reviewed_by = current_user.id
        match.reviewed_at = datetime.now(timezone.utc)
        
        # Deduct marks from answer
        answer = db_session.query(AssignmentAnswer).get(answer_id)
        if answer:
            answer.total_score = max(0, answer.total_score - penalty_marks)
            
            # Update submission total
            submission = answer.submission
            submission.total_score = sum(a.total_score for a in submission.answers)
        
        db_session.commit()
        
        return jsonify({'success': True, 'message': 'Penalty applied successfully'})
        
    except Exception as e:
        db_session.rollback()
        print(f"Error applying penalty: {e}")
        return jsonify({'success': False, 'message': str(e)})


@app.route('/guide/plagiarism/ignore', methods=['POST'])
@login_required
@role_required('Guide')
def ignore_plagiarism():
    """Mark plagiarism as false positive"""
    try:
        data = request.get_json()
        match_id = data.get('match_id')
        
        match = db_session.query(PlagiarismMatch).get(match_id)
        if not match:
            return jsonify({'success': False, 'message': 'Match not found'})
        
        match.status = 'ignored'
        match.reviewed_by = current_user.id
        match.reviewed_at = datetime.now(timezone.utc)
        
        db_session.commit()
        
        return jsonify({'success': True, 'message': 'Marked as false positive'})
        
    except Exception as e:
        db_session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@app.route('/guide/plagiarism/undo-penalty', methods=['POST'])
@login_required
@role_required('Guide')
def undo_plagiarism_penalty():
    """Undo plagiarism penalty"""
    try:
        data = request.get_json()
        match_id = data.get('match_id')
        answer_id = data.get('answer_id')
        
        match = db_session.query(PlagiarismMatch).get(match_id)
        if not match:
            return jsonify({'success': False, 'message': 'Match not found'})
        
        # Restore marks
        answer = db_session.query(AssignmentAnswer).get(answer_id)
        if answer:
            answer.total_score += match.penalty_marks
            
            # Update submission total
            submission = answer.submission
            submission.total_score = sum(a.total_score for a in submission.answers)
        
        # Reset match
        match.status = 'pending'
        match.penalty_marks = 0
        match.reviewed_by = None
        match.reviewed_at = None
        
        db_session.commit()
        
        return jsonify({'success': True, 'message': 'Penalty removed'})
        
    except Exception as e:
        db_session.rollback()
        return jsonify({'success': False, 'message': str(e)})


# ==================== GUIDE SUMMARIZER ROUTES ====================

@app.route('/guide/summarizer/upload', methods=['POST'])
@login_required
@role_required('Guide')
def guide_summarizer_upload():
    """Handle file upload for Guide Summarizer"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        original_filename = secure_filename(file.filename)
        file_ext = original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else ''
        
        # Check for unsupported formats
        allowed_extensions = {'pdf', 'docx', 'csv', 'txt', 'jpg', 'jpeg', 'png'}
        if file_ext not in allowed_extensions:
            return jsonify({
                'error': f'File type .{file_ext} not supported. Please convert to PDF, DOCX, TXT, CSV, JPG, or PNG'
            }), 400
        
        # Save file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{original_filename}"
        file_path = os.path.join(DOCUMENT_QA_UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        # Extract text based on file type
        text = ""
        page_count = 0
        
        if file_ext == 'pdf':
            text, page_count = extract_text_from_pdf_qa(file_path)
        elif file_ext == 'docx':
            text, page_count = extract_text_from_docx_qa(file_path)
        elif file_ext == 'txt':
            text, page_count = extract_text_from_txt_qa(file_path)
        elif file_ext == 'csv':
            text, page_count = extract_text_from_csv_qa(file_path)
        elif file_ext in ['jpg', 'jpeg', 'png']:
            text, page_count = extract_text_from_image_qa(file_path)
        
        # Generate summary and topics
        summary, topics = generate_document_summary(text)
        
        # Store in database
        file_size = os.path.getsize(file_path)
        course_id = request.form.get('course_id')  # Required course association
        if not course_id or not course_id.isdigit():
            # Delete the uploaded file if course_id is missing
            try:
                os.remove(file_path)
            except:
                pass
            return jsonify({'error': 'Course selection is required. Please select a course before uploading.'}), 400
        
        course_id = int(course_id)
        
        # Verify the guide has access to this course
        course = db_session.query(Course).filter_by(id=course_id, guide_id=current_user.id).first()
        if not course:
            # Delete the uploaded file if course is invalid
            try:
                os.remove(file_path)
            except:
                pass
            return jsonify({'error': 'Invalid course or you do not have access to this course'}), 403
        
        document = DocumentQA(
            user_id=current_user.id,
            course_id=course_id,
            filename=filename,
            original_filename=original_filename,
            file_type=file_ext,
            file_size=file_size,
            page_count=page_count,
            extracted_text=text,
            file_path=file_path,
            summary=summary,
            topics=json.dumps(topics) if topics else "[]"
        )
        db_session.add(document)
        db_session.commit()
        
        # Store embeddings asynchronously (in production, use background task)
        try:
            store_embeddings_qa(document.id, text)
        except Exception as e:
            print(f"Error storing embeddings: {e}")
        
        return jsonify({
            'success': True,
            'document_id': document.id,
            'filename': original_filename,
            'summary': summary,
            'topics': topics,
            'page_count': page_count,
            'word_count': len(text.split())
        })
    except Exception as e:
        db_session.rollback()
        print(f"Upload error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/guide/summarizer/documents', methods=['GET'])
@login_required
@role_required('Guide')
def guide_summarizer_documents():
    """Get list of uploaded documents for Guide"""
    try:
        documents = db_session.query(DocumentQA).filter_by(user_id=current_user.id).order_by(DocumentQA.upload_time.desc()).all()
        
        result = []
        for doc in documents:
            topics = []
            try:
                topics = json.loads(doc.topics) if doc.topics else []
            except:
                pass
            
            # Get course name if course_id exists
            course_name = None
            if doc.course_id:
                course = db_session.query(Course).get(doc.course_id)
                if course:
                    course_name = course.course_name
            
            result.append({
                'id': doc.id,
                'filename': doc.original_filename,
                'file_type': doc.file_type,
                'file_size': doc.file_size,
                'page_count': doc.page_count,
                'summary': doc.summary,
                'topics': topics,
                'course_id': doc.course_id,
                'course_name': course_name,
                'upload_time': doc.upload_time.isoformat() if doc.upload_time else None
            })
        
        return jsonify(result)
    except Exception as e:
        print(f"Error getting documents: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/guide/summarizer/delete/<int:document_id>', methods=['DELETE'])
@login_required
@role_required('Guide')
def guide_summarizer_delete(document_id):
    """Delete a document"""
    try:
        document = db_session.query(DocumentQA).filter_by(id=document_id, user_id=current_user.id).first()
        if not document:
            return jsonify({'error': 'Document not found'}), 404
        
        # Delete file
        if document.file_path and os.path.exists(document.file_path):
            try:
                os.remove(document.file_path)
            except:
                pass
        
        # Delete from database (cascades to chunks and chat)
        db_session.delete(document)
        db_session.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        db_session.rollback()
        return jsonify({'error': str(e)}), 500

# ==================== LEARNER DOCUMENT CHAT ROUTES ====================

@app.route('/learner/document-chat/courses', methods=['GET'])
@login_required
@role_required('Learner')
def learner_document_chat_courses():
    """Get courses available for learner"""
    try:
        # Get courses where learner is enrolled
        course_learners = db_session.query(CourseLearner).filter_by(learner_id=current_user.id).all()
        course_ids = [cl.course_id for cl in course_learners]
        
        courses = db_session.query(Course).filter(Course.id.in_(course_ids)).all()
        
        result = []
        for course in courses:
            result.append({
                'id': course.id,
                'name': course.course_name,
                'organization': course.organization
            })
        
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/learner/document-chat/documents/<int:course_id>', methods=['GET'])
@login_required
@role_required('Learner')
def learner_document_chat_documents(course_id):
    """Get documents available for a course"""
    try:
        # Verify learner is enrolled in course
        enrollment = db_session.query(CourseLearner).filter_by(
            course_id=course_id, learner_id=current_user.id
        ).first()
        
        if not enrollment:
            return jsonify({'error': 'Not enrolled in this course'}), 403
        
        # Get documents uploaded by guide for this course
        documents = db_session.query(DocumentQA).filter_by(course_id=course_id).order_by(DocumentQA.upload_time.desc()).all()
        
        result = []
        for doc in documents:
            topics = []
            try:
                topics = json.loads(doc.topics) if doc.topics else []
            except:
                pass
            
            result.append({
                'id': doc.id,
                'filename': doc.original_filename,
                'file_type': doc.file_type,
                'summary': doc.summary,
                'topics': topics,
                'upload_time': doc.upload_time.isoformat() if doc.upload_time else None
            })
        
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/learner/document-chat/chat', methods=['POST'])
@login_required
@role_required('Learner')
def learner_document_chat():
    """Handle chat queries for learner document Q&A"""
    try:
        # Debug logging
        print(f"DEBUG learner_document_chat - Content-Type: {request.content_type}")
        print(f"DEBUG learner_document_chat - Headers: {dict(request.headers)}")
        
        # Get JSON data
        if not request.is_json:
            return jsonify({'error': 'Request must be JSON'}), 400
        
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'error': 'Invalid JSON data'}), 400
        
        print(f"DEBUG learner_document_chat - Data keys: {list(data.keys()) if data else 'None'}")
        
        document_id = data.get('document_id')
        user_query = data.get('query')
        session_id = session.get('session_id', os.urandom(16).hex())
        session['session_id'] = session_id
        
        if not document_id or not user_query:
            return jsonify({'error': 'Document ID and query are required'}), 400
        
        if not GEMINI_API_KEY:
            return jsonify({'error': 'GEMINI_API_KEY not configured. Please configure it in your .env file.'}), 500
        
        # Verify document exists and learner has access
        document = db_session.query(DocumentQA).get(document_id)
        if not document:
            return jsonify({'error': 'Document not found'}), 404
        
        # Verify learner is enrolled in course
        if document.course_id:
            enrollment = db_session.query(CourseLearner).filter_by(
                course_id=document.course_id, learner_id=current_user.id
            ).first()
            if not enrollment:
                return jsonify({'error': 'Access denied'}), 403
        
        # Get relevant chunks using embeddings
        query_embedding = get_embedding_qa(user_query)
        if query_embedding is None:
            return jsonify({'error': 'Failed to generate query embedding'}), 500
        
        query_embedding = np.array(query_embedding, dtype=np.float32)
        query_norm = np.linalg.norm(query_embedding)
        if query_norm == 0:
            return jsonify({'error': 'Query embedding has zero norm'}), 500
        query_embedding = query_embedding / query_norm
        
        # Get all chunks for the document
        chunks = db_session.query(DocumentQAChunk).filter_by(document_id=document_id).all()
        
        if not chunks:
            return jsonify({'error': 'No embeddings found for this document'}), 400
        
        # Calculate similarities
        similarities = []
        for chunk in chunks:
            try:
                chunk_embedding = np.array(json.loads(chunk.embedding), dtype=np.float32)
                
                if len(chunk_embedding) == 0 or np.isnan(chunk_embedding).any() or np.isinf(chunk_embedding).any():
                    continue
                if len(chunk_embedding) != len(query_embedding):
                    continue
                
                chunk_norm = np.linalg.norm(chunk_embedding)
                if chunk_norm == 0:
                    continue
                chunk_embedding = chunk_embedding / chunk_norm
                
                similarity = np.dot(query_embedding, chunk_embedding)
                
                if not (np.isnan(similarity) or np.isinf(similarity)):
                    similarities.append({
                        'text': chunk.chunk_text,
                        'similarity': float(similarity)
                    })
            except Exception as e:
                print(f"Error processing chunk: {e}")
                continue
        
        # Get top 5 most relevant chunks
        similarities.sort(key=lambda x: x['similarity'], reverse=True)
        top_chunks = [s['text'] for s in similarities[:5]]
        context = "\n\n".join(top_chunks)
        
        # Get conversation history
        history = db_session.query(DocumentQAChat).filter_by(
            session_id=session_id, document_id=document_id
        ).order_by(DocumentQAChat.timestamp.desc()).limit(10).all()
        history.reverse()
        
        # Build conversation context
        conversation_context = ""
        current_topic = None
        last_main_question = None
        
        if history:
            recent_exchanges = history[-3:] if len(history) >= 3 else history
            for h in reversed(recent_exchanges):
                user_msg = h.user_message.lower()
                if any(word in user_msg for word in ['explain', 'what is', 'tell me about', 'describe', 'about', 'how does']):
                    last_main_question = h.user_message
                    words = h.user_message.split()
                    topic_keywords = []
                    for i, word in enumerate(words):
                        word_lower = word.lower().strip('.,!?')
                        if word_lower in ['explain', 'what', 'tell', 'describe', 'about', 'how']:
                            skip_words = ['is', 'are', 'the', 'a', 'an', 'this', 'that', 'file', 'storage', 'as', 'it']
                            j = i + 1
                            while j < len(words) and len(topic_keywords) < 6:
                                next_word = words[j].lower().strip('.,!?')
                                if next_word not in skip_words or len(topic_keywords) > 0:
                                    topic_keywords.append(words[j])
                                j += 1
                            if topic_keywords:
                                break
                    if topic_keywords:
                        current_topic = ' '.join(topic_keywords)
                        break
        
        for h in history:
            response_preview = h.bot_response[:300] + '...' if len(h.bot_response) > 300 else h.bot_response
            conversation_context += f"User: {h.user_message}\nAssistant: {response_preview}\n\n"
        
        query_lower = user_query.lower()
        is_followup = (
            len(user_query.split()) < 10 and 
            not any(word in query_lower for word in ['explain', 'what is', 'what are', 'tell me', 'describe', 'about', 'how does']) and
            any(word in query_lower for word in ['bullet', 'points', 'shorter', 'simpler', 'expand', 'longer', 'summary', 'notes', 'exam'])
        )
        
        if current_topic and is_followup and last_main_question:
            topic_context = f"""
CRITICAL CONTEXT AWARENESS:
- The user previously asked: "{last_main_question}"
- They are currently discussing the topic: "{current_topic}"
- The current question "{user_query}" is a FOLLOW-UP about the SAME TOPIC
- You MUST answer about "{current_topic}" even though the user didn't mention it again
"""
        else:
            topic_context = ""
        
        prompt = f"""You are an AI assistant helping users understand documents. Answer the question based ONLY on the provided context from the document.

Previous conversation:
{conversation_context}
{topic_context}

Document context:
{context}

User question: {user_query}

CRITICAL INSTRUCTIONS:
1. Answer ONLY based on the provided document context
2. If the answer isn't in the context, say so clearly
3. CONTEXT AWARENESS: If the user asks a follow-up question, maintain context from previous messages
4. Format your response using simple text formatting:
   - Use line breaks between paragraphs
   - For bullet points, use lines starting with "* "
   - DO NOT use markdown syntax like **bold** or __italic__
5. Response formatting rules:
   - If asked for "bullet points" or "points", format as bullet points (each point on a new line starting with "* ")
   - If asked to "Make it simpler", use simple language
   - If asked to "Expand", provide more detail

Answer in clear, well-formatted plain text with bullet points using "* " prefix (NO markdown syntax):"""
        
        response = chat_model.generate_content(prompt)
        bot_response = response.text
        bot_response_html = markdown_to_html_qa(bot_response)
        
        # Store in chat history
        chat_entry = DocumentQAChat(
            session_id=session_id,
            document_id=document_id,
            user_message=user_query,
            bot_response=bot_response
        )
        db_session.add(chat_entry)
        db_session.commit()
        
        return jsonify({
            'response': bot_response_html,
            'chunks_used': len(top_chunks)
        })
    except Exception as e:
        db_session.rollback()
        print(f"Chat error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Chat error: {str(e)}'}), 500

@app.route('/learner/document-chat/text-to-speech', methods=['POST'])
@login_required
@role_required('Learner')
def learner_document_chat_tts():
    """Convert text to speech for learner"""
    try:
        data = request.json
        text = data.get('text')
        language = data.get('language', 'en')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Generate speech
        tts = gTTS(text=text, lang=language, slow=False)
        fp = io.BytesIO()
        tts.write_to_fp(fp)
        fp.seek(0)
        audio_base64 = base64.b64encode(fp.read()).decode('utf-8')
        
        return jsonify({
            'audio': audio_base64,
            'format': 'mp3'
        })
    except Exception as e:
        return jsonify({'error': f'TTS error: {str(e)}'}), 500

@app.route('/transcribe', methods=['POST'])
@login_required
def transcribe_audio():
    """Transcribe audio to text using OpenAI Whisper with multi-language support"""
    if 'audio' not in request.files:
        return jsonify({'error': 'No audio file provided'}), 400
    
    audio_file = request.files['audio']
    
    if not whisper_model:
        return jsonify({'error': 'Whisper model not available. Please ensure openai-whisper is installed.'}), 500
    
    # Get optional language parameter (ISO 639-1 code)
    # If not provided, Whisper will auto-detect the language
    language = request.form.get('language', None)
    if language and language not in WHISPER_SUPPORTED_LANGUAGES:
        language = None  # Invalid language code, use auto-detect
    
    # Save temporarily
    temp_path = os.path.join(DOCUMENT_QA_UPLOAD_FOLDER, f'temp_audio_{datetime.now().timestamp()}.wav')
    try:
        audio_file.save(temp_path)
        
        # Transcribe with OpenAI Whisper
        # language=None means auto-detect, task="transcribe" is default
        transcribe_options = {
            'language': language if language else None,  # Auto-detect if not specified
            'task': 'transcribe',  # Options: 'transcribe' or 'translate' (translate to English)
            'fp16': False,  # Use FP32 for CPU compatibility
            'verbose': False  # Set to True for debugging
        }
        
        result = whisper_model.transcribe(temp_path, **transcribe_options)
        
        # Extract transcription and detected language
        text = result['text'].strip()
        detected_language = result.get('language', 'unknown')
        detected_language_name = WHISPER_SUPPORTED_LANGUAGES.get(detected_language, detected_language)
        
        # Clean up
        if os.path.exists(temp_path):
            os.remove(temp_path)
        
        return jsonify({
            'text': text,
            'detected_language': detected_language,
            'detected_language_name': detected_language_name,
            'confidence': result.get('no_speech_prob', 0)  # Lower is better (0 = high confidence)
        })
    except Exception as e:
        # Clean up on error
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except:
                pass
        print(f"Transcription error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Transcription error: {str(e)}'}), 500

@app.route('/transcribe/languages', methods=['GET'])
@login_required
def get_supported_languages():
    """Get list of supported languages for transcription"""
    return jsonify({
        'languages': WHISPER_SUPPORTED_LANGUAGES,
        'total': len(WHISPER_SUPPORTED_LANGUAGES),
        'auto_detect': True  # Whisper supports automatic language detection
    })

# ========================================
# DIAGRAM EVALUATION ROUTES
# ========================================

# Create upload folder for diagrams
DIAGRAM_UPLOAD_FOLDER = 'static/diagram_uploads'
os.makedirs(DIAGRAM_UPLOAD_FOLDER, exist_ok=True)

def detect_diagram_type(image_path):
    """Detect if diagram is hand-drawn or printed/downloaded using Gemini"""
    try:
        from PIL import Image
        img = Image.open(image_path)
        
        prompt = """Analyze this image and determine if it is:
1. HAND-DRAWN: Created by hand (pencil, pen, marker on paper, or digital drawing tablet with hand-drawn style)
2. PRINTED/DOWNLOADED: Computer-generated, printed from software, screenshot, or downloaded from internet

Look for these indicators:
- Hand-drawn: Irregular lines, natural imperfections, handwriting, paper texture, shadows from scanning, varying line thickness
- Printed: Perfect lines, computer fonts, uniform spacing, digital graphics, screenshot elements, watermarks

Respond in JSON format:
{
    "type": "hand-drawn" or "printed",
    "confidence": <0-100>,
    "reasoning": "brief explanation of why"
}"""
        
        if not GEMINI_API_KEY:
            return {'type': 'unknown', 'confidence': 0, 'reasoning': 'Gemini API not configured'}
        
        response = chat_model.generate_content([prompt, img])
        response_text = response.text.strip()
        
        # Parse JSON from response
        if '```json' in response_text:
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif '```' in response_text:
            response_text = response_text.split('```')[1].split('```')[0].strip()
        
        detection = json.loads(response_text)
        return detection
    except Exception as e:
        print(f"Diagram type detection error: {e}")
        return {'type': 'unknown', 'confidence': 0, 'reasoning': f'Error: {str(e)}'}

def evaluate_diagram_by_topic(student_image_path, topic, assignment_description, student_description):
    """Evaluate student diagram based on topic using Gemini"""
    try:
        from PIL import Image
        img = Image.open(student_image_path)
        
        prompt = f"""You are evaluating a student's diagram for the topic: "{topic}"

ASSIGNMENT DESCRIPTION:
{assignment_description if assignment_description else 'Create a diagram representing the topic'}

STUDENT'S DESCRIPTION:
{student_description}

Analyze the STUDENT'S DIAGRAM and evaluate:

1. TOPIC MATCH (30 points):
   - Does the diagram correctly represent the topic "{topic}"?
   - Are the key concepts of this topic present?
   - Is the diagram type appropriate (flowchart, body diagram, pie chart, graph, etc.)?

2. ACCURACY & COMPLETENESS (25 points):
   - Are all essential elements included?
   - Are labels correct and spelled properly?
   - Is the information accurate?

3. STRUCTURE & ORGANIZATION (20 points):
   - Is the diagram well-organized?
   - Are connections/relationships clear?
   - Is it easy to understand?

4. VISUAL QUALITY (15 points):
   - Is it neat and readable?
   - Are shapes and lines clear?
   - Is text legible?

5. TECHNICAL CORRECTNESS (10 points):
   - Are symbols/notations used correctly?
   - Does it follow standard conventions for this type of diagram?

IMPORTANT: Also determine if this diagram is:
- HAND-DRAWN: Drawn by hand (pencil, pen, markers, etc.)
- PRINTED/DIGITAL: Created using computer software or printed

Provide your evaluation in this JSON format:
{{
    "total_score": <0-100>,
    "topic_match_score": <0-30>,
    "accuracy_score": <0-25>,
    "structure_score": <0-20>,
    "visual_score": <0-15>,
    "technical_score": <0-10>,
    "is_correct": true/false,
    "is_hand_drawn": true/false,
    "drawing_type": "hand-drawn" or "printed/digital",
    "confidence_hand_drawn": <0-100>,
    "detailed_feedback": "detailed feedback about the diagram",
    "strengths": ["list", "of", "strengths"],
    "improvements": ["list", "of", "areas", "to", "improve"],
    "topic_match_explanation": "explain how well it matches the topic"
}}

Be thorough and fair in your evaluation."""
        
        if not GEMINI_API_KEY:
            return {
                'total_score': 0,
                'is_correct': False,
                'is_hand_drawn': False,
                'drawing_type': 'unknown',
                'detailed_feedback': 'Gemini API not configured'
            }
        
        response = chat_model.generate_content([prompt, img])
        response_text = response.text.strip()
        
        # Parse JSON from response
        if '```json' in response_text:
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif '```' in response_text:
            response_text = response_text.split('```')[1].split('```')[0].strip()
        
        evaluation = json.loads(response_text)
        return evaluation
    except Exception as e:
        print(f"Diagram evaluation error: {e}")
        import traceback
        traceback.print_exc()
        return {
            'total_score': 0,
            'is_correct': False,
            'is_hand_drawn': False,
            'drawing_type': 'unknown',
            'detailed_feedback': f'Error evaluating diagram: {str(e)}'
        }

@app.route('/guide/diagram-evaluation')
@login_required
@role_required('Guide')
def guide_diagram_evaluation():
    """Guide dashboard for diagram evaluation - redirects to main dashboard"""
    return redirect(url_for('guide_dashboard') + '#diagram-evaluation')

@app.route('/guide/diagram-evaluation/create', methods=['POST'])
@login_required
@role_required('Guide')
def create_diagram_assignment():
    """Create a new diagram assignment with comprehensive error handling"""
    try:
        # Get and validate form data
        course_id = safe_form_get('course_id')
        topic = safe_form_get('topic')
        description = safe_form_get('description', '').strip()
        
        # Validation checks
        if not course_id:
            return jsonify({
                'success': False, 
                'message': 'Please select a course',
                'field': 'course_id'
            }), 400
        
        if not topic:
            return jsonify({
                'success': False, 
                'message': 'Topic name is required',
                'field': 'topic'
            }), 400
        
        # Validate topic length
        if len(topic) < 3:
            return jsonify({
                'success': False, 
                'message': 'Topic name must be at least 3 characters long',
                'field': 'topic'
            }), 400
        
        if len(topic) > 200:
            return jsonify({
                'success': False, 
                'message': 'Topic name must be less than 200 characters',
                'field': 'topic'
            }), 400
        
        # Validate course_id is integer
        try:
            course_id_int = int(course_id)
        except (ValueError, TypeError):
            return jsonify({
                'success': False, 
                'message': 'Invalid course selected',
                'field': 'course_id'
            }), 400
        
        # Verify course belongs to guide
        course = db_session.query(Course).filter_by(id=course_id_int, guide_id=current_user.id).first()
        if not course:
            return jsonify({
                'success': False, 
                'message': 'Course not found or you do not have access to this course',
                'field': 'course_id'
            }), 403
        
        # Check for duplicate assignment (same topic in same course)
        existing = db_session.query(DiagramAssignment).filter_by(
            course_id=course_id_int,
            topic=topic.strip()
        ).first()
        
        if existing:
            return jsonify({
                'success': False, 
                'message': f'An assignment with topic "{topic}" already exists in this course',
                'field': 'topic'
            }), 409
        
        # Create assignment
        assignment = DiagramAssignment(
            course_id=course_id_int,
            topic=topic.strip(),
            description=description if description else None,
            created_by=current_user.id
        )
        
        db_session.add(assignment)
        db_session.commit()
        
        return jsonify({
            'success': True, 
            'message': 'Assignment created successfully!', 
            'assignment_id': assignment.id,
            'assignment': {
                'id': assignment.id,
                'topic': assignment.topic,
                'course_name': course.course_name
            }
        }), 201
        
    except ValueError as e:
        db_session.rollback()
        return jsonify({
            'success': False, 
            'message': f'Invalid input: {str(e)}',
            'field': 'general'
        }), 400
    except Exception as e:
        db_session.rollback()
        import traceback
        error_trace = traceback.format_exc()
        print(f"Error creating diagram assignment: {error_trace}")
        return jsonify({
            'success': False, 
            'message': f'An error occurred while creating the assignment. Please try again.',
            'field': 'general',
            'error': str(e) if app.debug else None
        }), 500

@app.route('/guide/diagram-evaluation/submissions')
@login_required
@role_required('Guide')
def view_diagram_submissions():
    """View all diagram submissions for guide's courses"""
    assignment_id = request.args.get('assignment_id')
    
    query = db_session.query(DiagramSubmission).join(DiagramAssignment).join(Course).filter(
        Course.guide_id == current_user.id
    )
    
    if assignment_id:
        query = query.filter(DiagramSubmission.assignment_id == int(assignment_id))
    
    submissions = query.order_by(DiagramSubmission.submitted_at.desc()).all()
    
    # Get assignment details
    assignments = db_session.query(DiagramAssignment).join(Course).filter(
        Course.guide_id == current_user.id
    ).all()
    
    return jsonify({
        'success': True,
        'submissions': [{
            'id': s.id,
            'assignment_id': s.assignment_id,
            'learner_name': s.learner.full_name or s.learner.username,
            'learner_email': s.learner.email,
            'assignment_topic': s.assignment.topic,
            'diagram_path': s.diagram_path,
            'total_score': s.total_score,
            'final_marks': s.final_marks,
            'is_hand_drawn': s.is_hand_drawn,
            'marks_status': s.marks_status,
            'submitted_at': s.submitted_at.isoformat() if s.submitted_at else None
        } for s in submissions],
        'assignments': [{'id': a.id, 'topic': a.topic, 'course_name': a.course.course_name} for a in assignments]
    })

@app.route('/guide/diagram-evaluation/submission/<int:submission_id>')
@login_required
@role_required('Guide')
def view_diagram_submission_detail(submission_id):
    """View detailed evaluation of a specific submission"""
    submission = db_session.query(DiagramSubmission).join(DiagramAssignment).join(Course).filter(
        DiagramSubmission.id == submission_id,
        Course.guide_id == current_user.id
    ).first()
    
    if not submission:
        return jsonify({'success': False, 'message': 'Submission not found'}), 404
    
    # Parse evaluation data
    evaluation_data = {}
    if submission.evaluation_data:
        try:
            evaluation_data = json.loads(submission.evaluation_data)
        except:
            pass
    
    return jsonify({
        'success': True,
        'submission': {
            'id': submission.id,
            'learner_name': submission.learner.full_name or submission.learner.username,
            'learner_email': submission.learner.email,
            'assignment_topic': submission.assignment.topic,
            'diagram_path': submission.diagram_path,
            'description': submission.description,
            'total_score': submission.total_score,
            'topic_match_score': submission.topic_match_score,
            'accuracy_score': submission.accuracy_score,
            'structure_score': submission.structure_score,
            'visual_score': submission.visual_score,
            'technical_score': submission.technical_score,
            'is_hand_drawn': submission.is_hand_drawn,
            'drawing_type': submission.drawing_type,
            'confidence_hand_drawn': submission.confidence_hand_drawn,
            'final_marks': submission.final_marks,
            'marks_status': submission.marks_status,
            'submitted_at': submission.submitted_at.isoformat(),
            'evaluation_data': evaluation_data
        }
    })

@app.route('/guide/diagram-evaluation/download-pdf/<int:assignment_id>')
@login_required
@role_required('Guide')
def download_diagram_pdf(assignment_id):
    """Download PDF report of all submissions for an assignment"""
    assignment = db_session.query(DiagramAssignment).join(Course).filter(
        DiagramAssignment.id == assignment_id,
        Course.guide_id == current_user.id
    ).first()
    
    if not assignment:
        return jsonify({'success': False, 'message': 'Assignment not found'}), 404
    
    submissions = db_session.query(DiagramSubmission).filter_by(
        assignment_id=assignment_id
    ).order_by(DiagramSubmission.submitted_at.desc()).all()
    
    # Generate PDF
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title = Paragraph(f"Diagram Evaluation Report: {assignment.topic}", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 0.2*inch))
    
    # Assignment info
    info_data = [
        ['Course:', assignment.course.course_name],
        ['Topic:', assignment.topic],
        ['Description:', assignment.description or 'N/A'],
        ['Created:', assignment.created_at.strftime('%Y-%m-%d %H:%M')],
        ['Total Submissions:', str(len(submissions))]
    ]
    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.grey),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (1, 0), (1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 0.3*inch))
    
    # Submissions table
    for idx, submission in enumerate(submissions, 1):
        learner = submission.learner
        elements.append(Paragraph(f"Submission #{idx}: {learner.full_name or learner.username}", styles['Heading2']))
        elements.append(Spacer(1, 0.1*inch))
        
        sub_data = [
            ['Learner:', learner.full_name or learner.username],
            ['Email:', learner.email],
            ['Submitted:', submission.submitted_at.strftime('%Y-%m-%d %H:%M')],
            ['Total Score:', f"{submission.total_score:.1f}/100"],
            ['Topic Match:', f"{submission.topic_match_score:.1f}/30"],
            ['Accuracy:', f"{submission.accuracy_score:.1f}/25"],
            ['Structure:', f"{submission.structure_score:.1f}/20"],
            ['Visual Quality:', f"{submission.visual_score:.1f}/15"],
            ['Technical:', f"{submission.technical_score:.1f}/10"],
            ['Drawing Type:', 'Hand-drawn' if submission.is_hand_drawn else 'Printed/Digital'],
            ['Final Marks:', f"{submission.final_marks:.1f}/100"],
            ['Status:', submission.marks_status or 'Evaluated']
        ]
        
        sub_table = Table(sub_data, colWidths=[2*inch, 4*inch])
        sub_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(sub_table)
        
        if submission.description:
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph(f"<b>Description:</b> {submission.description}", styles['Normal']))
        
        if submission.evaluation_data:
            try:
                eval_data = json.loads(submission.evaluation_data)
                if eval_data.get('detailed_feedback'):
                    elements.append(Spacer(1, 0.1*inch))
                    elements.append(Paragraph(f"<b>Feedback:</b> {eval_data['detailed_feedback']}", styles['Normal']))
            except:
                pass
        
        if idx < len(submissions):
            elements.append(PageBreak())
    
    doc.build(elements)
    buffer.seek(0)
    
    return send_file(buffer, mimetype='application/pdf', 
                    as_attachment=True, 
                    download_name=f'diagram_evaluation_{assignment.topic}_{datetime.now().strftime("%Y%m%d")}.pdf')

@app.route('/learner/diagram-evaluation')
@login_required
@role_required('Learner')
def learner_diagram_evaluation():
    """Learner dashboard for diagram evaluation - redirects to main dashboard"""
    return redirect(url_for('learner_dashboard') + '#diagram-evaluation')

@app.route('/learner/diagram-evaluation/submit', methods=['POST'])
@login_required
@role_required('Learner')
def submit_diagram():
    """Submit diagram for evaluation"""
    try:
        assignment_id = safe_form_get('assignment_id')
        description = safe_form_get('description', '')
        
        if 'diagram' not in request.files:
            return jsonify({'success': False, 'message': 'No file uploaded'})
        
        file = request.files['diagram']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'No file selected'})
        
        # Verify assignment exists and learner is enrolled
        assignment = db_session.query(DiagramAssignment).join(Course).join(CourseLearner).filter(
            DiagramAssignment.id == int(assignment_id),
            CourseLearner.learner_id == current_user.id
        ).first()
        
        if not assignment:
            return jsonify({'success': False, 'message': 'Assignment not found or access denied'}), 404
        
        # Save file
        filename = secure_filename(f"{current_user.username}_{datetime.now(timezone.utc).timestamp()}_{file.filename}")
        filepath = os.path.join(DIAGRAM_UPLOAD_FOLDER, filename)
        file.save(filepath)
        
        # Detect diagram type
        diagram_type_detection = detect_diagram_type(filepath)
        
        # Evaluate diagram
        evaluation = evaluate_diagram_by_topic(
            filepath,
            assignment.topic,
            assignment.description or '',
            description
        )
        
        # Auto-assign marks
        is_hand_drawn = evaluation.get('is_hand_drawn', False)
        if is_hand_drawn:
            final_marks = evaluation.get('total_score', 0)
            marks_status = 'Hand-drawn - Evaluated'
        else:
            final_marks = 0
            marks_status = 'Printed/Digital - Not Accepted'
        
        # Create submission
        submission = DiagramSubmission(
            assignment_id=int(assignment_id),
            learner_id=current_user.id,
            diagram_path=filepath,
            description=description,
            evaluation_data=json.dumps(evaluation),
            total_score=evaluation.get('total_score', 0),
            topic_match_score=evaluation.get('topic_match_score', 0),
            accuracy_score=evaluation.get('accuracy_score', 0),
            structure_score=evaluation.get('structure_score', 0),
            visual_score=evaluation.get('visual_score', 0),
            technical_score=evaluation.get('technical_score', 0),
            is_hand_drawn=is_hand_drawn,
            drawing_type=evaluation.get('drawing_type', 'unknown'),
            confidence_hand_drawn=evaluation.get('confidence_hand_drawn', 0),
            final_marks=final_marks,
            marks_status=marks_status
        )
        
        db_session.add(submission)
        db_session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Diagram submitted and evaluated successfully',
            'submission_id': submission.id,
            'evaluation': evaluation
        })
    except Exception as e:
        db_session.rollback()
        print(f"Submit diagram error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 400

@app.route('/diagram-uploads/<path:filename>')
@login_required
def serve_diagram(filename):
    """Serve diagram images"""
    return send_from_directory(DIAGRAM_UPLOAD_FOLDER, filename)

@app.route('/guide/plagiarism/unignore', methods=['POST'])
@login_required
@role_required('Guide')
def unignore_plagiarism():
    """Revert ignored plagiarism"""
    try:
        data = request.get_json()
        match_id = data.get('match_id')
        
        match = db_session.query(PlagiarismMatch).get(match_id)
        if not match:
            return jsonify({'success': False, 'message': 'Match not found'})
        
        match.status = 'pending'
        match.reviewed_by = None
        match.reviewed_at = None
        
        db_session.commit()
        
        return jsonify({'success': True})
        
    except Exception as e:
        db_session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/mcq-test')
@login_required
@role_required('Learner')
def mcq_test_page():
    """Render the secure MCQ test page"""
    return render_template('mcq_test.html', csrf_token=generate_csrf())

if __name__ == '__main__':
    # Run via SocketIO server (threading async_mode) to improve compatibility on macOS
    # Use 0.0.0.0 to allow access from other devices on the same network (like your phone)
    socketio.run(app, debug=True, host='0.0.0.0', port=5000)
